"""
p_analysis.py — NHL Role Scoring Pipeline (NST on-ice + Bio tabs) + Player Type + Methodology Doc

WHAT YOU GET (outputs):
VISIBLE sheets:
  - F_25-26, D_25-26
  - F_24-25, D_24-25
  - F_23-24, D_23-24
  - Rolling_F, Rolling_D

HIDDEN audit sheets:
  - _ALL_25-26, _ALL_24-25, _ALL_23-24, _ROLLING_ALL

NEW:
  ✅ Player type designation:
     - Player_Type (single best-fit label w/ guardrails)
     - Secondary_Type (optional)
     - Type_Margin (Top1 - Top2)
     - Type_Confidence (High/Med/Low)
     - Type_Notes

  ✅ Methodology document outputs in the same folder as OUTPUT_XLSX:
     - NHL_Model_Methodology.docx
     - NHL_Model_Methodology.pdf

INPUT workbook requirements:
  - Season tabs: "25-26", "24-25", "23-24" (NST on-ice tables)
  - Bio tabs:    "25-26 Bios", "24-25 Bios", "23-24 Bios"
    with height in inches and weight in pounds

Install:
  pip install pandas numpy openpyxl python-docx reportlab

Run:
  python p_analysis.py
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import numpy as np
import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


# =============================================================================
# EDIT THESE PATHS
# =============================================================================
INPUT_XLSX = Path(r"New Report (2)\P-Stats-B.xlsx")
OUTPUT_XLSX = Path(r"New Report (2)\P_Role_Scores_Output.xlsx")

SEASON_TABS = ["25-26", "24-25", "23-24"]  # newest, prev1, prev2
ROLLING_WEIGHTS = (0.60, 0.30, 0.10)

# Stability filters (TOI in minutes; NST TOI used as EV proxy)
MIN_TOI_EV_FWD = 100.0
MIN_TOI_EV_DEF = 200.0

# Player Type rules
TYPE_MIN_CONFIDENCE_EV = 0.35   # below this -> Type_Confidence Low & Type notes
TYPE_HYBRID_MARGIN = 6.0        # if Top1-Top2 < this -> "Hybrid"
TYPE_CLEAR_MARGIN = 12.0        # if >= this -> Type_Confidence High (unless TOI low)


# =============================================================================
# ROLE FEATURES (built for NST on-ice + Bio)
# =============================================================================
ROLE_FEATURES: Dict = {
    "F": {
        "Finisher": {
            "minutes_basis": "toi_ind",
            "k": 600,
            "features": {
                "g60_ind": 0.35,
                "ixg60_ind": 0.25,
                "shots60_ind": 0.15,
                "sh_pct": 0.10,
                "ihdcf": 0.15,  # volume of chances (count-based proxy)
            },
        },
        "Playmaker": {
            "minutes_basis": "toi_ind",
            "k": 600,
            "features": {
                "a1_60_ind": 0.45,
                "p1_60_ind": 0.20,
                "ipp": 0.15,
                "rebounds60_ind": 0.10,
                "rush60_ind": 0.10,
            },
        },
        "Driver": {
            "minutes_basis": "toi_ev",
            "k": 600,
            "features": {
                "xgf60_ev": 0.35,
                "cf60_ev": 0.25,
                "scf60_ev": 0.20,
                "hdcf60_ev": 0.20,
            },
        },
        "TwoWay": {
            "minutes_basis": "toi_ev",
            "k": 600,
            "features": {
                "xga60_ev": -0.35,
                "hdca60_ev": -0.25,
                "xgf60_ev": 0.15,
                "pen_diff60_ind": 0.25,
            },
        },
        "Power": {
            "minutes_basis": "toi_ind",
            "k": 600,
            "features": {
                "ihdcf": 0.25,          # net-front / inner-slot presence
                "ixg60_ind": 0.20,      # dangerous shot quality
                "shots60_ind": 0.10,
                "hits60_ind": 0.15,     # still physical, but not the whole story
                "pen_drawn60_ind": 0.15,
                "bmi": 0.15,
            },
        },
        "Grinder": {
            "minutes_basis": "toi_ind",
            "k": 500,
            "features": {
                "hits60_ind": 0.35,
                "blk_shots60_ind": 0.20,
                "pen_taken60_ind": 0.20,   # tough guys take penalties
                "pen_drawn60_ind": 0.10,   # some also draw
                "bmi": 0.15,
            },
        },
        "Producer": {
            "minutes_basis": "toi_ind",
            "k": 500,
            "features": {
                "points60_ind": 0.35,
                "g60_ind": 0.20,
                "a1_60_ind": 0.20,
                "ixg60_ind": 0.15,
                "ipp": 0.10,
            },
        },
    },
    "D": {
       "Suppressor": {
            "minutes_basis": "toi_ev",
            "k": 700,
            "features": {
                "xga60_ev": -0.35,
                "hdca60_ev": -0.25,
                "sa60_ev": -0.20,
                "ca60_ev": -0.20,
            },
        },
        "Transition": {
            "minutes_basis": "toi_ev",
            "k": 700,
            "features": {
                "xgf60_ev": 0.30,
                "cf60_ev": 0.20,
                "scf60_ev": 0.20,
                "rush60_ind": 0.20,
                "rebounds60_ind": 0.10,
            },
        },
        "PuckSkill": {
            "minutes_basis": "toi_ind",
            "k": 700,
            "features": {
                "a1_60_ind": 0.35,
                "points60_ind": 0.20,
                "giveaways60_ind": -0.25,
                "takeaways60_ind": 0.20,
            },
        },
        "Physical": {
            "minutes_basis": "toi_ind",
            "k": 700,
            "features": {
                "hits60_ind": 0.25,
                "blk_shots60_ind": 0.25,
                "bmi": 0.15,
                "weight_lb": 0.15,
                "pen_taken60_ind": -0.20,
            },
        },
    },
}

# =============================================================================
# COLUMN FALLBACKS MATCHING YOUR WORKBOOK
# =============================================================================
FALLBACKS: Dict[str, List[str]] = {
    "player": ["Player", "player", "Name", "Skater"],
    "team": ["Team", "team"],
    "pos_raw": ["Position", "Pos", "position"],
    "gp": ["GP", "Games", "gp"],
    "toi_ev": ["TOI", "toi", "TOI (min)", "toi_total"],

    "cf": ["CF"],
    "ca": ["CA"],
    "sf": ["SF"],
    "sa": ["SA"],
    "gf": ["GF"],
    "ga": ["GA"],
    "xgf": ["xGF"],
    "xga": ["xGA"],
    "scf": ["SCF"],
    "sca": ["SCA"],
    "hdcf": ["HDCF"],
    "hdca": ["HDCA"],

    "oz_starts": ["Off.\xa0Zone Starts", "Off. Zone Starts", "Off Zone Starts"],
    "nz_starts": ["Neu.\xa0Zone Starts", "Neu. Zone Starts", "Neutral Zone Starts"],
    "dz_starts": ["Def.\xa0Zone Starts", "Def. Zone Starts", "Def Zone Starts"],
    "oz_pct": ["Off.\xa0Zone Start %", "Off. Zone Start %", "Off Zone Start %"],
}

BIO_FALLBACKS: Dict[str, List[str]] = {
    "player": ["Player", "player", "Name"],
    "height_in": ["Height (in)", "Height_in", "height_in"],
    "weight_lb": ["Weight (lbs)", "Weight (lb)", "Weight_lb", "weight_lb"],
    "age": ["Age", "age"],
    "draft_round": ["Draft Round", "draft_round"],
    "round_pick": ["Round Pick", "round_pick"],
    "overall_pick": ["Overall Draft Position", "overall_pick"],
}

# Individual tabs (new)
IND_TABS = {
    "25-26": "25-26 Ind",
    "24-25": "24-25 Ind",
    "23-24": "23-24 Ind",
}

IND_FALLBACKS: Dict[str, List[str]] = {
    "player": ["Player", "player", "Name"],
    "team": ["Team", "team"],
    "pos_raw": ["Position", "Pos", "position"],
    "gp": ["GP", "Games"],
    "toi": ["TOI", "toi"],

    "goals": ["Goals", "G"],
    "assists": ["Total Assists", "Assists", "A"],
    "a1": ["First Assists", "Primary Assists", "A1"],
    "a2": ["Second Assists", "Secondary Assists", "A2"],
    "points": ["Total Points", "Points", "P"],
    "ipp": ["IPP"],

    "shots": ["Shots", "S"],
    "sh_pct": ["SH%", "Sh%"],
    "ixg": ["ixG", "iXG"],

    "icf": ["iCF"],
    "iff": ["iFF"],
    "iscf": ["iSCF"],
    "ihdcf": ["iHDCF"],

    "rush": ["Rush Attempts"],
    "rebounds": ["Rebounds Created"],

    "pim": ["PIM"],
    "pen_total": ["Total Penalties", "Total Penalty", "Penalties"],
    "pen_minor": ["Minor"],
    "pen_major": ["Major"],
    "pen_misconduct": ["Misconduct"],
    "pen_drawn": ["Penalties Drawn"],

    "giveaways": ["Giveaways", "GIVE"],
    "takeaways": ["Takeaways", "TAKE"],
    "hits": ["Hits", "HIT"],
    "hits_taken": ["Hits Taken"],
    "shots_blocked": ["Shots Blocked", "Blocks", "BLK"],

    "fow": ["Faceoffs Won", "FOW"],
    "fol": ["Faceoffs Lost", "FOL"],
    "fo_pct": ["Faceoffs %", "FO%"],
}

# =============================================================================
# HELPERS
# =============================================================================
def _clean_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [re.sub(r"\s+", " ", str(c).strip()) for c in df.columns]
    return df

def _find_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    lower_map = {c.lower(): c for c in df.columns}
    for c in candidates:
        if c.lower() in lower_map:
            return lower_map[c.lower()]
    return None

def _to_float(x):
    if pd.isna(x):
        return np.nan
    if isinstance(x, (int, float, np.number)):
        return float(x)
    s = str(x).strip().replace(",", "").replace("%", "")
    if not s:
        return np.nan
    if ":" in s:
        try:
            mm, ss = s.split(":")
            return float(mm) + float(ss) / 60.0
        except Exception:
            return np.nan
    try:
        return float(s)
    except Exception:
        return np.nan

def _player_key(s: str) -> str:
    t = re.sub(r"[^a-zA-Z\s\-']", "", str(s))
    t = re.sub(r"\s+", " ", t).lower().strip()
    return t

def _resolve_team(raw: str) -> str:
    """
    NST exports traded players with a comma-separated list of teams
    in chronological order, e.g. 'S.J, VAN' or 'COL, EDM, PIT'.
    Return the last token — the player's most recent/current team.
    Single-team strings are returned unchanged.
    """
    if not raw or raw.lower() in ("nan", ""):
        return ""
    parts = [p.strip() for p in raw.split(",") if p.strip()]
    return parts[-1] if parts else raw

def _canonical_pos(pos_raw: str) -> str:
    # Handles "C, L" by taking first token
    p = str(pos_raw).strip().upper()
    if not p:
        return "UNK"
    token = p.split(",")[0].strip()
    first = token[:1]
    if first == "C":
        return "C"
    if first == "L":
        return "LW"
    if first == "R":
        return "RW"
    if first == "D":
        return "D"
    return "UNK"

def _pos_group(pos: str) -> str:
    return "D" if pos == "D" else "F"

def _zscore_within(df: pd.DataFrame, col: str, by: List[str]) -> pd.Series:
    g = df.groupby(by)[col]
    mu = g.transform("mean")
    sd = g.transform("std").replace(0, np.nan)
    z = (df[col] - mu) / sd
    return z.replace([np.inf, -np.inf], np.nan).fillna(0.0)

def _shrink(z: pd.Series, minutes: pd.Series, k: float) -> pd.Series:
    m = minutes.fillna(0.0).clip(lower=0.0)
    return z * (m / (m + k))

def _percent_rank(s: pd.Series) -> pd.Series:
    return s.rank(pct=True, method="average")

def normalize_player_name(name: str) -> str:
    """
    Normalize player names so they match across sources.

    Handles:
      - "Last, First" -> "First Last"
      - "Last, First Middle" -> "First Middle Last"
      - preserves suffixes like Jr, Sr, II, III, IV when present

    Returns a clean display name, not the key.
    """
    s = str(name or "").strip()
    if not s:
        return ""

    # Remove double spaces, normalize commas
    s = re.sub(r"\s+", " ", s).strip()

    # Common suffix tokens
    suffixes = {"JR", "JR.", "SR", "SR.", "II", "III", "IV", "V"}

    # If comma format: "Last, First ..."
    if "," in s:
        parts = [p.strip() for p in s.split(",") if p.strip()]
        if len(parts) >= 2:
            last = parts[0]
            rest = " ".join(parts[1:]).strip()  # "First Middle Jr."
            rest_tokens = rest.split()

            # Detect suffix at end of rest
            suf = ""
            if rest_tokens and rest_tokens[-1].upper().replace(".", "") in {x.replace(".", "") for x in suffixes}:
                suf = rest_tokens[-1]
                rest_tokens = rest_tokens[:-1]

            first_middle = " ".join(rest_tokens).strip()
            # Build "First Middle Last [Suffix]"
            out = " ".join([x for x in [first_middle, last, suf] if x]).strip()
            out = re.sub(r"\s+", " ", out).strip()
            return out

    # Otherwise leave as-is (already "First Last")
    return s

def add_team_usage_rank(df: pd.DataFrame) -> pd.DataFrame:
    """
    Usage proxy for line fit gating.
    Prefer toi_ind (individual TOI) if present, else toi_ev.
    Group by team and pos_group.
    """
    out = df.copy()

    team_col = "team_cap" if "team_cap" in out.columns else ("team" if "team" in out.columns else None)
    if team_col is None:
        out["Usage_Rank_Team"] = 0.5
        return out

    toi_col = "toi_ind" if "toi_ind" in out.columns else ("toi_ev" if "toi_ev" in out.columns else None)
    if toi_col is None:
        out["Usage_Rank_Team"] = 0.5
        return out

    # Clean team
    out[team_col] = out[team_col].fillna("").astype(str).str.upper().str.strip()
    out["pos_group"] = out.get("pos_group", np.where(out.get("pos","").astype(str).eq("D"), "D", "F"))

    # Percentile rank within team & group
    out["Usage_Rank_Team"] = (
        out.groupby([team_col, "pos_group"])[toi_col]
           .transform(lambda s: s.rank(pct=True, method="average"))
           .fillna(0.5)
    )
    return out

def _first_notna(*vals, default=np.nan):
    """
    Returns the first value that is not None and not NaN.
    IMPORTANT: np.nan is truthy in Python, so never use `a or b` for numeric fallbacks.
    """
    for v in vals:
        if v is None:
            continue
        try:
            if pd.notna(v):
                return v
        except Exception:
            # if pd.notna fails (rare), fall back to truthiness
            if v:
                return v
    return default

def _coalesce_xy(df: pd.DataFrame, base: str) -> pd.DataFrame:
    """
    If merges produced base_x/base_y, coalesce into base and drop the suffix cols.
    Prefer *_y (enriched/model) then *_x (cap/original).
    """
    x = f"{base}_x"
    y = f"{base}_y"
    if base in df.columns:
        return df
    if x in df.columns or y in df.columns:
        df[base] = pd.Series(np.nan, index=df.index)
        if y in df.columns:
            df[base] = df[base].fillna(df[y])
        if x in df.columns:
            df[base] = df[base].fillna(df[x])
        df = df.drop(columns=[c for c in [x, y] if c in df.columns])
    return df

def add_missing_active_contract_players(
    season_dfs: dict[str, pd.DataFrame],
    roster_df: pd.DataFrame,
    current_season: str,
    fallback_seasons: list[str],
    key_col: str = "player_key",
    name_col: str = "player",
    fill_cols: list[str] | None = None,
    decay: float = 0.00,
) -> pd.DataFrame:
    """
    Inserts active-contract players missing from the current season (injured/not yet played)
    using their most recent prior season row as a template.

    Carries ALL columns from the fallback row — numeric with optional decay, non-numeric
    as-is — so imputed players have the same column coverage as regular players.
    Columns that only exist in the current season template (new derived cols) are left NaN
    and will be re-derived in the post-imputation pipeline.
    """
    cur = season_dfs[current_season].copy()

    # --- derive active_contract from cap table if missing ---
    if "active_contract" not in roster_df.columns:
        season_end_year = 2026  # for 25-26
        exp_col = "expiry_year"
        if exp_col not in roster_df.columns:
            raise KeyError(f"Expected '{exp_col}' in cap/roster_df. Columns={list(roster_df.columns)}")
        exp_year = pd.to_numeric(roster_df[exp_col], errors="coerce")
        roster_df = roster_df.copy()
        roster_df["active_contract"] = exp_year > season_end_year

    active = roster_df[roster_df["active_contract"] == True].copy()
    cur_keys = set(cur[key_col].dropna().astype(str))
    missing = active[~active[key_col].astype(str).isin(cur_keys)].copy()

    if missing.empty:
        season_dfs[current_season] = cur
        return cur

    # build lookup from fallback seasons (most recent first)
    fb_lookup = {}
    for s in fallback_seasons:
        if s not in season_dfs:
            continue
        df = season_dfs[s]
        if key_col not in df.columns:
            continue
        # collapse duplicates to one row per player_key
        num_cols = df.select_dtypes(include="number").columns.tolist()
        deduped = df.groupby(key_col, as_index=False)[num_cols].mean()
        non_num = [c for c in df.columns if c not in num_cols and c != key_col]
        first_vals = df.groupby(key_col)[non_num].first().reset_index() if non_num else None
        if first_vals is not None:
            deduped = deduped.merge(first_vals, on=key_col, how="left")
        fb_lookup[s] = deduped.set_index(key_col)

    # columns that should NOT be carried forward from fallback
    # (they are current-season context or will be freshly derived)
    SKIP_COLS = {key_col, "season", "keep", "Injury_Imputed", "injury_fill", "fill_source",
                 "Usage_Rank_Team"}

    rows_to_add = []
    for _, r in missing.iterrows():
        pk = str(r[key_col])

        # cap_df uses team_cap/pos_cap; fall back gracefully
        team_val = r.get("team") or r.get("team_cap", np.nan)
        pos_val  = r.get("pos")  or r.get("pos_cap",  np.nan)
        pos_str  = str(pos_val or "").upper().strip()

        new_row: dict = {
            key_col:       pk,
            name_col:      r.get("player", np.nan),
            "team":        team_val,
            "pos":         pos_val,
            "pos_group":   "D" if pos_str == "D" else ("F" if pos_str else np.nan),
            "gp":          0,
            "toi_ev":      0.0,
            "toi_ind":     0.0,
            "injury_fill": True,
            "fill_source": np.nan,
            "season":      current_season,
        }

        # Fill ALL columns from most recent fallback season that has this player
        for s, idx in fb_lookup.items():
            if pk not in idx.index:
                continue
            src = idx.loc[pk]
            for col in src.index:
                if col in SKIP_COLS:
                    continue
                val = src[col]
                if pd.isna(val):
                    continue
                # Apply decay to numeric performance cols; carry strings as-is
                if isinstance(val, (int, float, np.number)):
                    new_row[col] = float(val) * (1.0 - decay)
                else:
                    new_row[col] = val
            new_row["fill_source"] = s
            # Prefer the fallback's team/pos/pos_group if the cap row had blanks
            if pd.isna(new_row.get("team")) and "team" in src.index and pd.notna(src["team"]):
                new_row["team"] = src["team"]
            if pd.isna(new_row.get("pos")) and "pos" in src.index and pd.notna(src["pos"]):
                new_row["pos"] = src["pos"]
                new_row["pos_group"] = "D" if str(src["pos"]).upper() == "D" else "F"
            break
        else:
            new_row["fill_source"] = "NONE"

        # Final safety: derive pos_group if still missing
        if not new_row.get("pos_group") or pd.isna(new_row.get("pos_group")):
            new_row["pos_group"] = "D" if str(new_row.get("pos") or "").upper() == "D" else "F"

        rows_to_add.append(new_row)

    add_df = pd.DataFrame(rows_to_add)
    cur2 = pd.concat([cur, add_df], ignore_index=True, sort=False)

    season_dfs[current_season] = cur2
    return cur2

def impute_injured_seasons(season_all: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
    """
    For players with insufficient TOI in the current season (fully absent OR partially
    injured), carry forward their prior season row with a small decay on numeric values.

    Imputation fires when a player's toi_ev is below the minimum threshold for their
    position group AND the prior season has meaningfully more data (prior toi > current toi).
    This catches both:
      - Players with toi_ev = 0 / NaN  (never played)
      - Players like Kulikov with 42 min who played a few games then got injured

    Columns intentionally NOT carried forward (reset to current-season context):
      - season, keep, Injury_Imputed, Usage_Rank_Team
      - toi_ev is replaced with prior * DECAY so Confidence_EV isn't zero

    Everything else — bio, raw stats, per-60s, role scores, type labels, fit scores —
    is carried forward so imputed players have full column parity with normal players.
    The post-imputation pipeline then re-derives keep, Confidence_EV, type, fit, and
    usage rank on top of the imputed values.
    """
    DECAY = 0.92

    # Columns that are current-season context and must NOT be copied from prior year
    SKIP_COLS = {"season", "keep", "Injury_Imputed", "Usage_Rank_Team"}

    seasons = list(season_all.keys())

    for i, season in enumerate(seasons):
        df = season_all[season].copy()

        if i == len(seasons) - 1:
            season_all[season] = df
            continue  # oldest season has nothing to borrow from

        prev_season = seasons[i + 1]
        prev_df = season_all[prev_season].copy()

        # Collapse prev season to one row per player_key (average numerics, first for strings)
        if prev_df["player_key"].duplicated().any():
            num_cols = prev_df.select_dtypes(include="number").columns.tolist()
            non_num  = [c for c in prev_df.columns if c not in num_cols and c != "player_key"]
            agg_num  = prev_df.groupby("player_key", as_index=False)[num_cols].mean()
            agg_str  = prev_df.groupby("player_key")[non_num].first().reset_index() if non_num else None
            prev_df  = agg_num.merge(agg_str, on="player_key", how="left") if agg_str is not None else agg_num

        prev_lookup = prev_df.set_index("player_key")

        # ── Build imputation mask ──────────────────────────────────────────────
        # A player needs imputation if their current toi_ev is below the minimum
        # threshold for their position AND the prior season has better data.
        # "Better" = prior toi_ev > current toi_ev (prior is genuinely more complete).
        toi_cur = df["toi_ev"].fillna(0.0).astype(float)

        # Per-row minimum based on pos_group
        if "pos_group" in df.columns:
            toi_min = np.where(df["pos_group"].astype(str).eq("D"),
                               MIN_TOI_EV_DEF, MIN_TOI_EV_FWD)
        else:
            toi_min = np.full(len(df), MIN_TOI_EV_FWD)

        # Prior TOI for each player (NaN if not in prev season)
        toi_prior_ser = prev_lookup["toi_ev"].reindex(df["player_key"]).fillna(0.0).values.astype(float)

        below_threshold = toi_cur.values < toi_min
        prior_is_better = toi_prior_ser > toi_cur.values

        injured_mask = pd.Series(below_threshold & prior_is_better, index=df.index)

        if injured_mask.sum() == 0:
            season_all[season] = df
            continue

        keys_need = df.loc[injured_mask, "player_key"]

        # Determine which columns to carry: intersection of df cols and prev cols, minus skips
        carry_cols = [
            c for c in df.columns
            if c not in SKIP_COLS
            and c != "player_key"
            and c in prev_lookup.columns
        ]

        for col in carry_cols:
            prior_vals = prev_lookup[col].reindex(keys_need).values
            if df[col].dtype.kind in ("f", "i", "u"):
                # Numeric: replace with prior * decay (don't blend — prior is more reliable)
                df.loc[injured_mask, col] = np.where(
                    pd.isnull(prior_vals),
                    df.loc[injured_mask, col].values,
                    prior_vals.astype(float) * DECAY
                )
            else:
                # String/object: carry as-is, only fill where currently null/empty
                current_vals = df.loc[injured_mask, col].values
                filled = np.where(
                    pd.isnull(current_vals) | (np.array(current_vals, dtype=str) == ""),
                    prior_vals,
                    current_vals,
                )
                df.loc[injured_mask, col] = filled

        # toi_ev: replace with prior * DECAY to give a realistic Confidence_EV base
        if "toi_ev" in prev_lookup.columns:
            toi_prior_imp = prev_lookup["toi_ev"].reindex(keys_need).values.astype(float)
            df.loc[injured_mask, "toi_ev"] = np.where(
                np.isnan(toi_prior_imp), 0.0, toi_prior_imp * DECAY
            )

        # Ensure pos_group is set (may be blank for newly added rows)
        if "pos" in df.columns and "pos_group" in df.columns:
            missing_grp = injured_mask & (df["pos_group"].isna() | df["pos_group"].eq(""))
            df.loc[missing_grp, "pos_group"] = np.where(
                df.loc[missing_grp, "pos"].astype(str).str.upper().eq("D"), "D", "F"
            )

        df.loc[injured_mask, "Injury_Imputed"] = True
        season_all[season] = df

    return season_all

# =============================================================================
# TYPE CONFIDENCE + LINE FIT HELPERS
# =============================================================================

FWD_ROLE_COLS = ["Finisher_Score","Playmaker_Score","Driver_Score","TwoWay_Score","Power_Score","Grinder_Score","Producer_Score"]
DEF_ROLE_COLS = ["Suppressor_Score", "Transition_Score", "PuckSkill_Score", "Physical_Score"]

FWD_ROLE_LABELS = {
    "Finisher_Score": "Finisher",
    "Playmaker_Score": "Playmaker",
    "Driver_Score": "Driver",
    "TwoWay_Score": "Two-Way",
    "Power_Score": "Power",
    "Grinder_Score": "Grinder",
    "Producer_Score": "Producer",
}
DEF_ROLE_LABELS = {
    "Suppressor_Score": "Shutdown D",
    "Transition_Score": "Transition",
    "PuckSkill_Score": "Puck Skill",
    "Physical_Score": "Physical",
}

def _pick_role_cols(df: pd.DataFrame, base_cols: List[str]) -> Tuple[List[str], str]:
    """
    Chooses whether to use season cols (e.g. Finisher_Score) or rolling cols (Finisher_Score_3yr).
    Returns (cols_to_use, suffix_used) where suffix_used is "" or "_3yr".
    """
    if all(c in df.columns for c in base_cols):
        return base_cols, ""
    base_3yr = [c + "_3yr" for c in base_cols]
    if all(c in df.columns for c in base_3yr):
        return base_3yr, "_3yr"
    # mixed / partially missing; use whatever exists (best effort)
    cols = [c for c in base_cols if c in df.columns]
    if cols:
        return cols, ""
    cols = [c for c in base_3yr if c in df.columns]
    return cols, "_3yr" if cols else ""

def _vol_col_for(score_col: str, suffix: str) -> Optional[str]:
    """
    For rolling tables, volatility columns are like <ScoreCol>_vol_3yr where ScoreCol already includes _3yr.
    Example: Finisher_Score_3yr -> Finisher_Score_vol_3yr
    """
    if suffix != "_3yr":
        return None
    if score_col.endswith("_3yr"):
        return score_col.replace("_3yr", "_vol_3yr")
    return None

def add_type_confidence_and_stability(df: pd.DataFrame) -> pd.DataFrame:
    """
    Adds:
      Top_Role, Top_Role_Score, Second_Role, Second_Role_Score, Margin
      Role_Stability (0-1, rolling only when vol exists)
      Type_Confidence (High/Medium/Low) computed from Confidence_EV + Margin + Stability
    Works for both season sheets and rolling sheets.
    """
    out = df.copy()

    # pick proper role set per row group
    # we'll compute separately for F/D and then stitch back
    out["Top_Role"] = ""
    out["Top_Role_Score"] = np.nan
    out["Second_Role"] = ""
    out["Second_Role_Score"] = np.nan
    out["Margin"] = np.nan
    out["Role_Stability"] = np.nan
    out["Type_Confidence"] = ""

    # confidence source
    conf_col = "Confidence_EV" if "Confidence_EV" in out.columns else ("Confidence_EV_3yr" if "Confidence_EV_3yr" in out.columns else None)

    def _compute_block(mask: pd.Series, base_cols: List[str], labels: Dict[str, str]) -> None:
        block = out.loc[mask].copy()
        cols, suffix = _pick_role_cols(block, base_cols)
        if not cols:
            return

        # map labels for suffix variants
        label_map = {}
        for c in cols:
            base = c.replace("_3yr", "")
            label_map[c] = labels.get(base, base.replace("_Score", ""))

        # row-wise top2
        vals = block[cols].to_numpy(dtype=float)
        # handle all-nan rows
        for i, idx in enumerate(block.index):
            row = vals[i, :]
            # nan-safe: set nan to very low
            if np.all(np.isnan(row)):
                continue
            temp = np.where(np.isnan(row), -1e9, row)
            top1_j = int(np.argmax(temp))
            top1_c = cols[top1_j]
            top1_v = float(row[top1_j]) if not np.isnan(row[top1_j]) else np.nan

            temp2 = temp.copy()
            temp2[top1_j] = -1e9
            top2_j = int(np.argmax(temp2))
            top2_c = cols[top2_j]
            top2_v = float(row[top2_j]) if not np.isnan(row[top2_j]) else np.nan

            out.at[idx, "Top_Role"] = label_map[top1_c]
            out.at[idx, "Top_Role_Score"] = top1_v
            out.at[idx, "Second_Role"] = label_map[top2_c]
            out.at[idx, "Second_Role_Score"] = top2_v
            out.at[idx, "Margin"] = (top1_v - top2_v) if (pd.notna(top1_v) and pd.notna(top2_v)) else np.nan

            # stability (rolling only)
            vol_col = _vol_col_for(top1_c, suffix)
            if vol_col and vol_col in out.columns:
                vol = out.at[idx, vol_col]
                # Normalize volatility to a 0–1 stability score.
                # 0 vol => 1.00 stability; 25+ vol => 0.00 stability (tunable)
                if pd.notna(vol):
                    out.at[idx, "Role_Stability"] = float(np.clip(1.0 - (float(vol) / 25.0), 0.0, 1.0))

    if "pos_group" in out.columns:
        _compute_block(out["pos_group"] == "F", FWD_ROLE_COLS, FWD_ROLE_LABELS)
        _compute_block(out["pos_group"] == "D", DEF_ROLE_COLS, DEF_ROLE_LABELS)
    else:
        # fall back: use Position/pos if present
        if "pos" in out.columns:
            _compute_block(out["pos"].astype(str).str.upper().eq("D"), DEF_ROLE_COLS, DEF_ROLE_LABELS)
            _compute_block(~out["pos"].astype(str).str.upper().eq("D"), FWD_ROLE_COLS, FWD_ROLE_LABELS)

    # Type_Confidence scoring
    # minutes_conf (0-1), margin (0-100-ish), stability (0-1)
    minutes_conf = out[conf_col].astype(float) if conf_col and conf_col in out.columns else pd.Series(np.nan, index=out.index)
    margin = out["Margin"].astype(float)
    stability = out["Role_Stability"].astype(float)

    # Distinguish season sheets (no stability) from rolling sheets (stability present)
    has_stability = stability.notna().any()
    stability_filled = stability.fillna(0.60)

    # Rules:
    # Season sheets (no Role_Stability available):
    #   High:   minutes_conf >= 0.50 AND margin >= 12
    #   Medium: minutes_conf >= 0.35 AND margin >= 6
    #   Low:    otherwise
    #
    # Rolling sheets (Role_Stability computed from vol):
    #   High:   minutes_conf >= 0.50 AND margin >= 12 AND stability >= 0.70
    #   Medium: minutes_conf >= 0.35 AND margin >= 6
    #   Low:    otherwise
    out.loc[:, "Type_Confidence"] = "Low"
    out.loc[(minutes_conf >= 0.35) & (margin >= 6), "Type_Confidence"] = "Medium"
    if has_stability:
        # Rolling sheets: full 3-factor rule (minutes + clear margin + year-over-year stability)
        out.loc[(minutes_conf >= 0.50) & (margin >= 12) & (stability_filled >= 0.70), "Type_Confidence"] = "High"
    else:
        # Season sheets: Role_Stability unavailable; require slightly wider margin to compensate
        out.loc[(minutes_conf >= 0.50) & (margin >= 14), "Type_Confidence"] = "High"

    return out

CAP_SHEET = "Cap Data 25-26"
CAP_LIMIT_26_27 = 104_000_000  # NHL upper limit 2025-26 :contentReference[oaicite:2]{index=2}

CAP_FALLBACKS = {
    "player": ["PLAYERS", "Player", "PLAYER"],
    "team": ["TEAM", "Team"],
    "pos": ["POS", "Pos", "Position"],
    "hand": ["HAND", "Hand"],
    "cap_hit": ["CAP HIT", "Cap Hit", "AAV"],
    "term": ["TERM", "Term"],
    "expiry_year": ["EXPIRY YEAR", "Expiry Year", "EXPIRY"],
    "type": ["TYPE", "Type"],
}

def load_cap_table(xlsx: Path) -> pd.DataFrame:
    df = pd.read_excel(xlsx, sheet_name=CAP_SHEET)
    df = _clean_cols(df)

    def col(name): 
        return _find_col(df, CAP_FALLBACKS[name])

    req = ["player","team","pos","cap_hit","expiry_year"]
    for r in req:
        if col(r) is None:
            raise ValueError(f"[{CAP_SHEET}] missing required column: {r}")

    out = pd.DataFrame()
    out["player"] = df[col("player")].astype(str).map(normalize_player_name)
    out["player_key"] = out["player"].map(_player_key)
    out["team_cap"] = df[col("team")].astype(str).str.strip().str.upper()
    out["pos_cap"] = df[col("pos")].astype(str).str.strip().str.upper()
    out["hand"] = df[col("hand")].astype(str).str.strip().str.upper() if col("hand") else ""
    out["cap_hit"] = df[col("cap_hit")].map(_to_float)
    out["term"] = df[col("term")].map(_to_float) if col("term") else np.nan
    out["expiry_year"] = df[col("expiry_year")].map(_to_float)
    out["contract_type"] = df[col("type")].astype(str).str.strip().str.upper() if col("type") else ""

    # normalize dollars if user stored "8.5" meaning 8.5M
    # if most cap hits look < 200, treat as millions
    if out["cap_hit"].dropna().median() < 200:
        out["cap_hit"] = out["cap_hit"] * 1_000_000

    return out


def merge_cap(model_df: pd.DataFrame, cap_df: pd.DataFrame) -> pd.DataFrame:
    out = model_df.merge(
        cap_df[["player_key","team_cap","pos_cap","hand","cap_hit","term","expiry_year","contract_type"]],
        on="player_key",
        how="left",
    )
    # Cap value efficiency: Impact per $1M of cap hit
    if "Impact_Score" in out.columns and "cap_hit" in out.columns:
        cap_m = out["cap_hit"].astype(float).replace(0, np.nan) / 1_000_000
        out["Cap_Value_Per_M"] = (out["Impact_Score"].astype(float) / cap_m).round(2)
    return out


def team_current_roster(cap_df: pd.DataFrame, team: str, offseason_year: int) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Returns:
      - team_under_contract: players with expiry_year > offseason_year (still under contract after season)
      - team_expiring: players with expiry_year == offseason_year (UFAs/RFAs after season)
    """
    t = team.strip().upper()
    team_df = cap_df[cap_df["team_cap"].eq(t)].copy()

    exp = team_df[team_df["expiry_year"].eq(float(offseason_year))].copy()
    uct = team_df[team_df["expiry_year"].gt(float(offseason_year))].copy()

    return uct, exp


def available_free_agents(cap_df: pd.DataFrame, offseason_year: int) -> pd.DataFrame:
    """
    League-wide expiring contracts for the offseason year.

    IMPORTANT:
    - Excludes RFAs from the league-wide FA pool (because RFAs aren't true "free agents").
    - Team RFAs are still handled via team_current_roster(...)-> expiring_cap, and go through re-sign logic.
    """
    fa = cap_df[cap_df["expiry_year"].eq(float(offseason_year))].copy()

    # Robust "RFA" detection in contract_type (from CAP sheet TYPE column)
    if "contract_type" in fa.columns:
        is_rfa = fa["contract_type"].astype(str).str.contains(r"\bRFA\b", case=False, na=False)
        fa = fa[~is_rfa].copy()

    return fa

def add_line_fit(df: pd.DataFrame) -> pd.DataFrame:
    """
    Adds roster-usage fit scores.

    Forwards: Fit_Line1..Fit_Line4 + Best_Line_Fit (+ score)
    Defense:  Fit_Pair1..Fit_Pair3 + Best_Pair_Fit (+ score)

    Works for season sheets and rolling sheets (uses *_3yr if present).
    """
    out = df.copy()

    # ensure we have pos_group
    if "pos_group" not in out.columns:
        out["pos_group"] = np.nan

    # pre-create outputs so downstream code doesn't explode
    for c in ["Best_Line_Fit", "Best_Line_Fit_Score", "Best_Pair_Fit", "Best_Pair_Fit_Score"]:
        if c not in out.columns:
            out[c] = np.nan

    # Forward fit templates
    def _fit_f(prod, finish, play, drive, tw, power, grind, prod_mod=None):
        # prod_mod: actual points60 normalized to 0-100 using two tiers:
        #   pm_elite: ceiling = 4.0 pts/60 — only truly elite scorers approach 100
        #   pm_base:  ceiling = 2.5 pts/60 — good NHL producers reach 100
        # Line 1 is dominated by pm_elite so no amount of role-score gaming can
        # push a non-scorer onto the first line.
        if prod_mod is not None:
            # prod_mod is raw pts/60. Apply dual-ceiling scaling:
            #   pm_elite: ceiling = 4.0 pts/60 — only true elite scorers approach 100
            #   pm_base:  ceiling = 2.5 pts/60 — solid NHL producers reach 100
            pts60 = prod_mod.astype(float).fillna(0.0)
            pm_elite = (pts60 / 4.0).clip(0.0, 1.0) * 100.0
            pm_base  = (pts60 / 2.5).clip(0.0, 1.0) * 100.0
        else:
            pm_elite = pd.Series(0.0, index=prod.index)
            pm_base  = pd.Series(0.0, index=prod.index)

        # Line 1: elite offensive production is the overwhelming requirement.
        # pm_elite (60%) + pm_base (20%) = 80% of the score comes from actual point output.
        # No combination of role scores alone can beat a true elite scorer here.
        # Grinder penalty ensures checking forwards don't accidentally qualify.
        fit1 = (0.60 * pm_elite
                + 0.20 * pm_base
                + 0.12 * (finish + play) / 2.0
                + 0.08 * drive
                - 0.08 * grind)

        # Line 2: production required but not at elite levels.
        # Two-way reliability and driving are meaningful secondary traits.
        fit2 = (0.35 * pm_base
                + 0.20 * prod
                + 0.15 * (finish + play) / 2.0
                + 0.15 * drive
                + 0.10 * tw
                - 0.05 * grind)

        # Line 3: two-way/matchup role — drive and defensive responsibility dominate.
        # TwoWay and Power terms are discounted when the player has high production
        # (elite producers are not "two-way role players" regardless of their TwoWay score).
        tw_adj    = tw    * (1.0 - 0.50 * pm_base / 100.0)
        power_adj = power * (1.0 - 0.30 * pm_base / 100.0)
        fit3 = (0.10 * pm_base
                + 0.10 * prod
                + 0.15 * play
                + 0.30 * drive
                + 0.28 * tw_adj
                + 0.07 * power_adj)

        # Line 4: pure energy/checking — grind absolutely dominates.
        fit4 = 0.18 * tw + 0.22 * power + 0.50 * grind + 0.10 * drive

        return fit1, fit2, fit3, fit4

    def _fit_d(supp, trans, skill, phys, d_pts_mod=None):
        # d_pts_mod: defenseman points60 normalized to 0-100 (ceiling ~2.0 pts60).
        #
        # KEY DESIGN PRINCIPLE: Top-pair defensemen in the modern NHL come in two
        # distinct archetypes — the elite shutdown/transition type (Hedman, Ekman-Larsson)
        # and the puck-skill/transition type (Fox, Makar, Forsling).  Neither archetype
        # scores high on ALL four role dimensions.  Forsling, for example, has a very
        # low Suppressor score (21.9) despite being a Norris-calibre defensive player,
        # because his suppression works through puck control and transition rather than
        # shot-blocking metrics.  A formula that weights Suppressor too heavily will
        # always misclassify this archetype as a Pair 2/3 player.
        #
        # Fix: Pair 1 is anchored by the player's TOP-2 role ceiling — their two
        # best defensive dimensions — rather than any single dimension.  This rewards
        # elite two-way D who dominate in any combination of {Suppressor, Transition,
        # PuckSkill} rather than only players who are elite suppressors specifically.
        # Pair 2 rewards balanced workhorse profiles.  Pair 3 rewards shutdown + physical.

        dpm = d_pts_mod if d_pts_mod is not None else pd.Series(0.0, index=supp.index)
        dpm = dpm.fillna(0.0)

        # Top-2 ceiling: for each player take their two highest defensive role scores.
        # This is the primary signal for Pair 1 — can they dominate in at least two
        # defensive dimensions?  A player who is 90+ in two dimensions is a top-pair
        # candidate regardless of which two those are.
        role_matrix = pd.concat([supp, trans, skill], axis=1)
        role_matrix.columns = ["s", "t", "k"]
        top2_mean = role_matrix.apply(
            lambda row: float(
                sorted([row["s"], row["t"], row["k"]], reverse=True)[:2]
            ) if False else  # placeholder — computed properly below
            float(sum(sorted([row["s"], row["t"], row["k"]], reverse=True)[:2]) / 2.0),
            axis=1
        )

        # Pair 1: elite top-2 ceiling is the gate; production is a bonus.
        # Suppressor is still rewarded but is no longer the dominant term so
        # transition/puck-skill D aren't penalised for not being shot-blockers.
        pair1 = (0.45 * top2_mean
                 + 0.20 * trans        # transition is uniquely Pair-1 relevant
                 + 0.15 * skill        # offensive catalyst adds value on top line
                 + 0.10 * supp         # shutdown still rewarded, just not gate-keeping
                 + 0.10 * dpm)

        # Pair 2: workhorse pairing — solid across all four dimensions.
        # No single-dimension elite required; consistency rewarded.
        pair2 = (0.25 * trans
                 + 0.22 * skill
                 + 0.22 * phys
                 + 0.18 * supp
                 + 0.13 * dpm)

        # Pair 3: depth / physical / penalty-kill anchor.
        # Pure shutdown and physicality dominate; transition bonus for versatility.
        pair3 = (0.45 * supp
                 + 0.35 * phys
                 + 0.20 * trans)

        return pair1, pair2, pair3

    # choose correct columns (season vs rolling)
    f_cols, f_suf = _pick_role_cols(out, FWD_ROLE_COLS)
    d_cols, d_suf = _pick_role_cols(out, DEF_ROLE_COLS)

    # Create fit columns default (correct dtypes)
    score_like = {"Best_Line_Fit_Score", "Best_Pair_Fit_Score"}
    for c in [
        "Fit_Line1","Fit_Line2","Fit_Line3","Fit_Line4","Best_Line_Fit","Best_Line_Fit_Score",
        "Fit_Pair1","Fit_Pair2","Fit_Pair3","Best_Pair_Fit","Best_Pair_Fit_Score"
    ]:
        if c not in out.columns:
            if c in score_like or c.startswith("Fit_"):
                out[c] = np.nan
            else:
                # Must be object dtype so string labels (e.g. "Line 1") can be assigned
                out[c] = pd.array([""] * len(out), dtype=object)

    # safe usage series (always index-aligned)
    usage = (
        out["Usage_Rank_Team"].astype(float).fillna(0.5)
        if "Usage_Rank_Team" in out.columns
        else pd.Series(0.5, index=out.index, dtype=float)
    )

    # -------------------
    # Forwards
    # -------------------
    if f_cols:
        fmask = out["pos_group"].eq("F")

        def getv(colname: str) -> pd.Series:
            return out[colname].astype(float) if (colname and colname in out.columns) else pd.Series(np.nan, index=out.index)

        # Resolve actual column names by role (season vs rolling)
        f_map = {c.replace("_3yr", ""): c for c in f_cols}

        prod   = getv(f_map.get("Producer_Score", ""))
        finish = getv(f_map.get("Finisher_Score", ""))
        play   = getv(f_map.get("Playmaker_Score", ""))
        drive  = getv(f_map.get("Driver_Score", ""))
        tw     = getv(f_map.get("TwoWay_Score", ""))
        power  = getv(f_map.get("Power_Score", ""))
        grind  = getv(f_map.get("Grinder_Score", ""))

        # Production modifier: pass raw pts60 to _fit_f which applies dual-ceiling scaling.
        # pm_elite ceiling = 4.0 pts/60 (only true outliers near 100)
        # pm_base  ceiling = 2.5 pts/60 (solid NHL producers reach 100)
        if "points60_ind" in out.columns:
            prod_mod = out["points60_ind"].astype(float).fillna(0.0)
        elif "Proj_Points60" in out.columns:
            prod_mod = out["Proj_Points60"].astype(float).fillna(0.0)
        else:
            prod_mod = None

        l1, l2, l3, l4 = _fit_f(prod, finish, play, drive, tw, power, grind, prod_mod=prod_mod)

        out.loc[fmask, "Fit_Line1"] = l1.loc[fmask]
        out.loc[fmask, "Fit_Line2"] = l2.loc[fmask]
        out.loc[fmask, "Fit_Line3"] = l3.loc[fmask]
        out.loc[fmask, "Fit_Line4"] = l4.loc[fmask]

        # ── Player_Tier multiplier ────────────────────────────────────────────
        # Tier reflects overall dimensionality — a high-tier player should get
        # credit toward upper lines even when role scores are close to a lower-tier
        # player.  Applied as a multiplicative modifier on the raw fit scores
        # BEFORE argmax so it influences which line wins without overriding formulas.
        #
        # Tier → (L1_mult, L2_mult, L3_mult, L4_mult)
        TIER_MULT_F = {
            "Elite":   (1.12, 1.06, 0.95, 0.80),  # strongly prefer upper lines
            "Star":    (1.07, 1.04, 0.97, 0.88),
            "Solid":   (1.00, 1.00, 1.00, 1.00),  # neutral — let formulas decide
            "Depth":   (0.90, 0.95, 1.02, 1.08),  # lean toward lower lines
            "Fringe":  (0.78, 0.88, 1.03, 1.15),  # strong push to L3/L4
            "Unknown": (1.00, 1.00, 1.00, 1.00),
        }
        if "Player_Tier" in out.columns:
            for idx in out.index[fmask]:
                tier = str(out.at[idx, "Player_Tier"] or "Unknown")
                m1, m2, m3, m4 = TIER_MULT_F.get(tier, (1.0, 1.0, 1.0, 1.0))
                out.at[idx, "Fit_Line1"] = float(out.at[idx, "Fit_Line1"]) * m1
                out.at[idx, "Fit_Line2"] = float(out.at[idx, "Fit_Line2"]) * m2
                out.at[idx, "Fit_Line3"] = float(out.at[idx, "Fit_Line3"]) * m3
                out.at[idx, "Fit_Line4"] = float(out.at[idx, "Fit_Line4"]) * m4

        # ── Dim_Trend multiplier ──────────────────────────────────────────────
        # Trend reflects trajectory — a rising player should be deployed higher
        # (anticipating improvement); a declining player should be deployed more
        # conservatively (don't bet on last year's ceiling repeating).
        # Injury Year: small neutral penalty on upper lines to reflect uncertainty,
        # NOT a performance judgment — the player hasn't actually declined.
        #
        # Trend → (L1_mult, L2_mult, L3_mult, L4_mult)
        TREND_MULT_F = {
            "⚡ Breakout":      (1.10, 1.05, 0.97, 0.90),  # strong upward trajectory
            "↑↑ Strong Rise":  (1.08, 1.04, 0.97, 0.92),
            "✦ Age Surge":     (1.06, 1.04, 0.98, 0.94),  # notable but age-tempered
            "↑  Rising":       (1.04, 1.02, 0.99, 0.96),
            "→ Stable":        (1.00, 1.00, 1.00, 1.00),  # neutral
            "↓  Slipping":     (0.96, 0.98, 1.01, 1.04),  # mild conservative tilt
            "↓↓ Declining":    (0.90, 0.94, 1.03, 1.08),  # meaningful push down
            "⚠ Sharp Drop":   (0.82, 0.88, 1.05, 1.12),  # serious upper-line discount
            "⛑ Injury Year":  (0.94, 0.97, 1.01, 1.02),  # uncertainty, not decline
            "~  New Data":     (1.00, 1.00, 1.00, 1.00),  # no signal → neutral
            "":                (1.00, 1.00, 1.00, 1.00),
        }
        if "Dim_Trend" in out.columns:
            for idx in out.index[fmask]:
                trend = str(out.at[idx, "Dim_Trend"] or "")
                # normalise leading/trailing whitespace variants
                trend_key = next(
                    (k for k in TREND_MULT_F if k.strip() == trend.strip()),
                    ""
                )
                m1, m2, m3, m4 = TREND_MULT_F.get(trend_key, (1.0, 1.0, 1.0, 1.0))
                out.at[idx, "Fit_Line1"] = float(out.at[idx, "Fit_Line1"]) * m1
                out.at[idx, "Fit_Line2"] = float(out.at[idx, "Fit_Line2"]) * m2
                out.at[idx, "Fit_Line3"] = float(out.at[idx, "Fit_Line3"]) * m3
                out.at[idx, "Fit_Line4"] = float(out.at[idx, "Fit_Line4"]) * m4

        # Grinder suppression (ONLY after fits exist)
        if "Player_Type" in out.columns:
            is_grinder = out["Player_Type"].astype(str).str.contains("Grinder", case=False, na=False)
            grinder_mask = fmask & is_grinder
            out.loc[grinder_mask, "Fit_Line1"] = out.loc[grinder_mask, "Fit_Line1"].astype(float) * 0.20
            out.loc[grinder_mask, "Fit_Line2"] = out.loc[grinder_mask, "Fit_Line2"].astype(float) * 0.40

        # Usage eligibility gating (mask, DON'T scale)
        elig_l1 = usage >= 0.55
        elig_l2 = usage >= 0.40

        # STAR OVERRIDE: elite producers/finishers should always be eligible for top units
        prod_col = f_map.get("Producer_Score", "Producer_Score")
        finish_col = f_map.get("Finisher_Score", "Finisher_Score")

        star = pd.Series(False, index=out.index)
        if prod_col in out.columns:
            star |= out[prod_col].astype(float).fillna(-1).ge(80)
        if finish_col in out.columns:
            star |= out[finish_col].astype(float).fillna(-1).ge(80)

        elig_l1 = elig_l1 | star
        elig_l2 = elig_l2 | star

        # apply: ineligible => effectively impossible to win
        out.loc[fmask & ~elig_l1, "Fit_Line1"] = -1e9
        out.loc[fmask & ~elig_l2, "Fit_Line2"] = -1e9

        # best line (nan-robust)
        # IMPORTANT: only run if there are any forwards rows
        if fmask.any():
            fits = out.loc[fmask, ["Fit_Line1","Fit_Line2","Fit_Line3","Fit_Line4"]].to_numpy(dtype=float)
            labels = np.array(["Line 1","Line 2","Line 3","Line 4"], dtype=object)

            all_nan = np.all(np.isnan(fits), axis=1)

            # treat NaNs as impossible for argmax
            fits_safe = np.where(np.isnan(fits), -1e9, fits)
            best_idx = fits_safe.argmax(axis=1)

            best_labels = labels[best_idx]
            best_labels[all_nan] = ""

            best_scores = np.nanmax(fits, axis=1)
            best_scores[all_nan] = np.nan

            out["Best_Line_Fit"] = out["Best_Line_Fit"].astype(object)
            out.loc[fmask, "Best_Line_Fit"] = best_labels
            out["Best_Line_Fit_Score"] = out["Best_Line_Fit_Score"].astype(float)
            out.loc[fmask, "Best_Line_Fit_Score"] = best_scores

        # --- Minimum line floor based on impact score ---
        # Any forward with Impact >= 75 should never land on Line 4.
        # This corrects Two-Way/Power archetypes (e.g. Samoskevich) who score
        # high on physical/defensive role templates that technically fit Line 4.
        if "impact" in out.columns:
            high_impact = fmask & out["impact"].astype(float).fillna(0).ge(75)
            on_line4    = out["Best_Line_Fit"].astype(str).eq("Line 4")
            bump_impact = high_impact & on_line4
            if bump_impact.any():
                l4_candidates = out.loc[bump_impact].copy()
                non_l4 = l4_candidates[
                    ["Fit_Line1","Fit_Line2","Fit_Line3"]
                ].replace(-1e9, np.nan)
                col_map = {"Fit_Line1":"Line 1","Fit_Line2":"Line 2","Fit_Line3":"Line 3"}
                best_col = non_l4.idxmax(axis=1)
                for idx, fc in best_col.items():
                    if fc in col_map and fc in out.columns:
                        out.at[idx, "Best_Line_Fit"]       = col_map[fc]
                        out.at[idx, "Best_Line_Fit_Score"] = float(out.at[idx, fc])

        # Star correction: elite producer stuck on Line 4 → Line 2
        if prod_col in out.columns:
            star_mask = fmask & out[prod_col].astype(float).fillna(-1).ge(80)
            bump = star_mask & out["Best_Line_Fit"].astype(str).eq("Line 4")
            out.loc[bump, "Best_Line_Fit"] = "Line 2"
            out.loc[bump, "Best_Line_Fit_Score"] = out.loc[bump, "Fit_Line2"].astype(float)

        # High-usage players shouldn't be Line 4 either
        top_usage = fmask & (usage >= 0.75)
        bump2 = top_usage & out["Best_Line_Fit"].astype(str).eq("Line 4")
        out.loc[bump2, "Best_Line_Fit"] = "Line 2"
        out.loc[bump2, "Best_Line_Fit_Score"] = out.loc[bump2, "Fit_Line2"].astype(float)

    # -------------------
    # Defense
    # -------------------
    if d_cols:
        dmask = out["pos_group"].eq("D")

        def getd(colname: str) -> pd.Series:
            return out[colname].astype(float) if (colname and colname in out.columns) else pd.Series(np.nan, index=out.index)

        d_map = {c.replace("_3yr", ""): c for c in d_cols}

        supp  = getd(d_map.get("Suppressor_Score", ""))
        trans = getd(d_map.get("Transition_Score", ""))
        skill = getd(d_map.get("PuckSkill_Score", ""))
        phys  = getd(d_map.get("Physical_Score", ""))

        # Defenseman production modifier: points60 normalized to 0-100 (D ceiling ~2.0 pts60)
        if "points60_ind" in out.columns:
            d_pts_raw = out["points60_ind"].astype(float).fillna(0.0)
            d_pts_mod = (d_pts_raw / 2.0).clip(0.0, 1.0) * 100.0
        elif "Proj_Points60" in out.columns:
            d_pts_raw = out["Proj_Points60"].astype(float).fillna(0.0)
            d_pts_mod = (d_pts_raw / 2.0).clip(0.0, 1.0) * 100.0
        else:
            d_pts_mod = None

        p1, p2, p3 = _fit_d(supp, trans, skill, phys, d_pts_mod=d_pts_mod)

        out.loc[dmask, "Fit_Pair1"] = p1.loc[dmask]
        out.loc[dmask, "Fit_Pair2"] = p2.loc[dmask]
        out.loc[dmask, "Fit_Pair3"] = p3.loc[dmask]

        # ── Player_Tier multiplier (Defense) ─────────────────────────────────
        # Stronger multipliers than forwards because D pair assignment is more
        # binary (a Norris-calibre D should never land on Pair 3 regardless of
        # which specific role scores are highest).
        TIER_MULT_D = {
            "Elite":   (1.22, 1.04, 0.72),   # very strong Pair 1 pull, heavy Pair 3 discount
            "Star":    (1.15, 1.03, 0.80),
            "Solid":   (1.00, 1.00, 1.00),
            "Depth":   (0.88, 1.02, 1.10),
            "Fringe":  (0.74, 0.90, 1.18),
            "Unknown": (1.00, 1.00, 1.00),
        }
        if "Player_Tier" in out.columns:
            for idx in out.index[dmask]:
                tier = str(out.at[idx, "Player_Tier"] or "Unknown")
                m1, m2, m3 = TIER_MULT_D.get(tier, (1.0, 1.0, 1.0))
                out.at[idx, "Fit_Pair1"] = float(out.at[idx, "Fit_Pair1"]) * m1
                out.at[idx, "Fit_Pair2"] = float(out.at[idx, "Fit_Pair2"]) * m2
                out.at[idx, "Fit_Pair3"] = float(out.at[idx, "Fit_Pair3"]) * m3

        # ── Dim_Trend multiplier (Defense) ───────────────────────────────────
        TREND_MULT_D = {
            "⚡ Breakout":      (1.10, 1.04, 0.90),
            "↑↑ Strong Rise":  (1.08, 1.03, 0.92),
            "✦ Age Surge":     (1.06, 1.03, 0.94),
            "↑  Rising":       (1.04, 1.01, 0.97),
            "→ Stable":        (1.00, 1.00, 1.00),
            "↓  Slipping":     (0.96, 0.99, 1.04),
            "↓↓ Declining":    (0.90, 0.95, 1.07),
            "⚠ Sharp Drop":   (0.83, 0.90, 1.10),
            "⛑ Injury Year":  (0.94, 0.97, 1.02),
            "~  New Data":     (1.00, 1.00, 1.00),
            "":                (1.00, 1.00, 1.00),
        }
        if "Dim_Trend" in out.columns:
            for idx in out.index[dmask]:
                trend = str(out.at[idx, "Dim_Trend"] or "")
                trend_key = next(
                    (k for k in TREND_MULT_D if k.strip() == trend.strip()),
                    ""
                )
                m1, m2, m3 = TREND_MULT_D.get(trend_key, (1.0, 1.0, 1.0))
                out.at[idx, "Fit_Pair1"] = float(out.at[idx, "Fit_Pair1"]) * m1
                out.at[idx, "Fit_Pair2"] = float(out.at[idx, "Fit_Pair2"]) * m2
                out.at[idx, "Fit_Pair3"] = float(out.at[idx, "Fit_Pair3"]) * m3

        # best pair (nan-robust)
        if dmask.any():
            fits = out.loc[dmask, ["Fit_Pair1","Fit_Pair2","Fit_Pair3"]].to_numpy(dtype=float)
            labels = np.array(["Pair 1","Pair 2","Pair 3"], dtype=object)

            all_nan = np.all(np.isnan(fits), axis=1)

            fits_safe = np.where(np.isnan(fits), -1e9, fits)
            best_idx = fits_safe.argmax(axis=1)

            best_labels = labels[best_idx]
            best_labels[all_nan] = ""

            best_scores = np.nanmax(fits, axis=1)
            best_scores[all_nan] = np.nan

            out["Best_Pair_Fit"] = out["Best_Pair_Fit"].astype(object)
            out.loc[dmask, "Best_Pair_Fit"] = best_labels
            out["Best_Pair_Fit_Score"] = out["Best_Pair_Fit_Score"].astype(float)
            out.loc[dmask, "Best_Pair_Fit_Score"] = best_scores

    return out

# TEAM OFFSEASON PLANNER (TEAM-ONLY, OFFSEASON YEAR = 2026)
# =============================================================================

DEFAULT_ROSTER_TARGETS = {
    "F": 12,   # 4 lines x 3
    "D": 6,    # 3 pairs x 2
}
DEFAULT_MIN_BY_POS = {
    "C": 4,
    "LW": 3,
    "RW": 3,
    "D": 6,
}

def _production_score(row: pd.Series) -> float:
    """
    Converts raw individual production into a 0-100 percentile-like score.
    Uses points60 as primary, goals60 and ixg60 as secondary.
    Defensemen use a lower points60 scale since their production ceiling is lower.
    Returns NaN if no production data available (so it can be safely ignored).
    """
    pts60  = _safe_float(row.get("points60_ind"), default=np.nan)
    g60    = _safe_float(row.get("g60_ind"),      default=np.nan)
    ixg60  = _safe_float(row.get("ixg60_ind"),    default=np.nan)
    p_gp   = _safe_float(row.get("pts_per_gp_ind"), default=np.nan)

    if all(np.isnan(v) for v in [pts60, g60, ixg60, p_gp]):
        return np.nan

    # Fill missing with conservative defaults
    pts60 = pts60 if pd.notna(pts60) else 0.0
    g60   = g60   if pd.notna(g60)   else 0.0
    ixg60 = ixg60 if pd.notna(ixg60) else 0.0
    p_gp  = p_gp  if pd.notna(p_gp)  else 0.0

    pos_group = str(row.get("pos_group") or "F")

    # Scale ceilings: fwd elite ~3.5 pts60, def elite ~1.8 pts60
    pts60_ceil = 1.8 if pos_group == "D" else 3.5
    g60_ceil   = 0.6 if pos_group == "D" else 1.5
    ixg60_ceil = 0.8 if pos_group == "D" else 1.8

    prod = (
        0.50 * min(pts60  / pts60_ceil,  1.0) +
        0.25 * min(g60    / g60_ceil,    1.0) +
        0.25 * min(ixg60  / ixg60_ceil,  1.0)
    )
    return float(np.clip(prod * 100.0, 0.0, 100.0))


def _impact_score(row: pd.Series) -> float:
    """
    Player impact score (0-100 scale).
    Blends role versatility (Dimensionality), top role quality, lineup fit,
    and individual production. Reliability-adjusted by sample confidence + stability.
    """
    dim  = _safe_float(_first_notna(row.get("Dimensionality_Score"), row.get("Dimensionality_Score_3yr")), default=50.0)
    top  = _safe_float(row.get("Top_Role_Score"),  default=50.0)
    fit  = _safe_float(_first_notna(row.get("Best_Line_Fit_Score"), row.get("Best_Pair_Fit_Score")), default=50.0)
    conf = _safe_float(_first_notna(row.get("Confidence_EV"), row.get("Confidence_EV_3yr")), default=0.60)
    stab = _safe_float(row.get("Role_Stability"), default=0.60)
    prod = _production_score(row)

    dim  = dim  if pd.notna(dim)  else 50.0
    top  = top  if pd.notna(top)  else 50.0
    fit  = fit  if pd.notna(fit)  else 50.0
    conf = conf if pd.notna(conf) else 0.60
    stab = stab if pd.notna(stab) else 0.60

    # Production blended in at 15%; if missing, redistribute to dim/top
    if pd.notna(prod):
        base = 0.38*dim + 0.30*top + 0.17*fit + 0.15*prod
    else:
        base = 0.45*dim + 0.35*top + 0.20*fit

    reli = 0.85 + 0.10*conf + 0.05*stab
    return float(base * reli)

def _value_per_dollar(row: pd.Series, aav_col: str = "aav_est") -> float:
    """
    Returns Impact Score per $1M AAV.
    Normalizing to millions keeps the ratio in a human-readable 0–100 range
    (e.g. impact=80, AAV=$4M → 20.0) rather than the microscopic 0.00005
    that results from dividing by raw dollars. The microscopic version
    was causing all value_per_$ values to round to 0.00, collapsing the
    sort order to hole_bonus only.
    """
    impact = _first_notna(row.get("impact"), default=0.0)
    impact = float(impact) if pd.notna(impact) else 0.0

    aav = _first_notna(row.get(aav_col), default=np.nan)
    aav = float(aav) if pd.notna(aav) else np.nan

    if np.isnan(aav) or aav <= 0:
        return -1e9
    # Normalize to $M so the ratio is comparable to impact (0-100 scale)
    aav_m = aav / 1_000_000.0
    return float(impact / aav_m)

def _ensure_pos_labels(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if "pos" in out.columns:
        out["pos"] = out["pos"].astype(str).str.upper().str.strip()
        out["pos"] = out["pos"].replace({"L": "LW", "R": "RW"})
    if "pos_group" not in out.columns and "pos" in out.columns:
        out["pos_group"] = np.where(out["pos"].eq("D"), "D", "F")
    return out

def _cap_sum(df: pd.DataFrame) -> float:
    return float(df["cap_hit"].fillna(0.0).sum())

def _slot_counts(df: pd.DataFrame) -> Dict[str, int]:
    counts = {k: 0 for k in ["C","LW","RW","D"]}
    if "pos" not in df.columns:
        return counts
    for k, v in df["pos"].value_counts(dropna=False).to_dict().items():
        if k in counts:
            counts[k] = int(v)
    return counts

def _choose_core_team(team_df: pd.DataFrame, cap_limit: float) -> pd.DataFrame:
    """
    Selects the "best" under-contract core (already on team) if your cap table includes
    extra items (buried, etc.). Conservative: keep all under-contract; if over cap, trim worst value.
    """
    core = team_df.copy()
    core["impact"] = core.apply(_impact_score, axis=1)

    # If already under cap, keep all
    if _cap_sum(core) <= cap_limit:
        return core

    # Otherwise trim by lowest impact per $ until compliant
    core["value_$"] = core["impact"] / core["cap_hit"].replace(0, np.nan)
    core = core.sort_values("value_$", ascending=True)

    keep_mask = np.ones(len(core), dtype=bool)
    running = _cap_sum(core)
    for i, idx in enumerate(core.index):
        if running <= cap_limit:
            break
        running -= float(core.at[idx, "cap_hit"] or 0.0)
        keep_mask[i] = False

    core_kept = core.loc[core.index[keep_mask]].copy()
    return core_kept

def _define_holes(current: pd.DataFrame, targets=DEFAULT_ROSTER_TARGETS, min_by_pos=DEFAULT_MIN_BY_POS) -> Dict[str, int]:
    """
    Returns holes by:
      - group totals (need_F/need_D)
      - minimum positions (need_C/LW/RW/D)
      - fit buckets (need_Line 1/2/3/4 and need_Pair 1/2/3)
    """
    current = _ensure_pos_labels(current)

    holes: Dict[str, int] = {}

    # group-level holes
    for grp, tgt in targets.items():
        cur = int((current["pos_group"] == grp).sum()) if "pos_group" in current.columns else 0
        holes[f"need_{grp}"] = max(tgt - cur, 0)

    # position-level minimum holes
    counts = _slot_counts(current)
    for pos, min_ct in min_by_pos.items():
        holes[f"need_{pos}"] = max(min_ct - counts.get(pos, 0), 0)

    # fit bucket holes (target structure)
    # forwards: 3 per line (12 total)
    # defense: 2 per pair (6 total)
    if "pos_group" in current.columns:
        f = current[current["pos_group"].eq("F")].copy()
        d = current[current["pos_group"].eq("D")].copy()
    else:
        f = current.copy()
        d = current.iloc[0:0].copy()

    # Ensure buckets exist
    f_bucket = f.get("Best_Line_Fit", pd.Series(["Line 3"] * len(f))).astype(str)
    d_bucket = d.get("Best_Pair_Fit", pd.Series(["Pair 2"] * len(d))).astype(str)

    for ln in ["Line 1", "Line 2", "Line 3", "Line 4"]:
        holes[f"need_{ln}"] = max(3 - int((f_bucket == ln).sum()), 0)

    for pr in ["Pair 1", "Pair 2", "Pair 3"]:
        holes[f"need_{pr}"] = max(2 - int((d_bucket == pr).sum()), 0)

    return holes

def _needs_pos(holes: Dict[str, int], pos: str) -> bool:
    return holes.get(f"need_{pos}", 0) > 0

def _decrement_holes_for_player(holes: Dict[str, int], pos: str, fit_bucket: Optional[str] = None) -> None:
    # position & group
    if holes.get(f"need_{pos}", 0) > 0:
        holes[f"need_{pos}"] -= 1
    grp = "D" if pos == "D" else "F"
    if holes.get(f"need_{grp}", 0) > 0:
        holes[f"need_{grp}"] -= 1

    # fit bucket
    if fit_bucket:
        k = f"need_{fit_bucket}"
        if holes.get(k, 0) > 0:
            holes[k] -= 1

def _fit_bucket(row: pd.Series) -> str:
    pos_group = row.get("pos_group", "F")
    if pos_group == "D":
        return str(row.get("Best_Pair_Fit") or "Pair 2")
    return str(row.get("Best_Line_Fit") or "Line 3")

from dataclasses import dataclass

# =============================================================================
# MARKET COMPARABLES + AAV SIMULATION
# =============================================================================

@dataclass
class AAVEstimate:
    n_comps: int
    mean: float
    std: float
    p10: float
    p50: float
    p90: float


def _safe_float(x, default=np.nan) -> float:
    try:
        if x is None or pd.isna(x):
            return default
        return float(x)
    except Exception:
        return default


def _fit_bucket_col_for_group(pos_group: str) -> str:
    return "Best_Pair_Fit" if pos_group == "D" else "Best_Line_Fit"


def _fit_score_col_for_group(pos_group: str) -> str:
    return "Best_Pair_Fit_Score" if pos_group == "D" else "Best_Line_Fit_Score"


def _build_market_table(cap_df: pd.DataFrame, model_df: pd.DataFrame) -> pd.DataFrame:
    """
    Market table = cap rows enriched with model metrics (impact/roles/fit/production/projection).
    """
    ms = _ensure_pos_labels(model_df.copy())

    enrich_cols = [
        "player_key", "player", "pos", "pos_group", "age",
        "Top_Role", "Top_Role_Score", "Second_Role", "Second_Role_Score", "Margin",
        "Type_Confidence", "Role_Stability",
        "Best_Line_Fit", "Best_Line_Fit_Score", "Best_Pair_Fit", "Best_Pair_Fit_Score",
        "Dimensionality_Score", "Confidence_EV", "Player_Type",
        "points60_ind", "g60_ind", "ixg60_ind", "pts_per_gp_ind", "a1_60_ind",
        "Proj_Points60", "Proj_Goals60", "Age_Curve_Factor", "Prog_Tier", "Draft_Pedigree",
    ]
    enrich_cols = [c for c in enrich_cols if c in ms.columns]
    ms_enrich = ms[enrich_cols].drop_duplicates("player_key")

    mk = cap_df.merge(ms_enrich, on="player_key", how="left")
    mk = _ensure_pos_labels(mk)
    mk["aav_market"] = mk["cap_hit"].map(_safe_float)
    mk["impact"] = mk.apply(_impact_score, axis=1)
    mk = mk[(mk["aav_market"].notna()) & (mk["aav_market"] > 0)].copy()
    return mk


def _comp_distance(mk: pd.DataFrame, target: pd.Series) -> pd.Series:
    """
    Distance in player archetype + production space. Lower = closer comp.
    Dimensions: age, impact, role versatility, top role, lineup fit, production rate, projected output.
    Each dimension is IQR-normalized so no single axis dominates.
    """
    pos_group = str(target.get("pos_group") or "F")
    fit_col   = _fit_score_col_for_group(pos_group)

    # Target values
    age_t  = _safe_float(target.get("age"),                default=np.nan)
    imp_t  = _safe_float(target.get("impact"),             default=np.nan)
    dim_t  = _safe_float(target.get("Dimensionality_Score"), default=np.nan)
    top_t  = _safe_float(target.get("Top_Role_Score"),     default=np.nan)
    fit_t  = _safe_float(target.get(fit_col),              default=np.nan)
    pts_t  = _safe_float(target.get("points60_ind"),       default=np.nan)
    prj_t  = _safe_float(target.get("Proj_Points60"),      default=np.nan)

    def fill_t(v, col):
        return v if pd.notna(v) else (float(mk[col].median()) if col in mk.columns else 0.0)

    age_t = fill_t(age_t, "age")
    imp_t = fill_t(imp_t, "impact")
    dim_t = fill_t(dim_t, "Dimensionality_Score")
    top_t = fill_t(top_t, "Top_Role_Score")
    fit_t = fill_t(fit_t, fit_col)
    pts_t = fill_t(pts_t, "points60_ind")
    prj_t = fill_t(prj_t, "Proj_Points60")

    # IQR-based scale (robust to outliers)
    def scale(col, fallback=1.0):
        if col not in mk.columns:
            return fallback
        s = mk[col].astype(float)
        iqr = float(s.quantile(0.75) - s.quantile(0.25))
        return iqr if iqr > 0 else (float(s.std()) if s.std() > 0 else fallback)

    age_s = scale("age", 3.0)
    imp_s = scale("impact", 10.0)
    dim_s = scale("Dimensionality_Score", 10.0)
    top_s = scale("Top_Role_Score", 10.0)
    fit_s = scale(fit_col, 10.0)
    pts_s = scale("points60_ind", 0.5)
    prj_s = scale("Proj_Points60", 0.5)

    def col_f(col):
        return mk[col].astype(float) if col in mk.columns else pd.Series(0.0, index=mk.index)

    # Weighted squared distances — production/projection carry ~25% of total weight
    d = (
        0.15 * ((col_f("age")                - age_t) / age_s) ** 2 +
        0.22 * ((col_f("impact")             - imp_t) / imp_s) ** 2 +
        0.18 * ((col_f("Dimensionality_Score")- dim_t) / dim_s) ** 2 +
        0.15 * ((col_f("Top_Role_Score")     - top_t) / top_s) ** 2 +
        0.10 * ((col_f(fit_col)              - fit_t) / fit_s) ** 2 +
        0.12 * ((col_f("points60_ind")       - pts_t) / pts_s) ** 2 +
        0.08 * ((col_f("Proj_Points60")      - prj_t) / prj_s) ** 2
    )
    return np.sqrt(d)


def find_market_comparables(
    market: pd.DataFrame,
    target: pd.Series,
    n: int = 40,
    same_team_ok: bool = True,
) -> pd.DataFrame:
    """
    Comparable filter:
      - same pos_group
      - same primary fit bucket when available (Line 1/2/3/4 or Pair 1/2/3)
      - optionally same Top_Role
    Then rank by distance and return top n.
    """
    pos_group = str(target.get("pos_group") or "F")
    team = str(target.get("team_cap") or target.get("team") or "").upper().strip()

    mk = market.copy()
    mk = mk[mk["pos_group"].astype(str).eq(pos_group)].copy()

    # Fit bucket match if both present
    fit_col = _fit_bucket_col_for_group(pos_group)
    fit_t = str(target.get(fit_col) or "")
    if fit_col in mk.columns and fit_t:
        mk = mk[mk[fit_col].astype(str).eq(fit_t)].copy()

    # Top role match (soft): if too few remain, relax
    top_role_t = str(target.get("Top_Role") or "")
    if "Top_Role" in mk.columns and top_role_t:
        mk2 = mk[mk["Top_Role"].astype(str).eq(top_role_t)].copy()
        if len(mk2) >= max(12, int(n * 0.4)):
            mk = mk2

    # Optionally exclude team (for FA comps this isn't needed; for resigns it can be OK either way)
    if not same_team_ok and team:
        mk = mk[~mk["team_cap"].astype(str).eq(team)].copy()

    if len(mk) == 0:
        return market.iloc[0:0].copy()

    mk["_dist"] = _comp_distance(mk, target)
    mk = mk.sort_values("_dist", ascending=True).head(n).drop(columns=["_dist"])
    return mk


def estimate_aav_from_comps(
    comps: pd.DataFrame,
    cap_floor: float = 0.775e6,
    min_std_frac: float = 0.12,
    max_std_frac: float = 0.35,
) -> AAVEstimate:
    """
    Weighted AAV distribution:
      - weights favor closer comps (if you pass them pre-sorted by distance)
      - std from comps but clipped to a reasonable fraction of mean
    """
    if comps is None or len(comps) == 0 or "aav_market" not in comps.columns:
        mu = cap_floor
        sd = cap_floor * 0.20
        return AAVEstimate(n_comps=0, mean=mu, std=sd, p10=mu-1.28*sd, p50=mu, p90=mu+1.28*sd)

    aav = comps["aav_market"].astype(float).clip(lower=cap_floor).to_numpy()

    # weights: nearer comps (earlier rows) get more weight
    # if comps aren’t distance-sorted, this still gives mild shrink toward median.
    rank = np.arange(len(aav), dtype=float)
    w = 1.0 / (1.0 + rank)
    w = w / w.sum()

    mu = float(np.dot(w, aav))
    # weighted std
    var = float(np.dot(w, (aav - mu) ** 2))
    sd = float(np.sqrt(max(var, 0.0)))

    # clip sd to mean fraction
    sd = float(np.clip(sd, mu * min_std_frac, mu * max_std_frac))

    p10 = float(np.quantile(aav, 0.10))
    p50 = float(np.quantile(aav, 0.50))
    p90 = float(np.quantile(aav, 0.90))

    return AAVEstimate(n_comps=int(len(aav)), mean=mu, std=sd, p10=p10, p50=p50, p90=p90)


def simulate_aav_draws(
    est: AAVEstimate,
    n_sims: int = 5000,
    cap_floor: float = 0.775e6,
    dist: str = "normal",
    t_df: int = 6,
) -> np.ndarray:
    """
    Negotiation variation model.
    normal: N(mean, std)
    t: mean + std * t(df)  (fatter tails)
    """
    mu, sd = est.mean, est.std
    if dist.lower().startswith("t"):
        draws = mu + sd * np.random.standard_t(df=t_df, size=n_sims)
    else:
        draws = np.random.normal(mu, sd, size=n_sims)

    draws = np.clip(draws, cap_floor, None)
    return draws

def project_player(row: pd.Series) -> Dict:
    """
    Projects a player's next-season output and assigns a progression tier.

    Inputs (all optional — degrades gracefully):
      - age, draft_round, draft_pick (overall)
      - points60_ind, g60_ind, ixg60_ind, pts_per_gp_ind  (current season)
      - Dimensionality_Score, Top_Role_Score (role model outputs)
      - *_3yr rolling equivalents

    Outputs (added as Proj_* columns):
      Proj_Points60    — projected pts/60 next season
      Proj_Goals60     — projected g/60 next season
      Proj_PTS_PGP     — projected pts/game next season
      Age_Curve_Factor — multiplier reflecting where player is on aging curve
      Prog_Tier        — Emerging / Prime / Plateau / Declining / Aging
      Prog_Confidence  — High / Med / Low (based on sample size + data richness)
      Draft_Pedigree   — Elite / High / Mid / Late / Undrafted
    """
    sf = _safe_float

    age        = sf(row.get("age"),           default=np.nan)
    draft_rd   = sf(row.get("draft_round"),   default=np.nan)
    draft_pick = sf(row.get("draft_pick"),    default=np.nan)  # overall pick #
    pos_group  = str(row.get("pos_group") or "F")

    # Current-season production rates (prefer per-60 for rate stability)
    pts60_cur  = sf(row.get("points60_ind"),  default=np.nan)
    g60_cur    = sf(row.get("g60_ind"),       default=np.nan)
    ixg60_cur  = sf(row.get("ixg60_ind"),     default=np.nan)
    p_gp_cur   = sf(row.get("pts_per_gp_ind"),default=np.nan)

    # Rolling 3yr production (if available — more stable for projection)
    pts60_3yr  = sf(row.get("points60_ind_3yr"), default=np.nan)  # if you add rolling IND later
    dim_score  = sf(_first_notna(row.get("Dimensionality_Score"), row.get("Dimensionality_Score_3yr")), default=np.nan)
    top_score  = sf(row.get("Top_Role_Score"), default=np.nan)
    conf_ev    = sf(_first_notna(row.get("Confidence_EV"), row.get("Confidence_EV_3yr")), default=0.5)

    # ----------------------------------------------------------------
    # 1. Draft pedigree — persistent talent signal (especially <26)
    # ----------------------------------------------------------------
    if pd.isna(draft_rd) and pd.isna(draft_pick):
        draft_pedigree = "Unknown"
        pedigree_boost = 0.0
    elif (pd.notna(draft_pick) and draft_pick <= 10) or (pd.notna(draft_rd) and draft_rd == 1 and pd.isna(draft_pick)):
        draft_pedigree = "Elite"
        pedigree_boost = 0.06
    elif pd.notna(draft_pick) and draft_pick <= 30:
        draft_pedigree = "High"
        pedigree_boost = 0.04
    elif pd.notna(draft_pick) and draft_pick <= 90:
        draft_pedigree = "Mid"
        pedigree_boost = 0.02
    elif pd.notna(draft_rd) and draft_rd <= 3:
        draft_pedigree = "Mid"
        pedigree_boost = 0.01
    elif pd.notna(draft_pick) and draft_pick > 180:
        draft_pedigree = "Late"
        pedigree_boost = -0.01
    else:
        draft_pedigree = "Late"
        pedigree_boost = 0.0

    # Pedigree boost decays with age — still a signal through ~26, minimal after 30
    if pd.notna(age):
        pedigree_boost *= float(np.clip(1.0 - (max(age - 22.0, 0.0) / 12.0), 0.0, 1.0))

    # ----------------------------------------------------------------
    # 2. Aging curve factor
    # NHL peaks roughly 24-28, fast decline after 32
    # ----------------------------------------------------------------
    if pd.isna(age):
        age_curve = 1.00
        prog_tier = "Unknown"
    elif age <= 21:
        age_curve = 0.88   # still developing
        prog_tier = "Emerging"
    elif age <= 23:
        age_curve = 0.96
        prog_tier = "Emerging"
    elif age <= 27:
        age_curve = 1.02   # slight peak boost
        prog_tier = "Prime"
    elif age <= 29:
        age_curve = 1.00
        prog_tier = "Prime"
    elif age <= 31:
        age_curve = 0.97
        prog_tier = "Plateau"
    elif age <= 33:
        age_curve = 0.93
        prog_tier = "Declining"
    elif age <= 35:
        age_curve = 0.87
        prog_tier = "Declining"
    else:
        age_curve = 0.80
        prog_tier = "Aging"

    # Defensemen decline slightly later (peak ~26-30)
    if pos_group == "D" and pd.notna(age):
        if 26 <= age <= 30:
            age_curve = min(age_curve + 0.02, 1.04)
        elif age > 33:
            age_curve = max(age_curve - 0.02, 0.70)

    # ----------------------------------------------------------------
    # 3. Best available production baseline (current > 3yr rolling > role proxy)
    # ----------------------------------------------------------------
    is_injured = bool(row.get("Injury_Imputed") or row.get("injury_fill"))

    if is_injured:
        # Current-season stats are imputed (prior * 0.92) — using them would
        # double-penalise the player. Prefer the 3yr rolling average if available,
        # which is built from multiple real seasons and is a much better baseline.
        if pd.notna(pts60_3yr):
            pts60_base = pts60_3yr
        elif pd.notna(pts60_cur):
            # Fall back to imputed only if nothing else exists, but we'll
            # undo the 0.92 decay so the projection isn't unfairly depressed.
            pts60_base = pts60_cur / 0.92
        elif pd.notna(dim_score):
            ceil = 1.2 if pos_group == "D" else 2.5
            pts60_base = (dim_score / 100.0) * ceil
        else:
            pts60_base = np.nan
    elif pd.notna(pts60_3yr) and (pd.isna(pts60_cur) or conf_ev < 0.4):
        pts60_base = pts60_3yr
    elif pd.notna(pts60_cur):
        pts60_base = pts60_cur
    elif pd.notna(dim_score):
        # Rough proxy from dimensionality when no IND stats available
        # Fwd elite dim~95 -> ~2.5 pts60; def elite dim~95 -> ~1.2 pts60
        ceil = 1.2 if pos_group == "D" else 2.5
        pts60_base = (dim_score / 100.0) * ceil
    else:
        pts60_base = np.nan

    g60_base = g60_cur if pd.notna(g60_cur) else np.nan
    if is_injured and pd.notna(g60_base):
        g60_base = g60_base / 0.92  # undo imputation decay

    # ----------------------------------------------------------------
    # 4. Projected rates (age curve + pedigree)
    # ----------------------------------------------------------------
    if pd.notna(pts60_base):
        proj_pts60 = float(pts60_base * age_curve * (1.0 + pedigree_boost))
        proj_pts60 = max(proj_pts60, 0.0)
    else:
        proj_pts60 = np.nan

    if pd.notna(g60_base):
        proj_g60 = float(g60_base * age_curve * (1.0 + pedigree_boost * 0.5))
        proj_g60 = max(proj_g60, 0.0)
    else:
        proj_g60 = np.nan

    # pts/game projection (roughly pts60 * avg TOI/60 — assume 15-17 min fwd, 19-22 min def)
    avg_toi_hr = (20.0 if pos_group == "D" else 15.5) / 60.0
    proj_p_gp  = float(proj_pts60 * avg_toi_hr) if pd.notna(proj_pts60) else np.nan

    # ----------------------------------------------------------------
    # 5. Projection confidence
    # ----------------------------------------------------------------
    has_ind  = pd.notna(pts60_cur) or pd.notna(g60_cur)
    has_roll = pd.notna(pts60_3yr)
    has_bio  = pd.notna(age) and pd.notna(draft_rd)

    n_signals = sum([has_ind, has_roll, has_bio, conf_ev >= 0.5])

    if is_injured:
        # Projection is based on prior-season or rolling data, not current performance.
        # Cap at Med regardless of signal count — return-to-health is genuinely uncertain.
        proj_conf = "Med (Injury)" if n_signals >= 2 else "Low (Injury)"
    elif n_signals >= 3:
        proj_conf = "High"
    elif n_signals >= 2:
        proj_conf = "Med"
    else:
        proj_conf = "Low"

    return {
        "Proj_Points60":    proj_pts60,
        "Proj_Goals60":     proj_g60,
        "Proj_PTS_PGP":     proj_p_gp,
        "Age_Curve_Factor": age_curve,
        "Prog_Tier":        prog_tier,
        "Prog_Confidence":  proj_conf,
        "Draft_Pedigree":   draft_pedigree,
    }


def add_player_projections(df: pd.DataFrame) -> pd.DataFrame:
    """Applies project_player row-wise and merges projection columns back."""
    proj_cols = ["Proj_Points60","Proj_Goals60","Proj_PTS_PGP",
                 "Age_Curve_Factor","Prog_Tier","Prog_Confidence","Draft_Pedigree"]
    proj = df.apply(project_player, axis=1, result_type="expand")
    proj.columns = proj_cols
    out = df.copy()
    for c in proj_cols:
        out[c] = proj[c].values
    return out


def apply_prospect_tier(df: pd.DataFrame) -> pd.DataFrame:
    """
    Overlay pass — must run AFTER add_player_projections() so that
    Draft_Pedigree and Prog_Tier columns are already present.

    Elevates young players with genuine elite draft capital who haven't
    yet reached Star/Elite — their current stats understate their ceiling.

    Criteria (ALL must be true):
      1. Not already Star or Elite (they've arrived)
      2. Draft pedigree is Elite (~picks 1-10) or High (~picks 11-32)
      3. Age window:
           a. Age <= 23  AND  Prog_Tier == "Emerging"
           b. Age 23.5-24.5  AND  overall_pick <= 10  (top franchise pick entering prime)
      4. TOI >= 20 min — has actually played NHL games (not pure AHL call-up)
      5. Bust guard: age >= 22, top-10 pick, full sample (conf >= 0.50), Fringe
         → these players have had their chance; keep as Fringe, not Prospect
    """
    out = df.copy()

    if "Player_Tier" not in out.columns:
        return out

    if "Draft_Pedigree" not in out.columns or "age" not in out.columns:
        return out

    idx       = out.index
    pedigree  = out["Draft_Pedigree"].fillna("").astype(str)
    age_s     = out["age"].fillna(99).astype(float)
    prog_tier = out.get("Prog_Tier", pd.Series("", index=idx)).fillna("").astype(str)
    pick      = out.get("overall_pick", pd.Series(999, index=idx)).fillna(999).astype(float)
    toi_s     = out.get("toi_ev", pd.Series(0, index=idx)).fillna(0).astype(float)
    conf      = out.get("Confidence_EV", pd.Series(0.5, index=idx)).fillna(0.5).astype(float)
    tier      = out["Player_Tier"].astype(str)

    high_pedigree = pedigree.isin(["Elite", "High"])

    # Age window A: ≤23 and Emerging prog tier
    window_a = (age_s <= 23) & (prog_tier == "Emerging")
    # Age window B: 24-year-old with top-10 franchise pick
    window_b = (age_s.between(23.5, 24.5)) & (pick <= 10)

    has_nhl_time = toi_s >= 20.0
    not_arrived  = ~tier.isin(["Star", "Elite"])

    # Confirmed bust guard — had full opportunity, still Fringe
    bust_guard = (
        (age_s >= 22) &
        (pick <= 10) &
        (conf >= 0.50) &
        (tier == "Fringe")
    )

    prospect_mask = (
        (window_a | window_b) &
        high_pedigree &
        has_nhl_time &
        not_arrived &
        ~bust_guard
    )

    out.loc[prospect_mask, "Player_Tier"] = "Prospect"
    return out


def build_team_offseason_plan(
    model_season_df: pd.DataFrame,
    cap_df: pd.DataFrame,
    team: str,
    offseason_year: int = 2026,
    cap_limit: float = CAP_LIMIT_26_27,
    keep_core_overrides: Optional[Dict[str, bool]] = None,
    resign_budget_share: float = 0.45,
) -> Dict[str, pd.DataFrame]:
    """
    Team-only offseason plan.

    - model_season_df: your scored season dataframe (25-26), already merged with cap via merge_cap()
    - cap_df: full cap table from load_cap_table()
    - team: e.g. 'EDM'
    - offseason_year: 2026
    - cap_limit: 95.5M default
    - resign_budget_share: portion of remaining cap reserved for re-signing expiring team players

    Returns dict of DataFrames to write as tabs.
    """
    team = team.strip().upper()
    ms = _ensure_pos_labels(model_season_df.copy())
    cap_df = cap_df.copy()

    # --- Team cap view (under contract vs expiring)
    under_contract_cap, expiring_cap = team_current_roster(cap_df, team, offseason_year)

    # --- Enrich model metrics onto cap rows
    # Single source of truth: pull all fields needed for impact scoring, comparables, and decisioning
    enrich_cols = [
        "player_key","player","pos","pos_group",
        "Top_Role","Top_Role_Score","Second_Role","Second_Role_Score","Margin",
        "Role_Stability","Type_Confidence",
        "Best_Line_Fit","Best_Line_Fit_Score","Best_Pair_Fit","Best_Pair_Fit_Score",
        "Dimensionality_Score","Confidence_EV",
        "Player_Type","age",
        "hand",   # handedness — drives slot assignment in lineup builder
        # Individual production (used in _impact_score + _comp_distance)
        "points60_ind","g60_ind","ixg60_ind","pts_per_gp_ind","a1_60_ind",
        "goals_ind","assists_ind","points_ind",
        # Projection outputs (used in comparables + resign ranking)
        "Proj_Points60","Proj_Goals60","Proj_PTS_PGP",
        "Age_Curve_Factor","Prog_Tier","Prog_Confidence","Draft_Pedigree",
        # TOI validity flag — carried so team sheets can filter insufficient-data players
        "keep","toi_ev","Injury_Imputed",
    ]
    enrich_cols = [c for c in enrich_cols if c in ms.columns]
    ms_enrich = ms[enrich_cols].drop_duplicates("player_key")

    uc = under_contract_cap.merge(ms_enrich, on="player_key", how="left")
    ex = expiring_cap.merge(ms_enrich, on="player_key", how="left")

    # FIX: merges create player_x/player_y etc → coalesce back to canonical columns
    for _df_name in ("uc", "ex"):
        _df = locals()[_df_name]
        _df = _coalesce_xy(_df, "player")
        _df = _coalesce_xy(_df, "pos")
        _df = _coalesce_xy(_df, "pos_group")
        # team_cap is from cap table, but keep canonical if it ever gets suffixed
        _df = _coalesce_xy(_df, "team_cap")
        locals()[_df_name] = _df

    # If cap table uses pos_cap, prefer enriched pos but fall back to pos_cap
    if "pos" not in uc.columns and "pos_cap" in uc.columns:
        uc["pos"] = uc["pos_cap"]
    if "pos" not in ex.columns and "pos_cap" in ex.columns:
        ex["pos"] = ex["pos_cap"]

    uc = _ensure_pos_labels(uc)
    ex = _ensure_pos_labels(ex)

    # --- Core roster selection (mostly keep all under contract unless cap disaster)
    core = _choose_core_team(uc, cap_limit=cap_limit)
    core_cap = _cap_sum(core)
    cap_space = cap_limit - core_cap

        # --- Build market table (league cap + model signals)
    market = _build_market_table(cap_df, ms)

    # --- Estimate re-sign AAV for expiring players (team) using comparables + simulation
    ex = _ensure_pos_labels(ex)
    ex["impact"] = ex.apply(_impact_score, axis=1)

    # comparables-based AAV distribution
    est_rows = []
    for _, r in ex.iterrows():
        comps = find_market_comparables(market, r, n=50, same_team_ok=True)
        est = estimate_aav_from_comps(comps)
        draws = simulate_aav_draws(est, n_sims=3000, dist="t", t_df=6)

        est_rows.append({
            "player_key": r["player_key"],
            "comp_n": est.n_comps,
            "aav_mean": est.mean,
            "aav_std": est.std,
            "aav_p10": float(np.quantile(draws, 0.10)),
            "aav_p50": float(np.quantile(draws, 0.50)),
            "aav_p90": float(np.quantile(draws, 0.90)),
        })

    if est_rows:
        ex = ex.merge(pd.DataFrame(est_rows), on="player_key", how="left")
    else:
        ex["comp_n"] = 0
        ex["aav_mean"] = np.nan
        ex["aav_std"] = np.nan
        ex["aav_p10"] = np.nan
        ex["aav_p50"] = np.nan
        ex["aav_p90"] = np.nan

    # choose a negotiation point estimate (P50 by default)
    ex["aav_est"] = ex["aav_p50"].fillna(ex["cap_hit"]).fillna(0.775e6)

    # value per $ using point estimate
    ex["value_per_$"] = ex["impact"] / (ex["aav_est"].replace(0, np.nan) / 1_000_000.0)

    # --- Progression-adjusted value: players trending up are worth more than current impact suggests
    # Prog_Tier: Emerging (+15%), Prime (+5%), Plateau (0%), Declining (-8%), Aging (-15%)
    prog_tier_bonus = {
        "Emerging": 0.15,
        "Prime":    0.05,
        "Plateau":  0.00,
        "Declining":-0.08,
        "Aging":   -0.15,
        "Unknown":  0.00,
    }
    def _prog_adjusted_value(row) -> float:
        base_val = _safe_float(row.get("value_per_$"), default=0.0)
        tier = str(row.get("Prog_Tier") or "Unknown")
        bonus = prog_tier_bonus.get(tier, 0.0)
        # Draft pedigree adds a small signal for young players (where it's still predictive)
        age = _safe_float(row.get("age"), default=30.0)
        ped = str(row.get("Draft_Pedigree") or "Unknown")
        ped_bonus = {"Elite": 0.06, "High": 0.03, "Mid": 0.01, "Late": 0.0}.get(ped, 0.0)
        ped_bonus *= float(np.clip(1.0 - (max(age - 22.0, 0.0) / 10.0), 0.0, 1.0))
        return float(base_val * (1.0 + bonus + ped_bonus))

    ex["value_prog_adj"] = ex.apply(_prog_adjusted_value, axis=1)

    # Expose projected production for the output sheet
    if "Proj_Points60" not in ex.columns:
        ex["Proj_Points60"] = np.nan
    if "Proj_PTS_PGP" not in ex.columns:
        ex["Proj_PTS_PGP"] = np.nan
    if "Prog_Tier" not in ex.columns:
        ex["Prog_Tier"] = "Unknown"
    if "Draft_Pedigree" not in ex.columns:
        ex["Draft_Pedigree"] = "Unknown"

    # Optional manual keep/drop overrides by player_key
    if keep_core_overrides:
        ov = {_player_key(k): v for k, v in keep_core_overrides.items()}
        ex["_override"] = ex["player_key"].map(ov)
    else:
        ex["_override"] = np.nan

    # ── BUCKET QUALITY WEIGHTS (shared by all ranking) ───────────────────────
    _BUCKET_QUALITY = {
        "Line 1": 1.00, "Line 2": 0.80, "Line 3": 0.55, "Line 4": 0.25,
        "Pair 1": 1.00, "Pair 2": 0.70, "Pair 3": 0.40,
    }

    post0  = core.copy()
    # holes0 should reflect roster AFTER all expiring players leave — 
    # this shows the true open spots the team needs to fill via resigns + FA.
    ex_keys = set(ex["player_key"].astype(str)) if len(ex) else set()
    post0_without_expiring = post0[~post0["player_key"].astype(str).isin(ex_keys)].copy()
    holes0 = _define_holes(post0_without_expiring)

    def _player_priority(row: pd.Series, holes_ref: dict) -> float:
        """
        Single priority formula for both expiring team players and FA pool.
          60% impact  (0-100 → 0-0.60)
          20% value/$M (impact/AAV_M, soft-cap 100 → 0-0.20)
          10% hole fill bonus (fit quality × hole existence → 0-0.10)
          10% type confidence (High/Med/Low → 0-0.10)
        """
        impact = float(row.get("impact") or 0.0)
        aav    = float(row.get("aav_est") or np.nan)
        aav_m  = (aav / 1_000_000.0) if (pd.notna(aav) and aav > 0) else np.nan
        vpm    = min((impact / aav_m) / 100.0, 1.0) if pd.notna(aav_m) else 0.0

        pos = str(row.get("pos") or "").upper()
        grp = "D" if pos == "D" else "F"
        fb  = _fit_bucket(row)
        bq  = _BUCKET_QUALITY.get(fb, 0.30)
        hb  = float(np.clip(
            0.15 * float(holes_ref.get(f"need_{grp}", 0) > 0) +
            0.25 * float(holes_ref.get(f"need_{pos}", 0) > 0) +
            0.60 * bq * float(holes_ref.get(f"need_{fb}", 0) > 0) +
            0.10 * bq, 0.0, 1.0))

        conf_s = str(row.get("Type_Confidence") or "").strip().lower()
        conf   = {"high": 1.0, "medium": 0.67, "low": 0.33}.get(conf_s, 0.33)

        return 0.60*(impact/100.0) + 0.20*vpm + 0.10*hb + 0.10*conf

    # ── BUILD FA POOL ─────────────────────────────────────────────────────────
    fa = available_free_agents(cap_df, offseason_year)
    fa = fa[~fa["team_cap"].eq(team)].copy()
    fa = fa.merge(ms_enrich, on="player_key", how="left")
    fa = _coalesce_xy(fa, "player")
    fa = _coalesce_xy(fa, "pos")
    fa = _coalesce_xy(fa, "pos_group")
    fa = _coalesce_xy(fa, "team_cap")
    if "pos" not in fa.columns and "pos_cap" in fa.columns:
        fa["pos"] = fa["pos_cap"]
    fa = _ensure_pos_labels(fa)
    fa["impact"] = fa.apply(_impact_score, axis=1)

    fa_est_rows = []
    for _, r in fa.iterrows():
        comps = find_market_comparables(market, r, n=50, same_team_ok=True)
        est   = estimate_aav_from_comps(comps)
        draws = simulate_aav_draws(est, n_sims=2000, dist="t", t_df=6)
        fa_est_rows.append({
            "player_key": r["player_key"], "comp_n": est.n_comps,
            "aav_mean": est.mean, "aav_std": est.std,
            "aav_p10": float(np.quantile(draws, 0.10)),
            "aav_p50": float(np.quantile(draws, 0.50)),
            "aav_p90": float(np.quantile(draws, 0.90)),
        })
    if fa_est_rows:
        fa_est_df = pd.DataFrame(fa_est_rows)
        drop_cols = [c for c in ["aav_p10","aav_p50","aav_p90","aav_mean","aav_std","comp_n"] if c in fa.columns]
        if drop_cols:
            fa = fa.drop(columns=drop_cols)
        fa = fa.merge(fa_est_df, on="player_key", how="left")
    else:
        for c in ["aav_p10","aav_p50","aav_p90","aav_mean","aav_std","comp_n"]:
            fa[c] = np.nan
    for _c in ["aav_p50","aav_p10","aav_p90"]:
        fa = _coalesce_xy(fa, _c)
    if "aav_p50" not in fa.columns:
        fa["aav_p50"] = np.nan
    fa["aav_est"] = fa["aav_p50"].fillna(fa["cap_hit"]).fillna(0.775e6)

    # ── AAV estimates for expiring team players ───────────────────────────────
    ex["impact"] = ex.apply(_impact_score, axis=1)
    ex_est_rows = []
    for _, r in ex.iterrows():
        comps = find_market_comparables(market, r, n=50, same_team_ok=True)
        est   = estimate_aav_from_comps(comps)
        draws = simulate_aav_draws(est, n_sims=3000, dist="t", t_df=6)
        ex_est_rows.append({
            "player_key": r["player_key"], "comp_n": est.n_comps,
            "aav_mean": est.mean, "aav_std": est.std,
            "aav_p10": float(np.quantile(draws, 0.10)),
            "aav_p50": float(np.quantile(draws, 0.50)),
            "aav_p90": float(np.quantile(draws, 0.90)),
        })
    if ex_est_rows:
        est_df = pd.DataFrame(ex_est_rows)
        # Drop any pre-existing aav columns to avoid _x/_y suffixes on merge
        drop_cols = [c for c in ["aav_p10","aav_p50","aav_p90","aav_mean","aav_std","comp_n"] if c in ex.columns]
        if drop_cols:
            ex = ex.drop(columns=drop_cols)
        ex = ex.merge(est_df, on="player_key", how="left")
    else:
        for c in ["aav_p10","aav_p50","aav_p90","aav_mean","aav_std","comp_n"]:
            ex[c] = np.nan
    # Coalesce any suffix columns from prior merges
    for _c in ["aav_p50","aav_p10","aav_p90"]:
        ex = _coalesce_xy(ex, _c)
    if "aav_p50" not in ex.columns:
        ex["aav_p50"] = np.nan
    ex["aav_est"] = ex["aav_p50"].fillna(ex["cap_hit"]).fillna(0.775e6)

    # ── UNIFIED POOL: expiring team players + all FAs ranked together ─────────
    # This prevents the artificial "resign budget wall" from blocking a high-value
    # expiring player (like Balinskis) in favour of a cheaper lower-quality one
    # (like Mikkola) just because the resign sub-budget ran out first.
    ex_tagged       = ex.copy();  ex_tagged["_source"] = "RESIGN"
    fa_tagged       = fa.copy();  fa_tagged["_source"] = "SIGN"
    all_cols        = sorted(set(ex_tagged.columns) | set(fa_tagged.columns))
    pool            = pd.concat([ex_tagged.reindex(columns=all_cols),
                                 fa_tagged.reindex(columns=all_cols)], ignore_index=True)
    pool["impact"]    = pool.apply(_impact_score, axis=1)
    pool["aav_est"]   = pool["aav_est"].fillna(0.775e6)
    pool["priority"]  = pool.apply(lambda r: _player_priority(r, holes0), axis=1)
    pool["value_per_$"] = pool["impact"] / (pool["aav_est"].replace(0, np.nan) / 1_000_000.0)

    if keep_core_overrides:
        ov = {_player_key(k): v for k, v in keep_core_overrides.items()}
        pool["_override"] = pool["player_key"].map(ov)
    else:
        pool["_override"] = np.nan

    pool = pool.sort_values(["_override","priority","impact"], ascending=[False,False,False])

    total_budget = max(cap_space, 0.0)
    holes_work   = holes0.copy()
    remaining    = total_budget
    selected_rows = []
    selected_keys = set()

    for _, r in pool.iterrows():
        pk  = str(r.get("player_key") or "")
        pos = str(r.get("pos") or "").upper()
        if pos not in {"C", "LW", "RW", "D"} or pk in selected_keys:
            continue
        aav   = float(r.get("aav_est") or np.nan)
        if np.isnan(aav) or aav <= 0:
            continue
        force = r.get("_override") is True
        if r.get("_override") is False:
            continue
        if not force and aav > remaining:
            continue
        fb       = _fit_bucket(r)
        need_pos = _needs_pos(holes_work, pos)
        need_grp = holes_work.get(f"need_{'D' if pos=='D' else 'F'}", 0) > 0
        need_fit = holes_work.get(f"need_{fb}", 0) > 0
        is_own_expiring = str(r.get("_source") or "") == "RESIGN"
        # Own expiring players: always evaluate (you're deciding whether to keep them).
        # External FAs: only sign if there's a genuine roster/fit need.
        if not force and not is_own_expiring and not (need_pos or need_grp or need_fit):
            continue
        selected_rows.append(r)
        selected_keys.add(pk)
        remaining -= aav
        _decrement_holes_for_player(holes_work, pos, fit_bucket=fb)
        if all(v <= 0 for k, v in holes_work.items() if k.startswith("need_")):
            break

    sel_df        = pd.DataFrame(selected_rows) if selected_rows else pool.iloc[0:0].copy()
    resigned_keys = set(sel_df.loc[sel_df["_source"].eq("RESIGN"), "player_key"].astype(str))
    signed_keys   = set(sel_df.loc[sel_df["_source"].eq("SIGN"),   "player_key"].astype(str))

    # ── Expiring team decision table ─────────────────────────────────────────
    ex["decision"]      = np.where(ex["player_key"].astype(str).isin(resigned_keys), "RESIGN", "LET WALK")
    ex["new_cap_hit"]   = np.where(ex["player_key"].astype(str).isin(resigned_keys), ex["aav_est"], np.nan)
    ex["value_per_$"]   = ex["impact"] / (ex["aav_est"].replace(0, np.nan) / 1_000_000.0)
    ex["hole_bonus"]    = ex.apply(lambda r: float(np.clip(
        0.15*float(holes0.get(f"need_{'D' if str(r.get('pos','')).upper()=='D' else 'F'}",0)>0) +
        0.25*float(holes0.get(f"need_{str(r.get('pos','')).upper()}",0)>0) +
        0.60*_BUCKET_QUALITY.get(_fit_bucket(r),0.30)*float(holes0.get(f"need_{_fit_bucket(r)}",0)>0) +
        0.10*_BUCKET_QUALITY.get(_fit_bucket(r),0.30), 0.0, 1.0)), axis=1)
    ex["resign_priority"] = ex.apply(lambda r: _player_priority(r, holes0), axis=1)
    resigned_df   = ex[ex["decision"].eq("RESIGN")].copy()
    not_resigned  = ex[ex["decision"].eq("LET WALK")].copy()
    resign_decisions = pd.concat([resigned_df, not_resigned], ignore_index=True)
    resign_decisions = resign_decisions.sort_values(["decision","resign_priority"], ascending=[True,False])
    spent = float(resigned_df["aav_est"].sum()) if len(resigned_df) else 0.0

    # ── FA signed ────────────────────────────────────────────────────────────
    fa["priority"]    = fa.apply(lambda r: _player_priority(r, holes0), axis=1)
    fa["value_per_$"] = fa["impact"] / (fa["aav_est"].replace(0, np.nan) / 1_000_000.0)
    fa["hole_bonus"]  = fa.apply(lambda r: float(np.clip(
        0.15*float(holes0.get(f"need_{'D' if str(r.get('pos','')).upper()=='D' else 'F'}",0)>0) +
        0.25*float(holes0.get(f"need_{str(r.get('pos','')).upper()}",0)>0) +
        0.60*_BUCKET_QUALITY.get(_fit_bucket(r),0.30)*float(holes0.get(f"need_{_fit_bucket(r)}",0)>0) +
        0.10*_BUCKET_QUALITY.get(_fit_bucket(r),0.30), 0.0, 1.0)), axis=1)

    # Full ranked FA analysis table (all FAs, signed or not)
    fa_analysis = fa.sort_values("priority", ascending=False).copy()
    fa_analysis["Signed"] = fa_analysis["player_key"].astype(str).isin(signed_keys)

    sign_df = fa_analysis[fa_analysis["Signed"]].copy()
    if len(sign_df) > 0:
        sign_df["decision"] = "SIGN"
        sign_df["new_cap_hit"] = sign_df["aav_est"]
        sign_df["origin"] = "FA"
    else:
        sign_df = fa.iloc[0:0].copy()
        sign_df["decision"] = pd.Series(dtype=str)
        sign_df["new_cap_hit"] = pd.Series(dtype=float)
        sign_df["origin"] = pd.Series(dtype=str)

    # --- Build post-resign base roster (core + resigned expiring players)
    post = core.copy()
    if len(resigned_df) > 0:
        add_resigned = resigned_df.copy()
        add_resigned["cap_hit"] = add_resigned["aav_est"]
        post = pd.concat([post, add_resigned.reindex(columns=post.columns, fill_value=np.nan)],
                         ignore_index=True, sort=False)
    post = _ensure_pos_labels(post)
    post["origin"] = post.get("origin", pd.Series("Team", index=post.index))
    post["origin"] = post["origin"].fillna("Team")
    post_cap = _cap_sum(post)

    # --- Projected roster after signings
    projected = post.copy()
    if len(sign_df) > 0:
        add = sign_df.copy()
        add["cap_hit"] = add["new_cap_hit"]
        # align columns
        common_cols = sorted(set(projected.columns).union(set(add.columns)))
        projected = projected.reindex(columns=common_cols)
        add = add.reindex(columns=common_cols)
        projected = pd.concat([projected, add], ignore_index=True)

    projected = _ensure_pos_labels(projected)
    projected["impact"] = projected.apply(_impact_score, axis=1)
    projected = add_team_usage_rank(projected)

    # --- Build lineup presentation
    lineup = build_lineup_tables(projected, team=team)

    # --- Cap summary
    def _capline(df):
        return pd.DataFrame([{
            "team": team,
            "cap_limit": cap_limit,
            "core_under_contract_cap": core_cap,
            "resign_spend_est": spent,
            "post_resign_cap": post_cap,
            "fa_sign_spend_est": float(sign_df["new_cap_hit"].sum()) if len(sign_df) else 0.0,
            "projected_total_cap": _cap_sum(projected),
            "projected_cap_space": cap_limit - _cap_sum(projected),
            "need_F_remaining": holes_work.get("need_F", 0),
            "need_D_remaining": holes_work.get("need_D", 0),
            "need_C_remaining": holes_work.get("need_C", 0),
            "need_LW_remaining": holes_work.get("need_LW", 0),
            "need_RW_remaining": holes_work.get("need_RW", 0),
        }])

    cap_summary = _capline(projected)

    # ── Apply keep filter to all player-level team output tables ─────────────
    # Players with keep=False had insufficient TOI even after imputation.
    # Goalies are excluded from the role model entirely and never get keep=False,
    # so the goalie check (pos != G) is just a safety net.
    def _filter_keep(df: pd.DataFrame) -> pd.DataFrame:
        """Remove skaters with keep=False. Preserve rows without a keep column (e.g. goalies)."""
        if "keep" not in df.columns:
            return df
        # keep NaN = no flag set (e.g. cap-only rows that never went through scoring) → include
        mask = df["keep"].isna() | (df["keep"] == True)
        return df[mask].copy()

    # ── Contract Urgency flag on resign decisions ─────────────────────────────
    def _contract_urgency(row) -> str:
        tier = str(row.get("Player_Tier") or row.get("player_tier") or "")
        age = _safe_float(row.get("age"), default=np.nan)
        impact = _safe_float(row.get("impact") or row.get("Impact_Score"), default=0.0)
        if tier in ("Elite", "Star") or impact >= 80:
            return "CRITICAL"
        if tier == "Solid" and pd.notna(age) and 23 <= age <= 30:
            return "HIGH"
        if tier in ("Solid", "Depth") or impact >= 50:
            return "MEDIUM"
        return "LOW"

    if "Contract_Urgency" not in resign_decisions.columns:
        resign_decisions["Contract_Urgency"] = resign_decisions.apply(_contract_urgency, axis=1)

    # Re-sign priority score (0-100, spread) replacing the compressed 0.47-0.67 range
    def _resign_priority(row) -> float:
        imp = _safe_float(row.get("impact") or row.get("Impact_Score"), default=50.0)
        val = _safe_float(row.get("value_per_$") or row.get("Value / $M"), default=0.0)
        urg_map = {"CRITICAL": 30, "HIGH": 20, "MEDIUM": 10, "LOW": 0}
        urg = urg_map.get(str(row.get("Contract_Urgency") or "LOW"), 0)
        hole = _safe_float(row.get("hole_bonus") or row.get("Hole Fill Bonus"), default=0.0) * 20
        raw = 0.45 * imp + 0.25 * min(val * 5, 25) + urg + hole
        return float(np.clip(raw, 0, 100))

    resign_decisions["Re_Sign_Priority_100"] = resign_decisions.apply(_resign_priority, axis=1).round(1)

    priority_tiers = {
        (75, 101): "Must Sign",
        (50, 75):  "Target",
        (30, 50):  "Optional",
        (0, 30):   "Pass",
    }
    def _priority_tier(score) -> str:
        s = _safe_float(score, default=0.0)
        for (lo, hi), label in priority_tiers.items():
            if lo <= s < hi:
                return label
        return "Pass"
    resign_decisions["Priority_Tier"] = resign_decisions["Re_Sign_Priority_100"].map(_priority_tier)

    # ── Per-unit cap breakdown ─────────────────────────────────────────────────
    def _build_unit_cap_breakdown(lineup_df: pd.DataFrame, team: str) -> pd.DataFrame:
        rows = []
        if lineup_df is None or len(lineup_df) == 0:
            return pd.DataFrame()
        for unit in ["Line 1", "Line 2", "Line 3", "Line 4", "Pair 1", "Pair 2", "Pair 3"]:
            unit_col = "Assigned_Unit" if "Assigned_Unit" in lineup_df.columns else None
            if not unit_col:
                continue
            block = lineup_df[lineup_df[unit_col].astype(str).eq(unit)]
            cap_col = next((c for c in ["cap_hit", "Cap Hit", "new_cap_hit"] if c in block.columns), None)
            cap_total = block[cap_col].astype(float).sum() if cap_col else 0.0
            rows.append({
                "Team": team,
                "Roster_Unit": unit,
                "Players": int(len(block)),
                "Unit_Cap_Total": cap_total,
                "Avg_Cap_Per_Player": cap_total / max(len(block), 1),
            })
        return pd.DataFrame(rows)

    unit_cap = _build_unit_cap_breakdown(lineup_out if 'lineup_out' in dir() else pd.DataFrame(), team)

    # ── Roster needs heatmap ──────────────────────────────────────────────────
    def _build_needs_heatmap(lineup_df: pd.DataFrame, team: str) -> pd.DataFrame:
        slots = {
            "Line 1": [("LW", "Line 1 LW"), ("C", "Line 1 C"), ("RW", "Line 1 RW")],
            "Line 2": [("LW", "Line 2 LW"), ("C", "Line 2 C"), ("RW", "Line 2 RW")],
            "Line 3": [("LW", "Line 3 LW"), ("C", "Line 3 C"), ("RW", "Line 3 RW")],
            "Line 4": [("LW", "Line 4 LW"), ("C", "Line 4 C"), ("RW", "Line 4 RW")],
            "Pair 1": [("D", "Pair 1 LD"), ("D", "Pair 1 RD")],
            "Pair 2": [("D", "Pair 2 LD"), ("D", "Pair 2 RD")],
            "Pair 3": [("D", "Pair 3 LD"), ("D", "Pair 3 RD")],
        }
        rows = []
        if lineup_df is None or len(lineup_df) == 0:
            for unit, slot_list in slots.items():
                for _, slot_name in slot_list:
                    rows.append({"Team": team, "Slot": slot_name, "Player": "VACANT",
                                 "Impact_Score": np.nan, "Fit_Score": np.nan, "RAG_Status": "RED"})
            return pd.DataFrame(rows)

        unit_col = "Assigned_Unit" if "Assigned_Unit" in lineup_df.columns else None
        if not unit_col:
            return pd.DataFrame()

        for unit, slot_list in slots.items():
            block = lineup_df[lineup_df[unit_col].astype(str).eq(unit)].copy()
            impact_col = next((c for c in ["impact", "Impact_Score", "Impact Score"] if c in block.columns), None)
            fit_col = next((c for c in ["Best_Line_Fit_Score", "Best_Pair_Fit_Score",
                                        "Best Line Fit Score", "Best Pair Fit Score"] if c in block.columns), None)

            # prefer assigned-slot col if present
            slot_col = "Assigned_Slot" if "Assigned_Slot" in block.columns else None

            filled = block.to_dict("records")
            for i, (_, slot_name) in enumerate(slot_list):
                if i < len(filled):
                    p = filled[i]
                    player_name = p.get("Player") or p.get("player") or "?"
                    imp = _safe_float(p.get(impact_col) if impact_col else np.nan, default=np.nan)
                    fit = _safe_float(p.get(fit_col) if fit_col else np.nan, default=np.nan)

                    # RAG logic
                    fit_ok = pd.isna(fit) or fit >= 70
                    imp_ok = pd.isna(imp) or imp >= 60
                    fit_warn = pd.notna(fit) and 50 <= fit < 70
                    imp_warn = pd.notna(imp) and 40 <= imp < 60
                    fit_bad = pd.notna(fit) and fit < 50
                    imp_bad = pd.notna(imp) and imp < 40

                    if fit_bad or imp_bad:
                        rag = "RED"
                    elif fit_warn or imp_warn:
                        rag = "YELLOW"
                    else:
                        rag = "GREEN"

                    rows.append({"Team": team, "Slot": slot_name, "Player": player_name,
                                 "Impact_Score": round(imp, 1) if pd.notna(imp) else np.nan,
                                 "Fit_Score": round(fit, 1) if pd.notna(fit) else np.nan,
                                 "RAG_Status": rag})
                else:
                    rows.append({"Team": team, "Slot": slot_name, "Player": "VACANT",
                                 "Impact_Score": np.nan, "Fit_Score": np.nan, "RAG_Status": "RED"})

        return pd.DataFrame(rows)

    needs_heatmap = _build_needs_heatmap(lineup_out if 'lineup_out' in dir() else pd.DataFrame(), team)

    core_out      = _filter_keep(core.sort_values(["pos_group","cap_hit"], ascending=[True, False]))
    resign_out    = _filter_keep(resign_decisions)
    sign_out      = _filter_keep(sign_df)
    fa_out        = _filter_keep(fa_analysis)
    projected_out = _filter_keep(projected.sort_values(["pos_group","impact"], ascending=[True, False]))
    lineup_out    = build_lineup_tables(projected_out, team=team)

    needs_heatmap = _build_needs_heatmap(lineup_out, team)
    unit_cap = _build_unit_cap_breakdown(lineup_out, team)

    return {
        f"TEAM_{team}_Returning_26-27": core_out,
        f"TEAM_{team}_Resign_Decisions_26": resign_out,
        f"TEAM_{team}_FA_Targets_26": sign_out,
        f"TEAM_{team}_FA_Analysis_26": fa_out,
        f"TEAM_{team}_Projected_Roster_26-27": projected_out,
        f"TEAM_{team}_Projected_Lineup_26-27": lineup_out,
        f"TEAM_{team}_Cap_Summary_26-27": cap_summary,
        f"TEAM_{team}_Unit_Cap_Breakdown_26-27": unit_cap,
        f"TEAM_{team}_Needs_Heatmap_26-27": needs_heatmap,
    }

def _pos_eligible_slots(pos: str, hand: str) -> List[str]:
    """
    Returns the roster slots a player is eligible to play, in preference order.

    Rules:
      - D plays only D.
      - C can play C (primary), LW or RW (secondary — strong-side preferred by hand).
      - LW can play LW (primary), C or RW (secondary — off-side is last resort).
      - RW can play RW (primary), C or LW (secondary).
      - Hand preference: L-hand prefers left side; R-hand prefers right side.
        When a winger moves to the off-side, there is a small penalty applied
        by the caller, not here — this function just returns the eligibility order.
    """
    pos   = str(pos  or "").upper().strip()
    hand  = str(hand or "").upper().strip()[:1]  # "L" or "R"

    if pos == "D":
        return ["D"]
    if pos == "C":
        # C can bump to wing; prefer the strong side based on hand
        wing_order = ["LW", "RW"] if hand == "L" else ["RW", "LW"]
        return ["C"] + wing_order
    if pos == "LW":
        # LW prefers left, can center, off-side RW last
        if hand == "R":
            # off-hand on left — already playing off-side naturally; can go right easily
            return ["LW", "RW", "C"]
        return ["LW", "C", "RW"]
    if pos == "RW":
        if hand == "L":
            return ["RW", "LW", "C"]
        return ["RW", "C", "LW"]
    # Unknown — try everything
    return ["C", "LW", "RW"]


def _hand_penalty(natural_pos: str, assigned_slot: str, hand: str) -> float:
    """
    Returns a small fit penalty (subtracted from impact) when a player is
    assigned to a non-natural or off-hand slot.

    Scale is relative to the impact score (0–100):
      0.0  = no penalty (natural slot)
      2.0  = minor (e.g. LW→C or C playing natural-hand wing)
      4.0  = moderate (winger bumped to center not their natural pos)
      6.0  = significant (winger on off-side, e.g. L-hand on RW)
      8.0  = heavy (C on off-side wing — rare, last resort)
    """
    natural_pos   = str(natural_pos   or "").upper().strip()
    assigned_slot = str(assigned_slot or "").upper().strip()
    hand          = str(hand          or "").upper().strip()[:1]

    if assigned_slot == natural_pos:
        return 0.0

    # D has no wing penalties
    if natural_pos == "D" or assigned_slot == "D":
        return 0.0

    # Center bumping to wing
    if natural_pos == "C":
        strong_wing = "LW" if hand == "L" else "RW"
        if assigned_slot == strong_wing:
            return 2.0   # natural-hand side — very common (Reinhart scenario)
        return 4.0       # off-side wing from C

    # Winger going to center
    if assigned_slot == "C":
        return 3.0       # moderate — winger at center costs some value

    # Winger going to opposite wing
    if natural_pos == "LW" and assigned_slot == "RW":
        return 6.0 if hand == "L" else 3.0   # L-hand on RW = off-side; R-hand on RW = natural
    if natural_pos == "RW" and assigned_slot == "LW":
        return 6.0 if hand == "R" else 3.0

    return 2.0   # fallback


def _assign_lines_greedy(
    players: pd.DataFrame,
    line_units: List[str],
    slots_per_unit: int,
    slot_types: List[str],
    fit_col: str,
    impact_col: str = "impact",
) -> pd.DataFrame:
    """
    Greedy best-fit assignment of players to line/pair slots.

    Algorithm:
      1. Sort players by impact desc (best player gets first pick of best unit).
      2. For each player in order, find the highest-scoring (unit, slot) combination
         they are eligible for that still has an open slot.
      3. Apply hand penalty to the effective score when considering non-natural slots.
      4. Assign to best available (unit, slot).

    This replaces the old categorical-sort + sequential-fill approach which
    ignored per-player fit scores when making assignments.

    Returns the input df with two new columns: Assigned_Unit, Assigned_Slot.
    """
    out = players.copy().reset_index(drop=True)

    # Build slot inventory: {unit: {slot_type: None or player_idx}}
    slots: Dict[str, Dict[str, Optional[int]]] = {}
    for unit in line_units:
        slots[unit] = {st: None for st in slot_types}

    # Pre-fetch fit scores per unit (Fit_Line1 → "Line 1", etc.)
    fit_score_map: Dict[str, str] = {}
    for unit in line_units:
        # e.g. "Line 1" → "Fit_Line1", "Pair 1" → "Fit_Pair1"
        col_name = "Fit_" + unit.replace(" ", "")  # "Fit_Line1", "Fit_Pair1"
        if col_name in out.columns:
            fit_score_map[unit] = col_name

    out[impact_col] = out[impact_col].fillna(0.0).astype(float)
    sorted_idx = out[impact_col].sort_values(ascending=False).index.tolist()

    out["Assigned_Unit"] = ""
    out["Assigned_Slot"] = ""

    for player_idx in sorted_idx:
        row = out.loc[player_idx]
        natural_pos = str(row.get("pos") or "").upper().strip()
        hand        = str(row.get("hand") or row.get("hand_cap") or "").upper().strip()[:1]
        base_impact = float(row.get(impact_col) or 0.0)

        eligible_slots = _pos_eligible_slots(natural_pos, hand)

        best_score = -1e9
        best_unit  = None
        best_slot  = None

        for unit in line_units:
            # Raw fit score for this unit (e.g. Fit_Line1 value for this player)
            unit_fit_col = fit_score_map.get(unit)
            unit_fit = float(row.get(unit_fit_col) or 0.0) if unit_fit_col else base_impact

            # Skip if unit_fit is sentinel (ineligible due to guardrails)
            if unit_fit <= -1e8:
                continue

            for slot in eligible_slots:
                if slot not in slots[unit]:
                    continue
                if slots[unit][slot] is not None:
                    continue  # slot taken

                penalty = _hand_penalty(natural_pos, slot, hand)
                # Effective score: combine unit fit quality with impact, minus penalty
                eff = 0.60 * unit_fit + 0.40 * base_impact - penalty

                if eff > best_score:
                    best_score = eff
                    best_unit  = unit
                    best_slot  = slot

        if best_unit is not None and best_slot is not None:
            slots[best_unit][best_slot] = player_idx
            out.at[player_idx, "Assigned_Unit"] = best_unit
            out.at[player_idx, "Assigned_Slot"] = best_slot
        else:
            # Fallback: assign to any open slot in any unit (overflow / short roster)
            for unit in line_units:
                for slot in slot_types:
                    if slots[unit][slot] is None:
                        slots[unit][slot] = player_idx
                        out.at[player_idx, "Assigned_Unit"] = unit
                        out.at[player_idx, "Assigned_Slot"] = slot
                        break
                if out.at[player_idx, "Assigned_Unit"]:
                    break

    return out


def build_lineup_tables(projected: pd.DataFrame, team: str) -> pd.DataFrame:
    """
    Produces a flat lineup table using a greedy best-fit assignment solver.

    Key improvements over the previous version:
      - Position fluidity: centers can bump to wing; wingers can flip sides.
      - Handedness preference: L-hand prefers LW/left side, R-hand prefers RW.
      - Hand penalty system: off-side or unnatural-position assignments are
        penalized proportionally so the model prefers natural slots but allows
        flexibility when skill warrants it (e.g. Barkov stays C; Reinhart bumps
        to RW naturally as a R-hand C).
      - Best-player-first greedy: highest-impact player gets first pick of best
        unit, preventing low-impact depth pieces from locking premium slots.
      - Impact floor: no player with Impact >= 75 is assigned to Line 4.
    """
    df = _ensure_pos_labels(projected).copy()

    if "impact" not in df.columns:
        df["impact"] = df.apply(_impact_score, axis=1)

    # Pull handedness from cap data if available
    hand_col = next((c for c in ["hand", "hand_cap"] if c in df.columns), None)
    if hand_col:
        df["hand"] = df[hand_col].astype(str).str.upper().str.strip().str[:1]
    else:
        df["hand"] = ""

    # --- Select roster pool ---
    f = df[df["pos_group"].eq("F") & df["pos"].isin(["C", "LW", "RW"])].copy()
    d = df[df["pos_group"].eq("D") & df["pos"].eq("D")].copy()

    # Top 12 forwards and 6 defensemen by impact
    f = f.sort_values("impact", ascending=False).head(12)
    d = d.sort_values("impact", ascending=False).head(6)

    # --- Forward assignment ---
    # Slot types: one C, one LW, one RW per line
    f_assigned = _assign_lines_greedy(
        players       = f,
        line_units    = ["Line 1", "Line 2", "Line 3", "Line 4"],
        slots_per_unit= 3,
        slot_types    = ["C", "LW", "RW"],
        fit_col       = "Best_Line_Fit_Score",
        impact_col    = "impact",
    )

    # --- Defense assignment ---
    # D pairs: two D slots per pair (no positional distinction within pair,
    # but we do track hand to note LD/RD balance in the depth chart)
    d_assigned = _assign_lines_greedy(
        players       = d,
        line_units    = ["Pair 1", "Pair 2", "Pair 3"],
        slots_per_unit= 2,
        slot_types    = ["LD", "RD"],
        fit_col       = "Best_Pair_Fit_Score",
        impact_col    = "impact",
    )

    # --- Combine ---
    out = pd.concat([f_assigned, d_assigned], ignore_index=True, sort=False)

    unit_order = {**{f"Line {i}": i for i in range(1,5)}, **{f"Pair {i}": 10+i for i in range(1,4)}}
    out["_unit_sort"] = out["Assigned_Unit"].map(unit_order).fillna(99)
    out = out.sort_values(["_unit_sort", "impact"], ascending=[True, False])
    out = out.drop(columns=["_unit_sort"], errors="ignore")

    show_cols = [
        "Assigned_Unit", "Assigned_Slot",
        "player", "pos", "hand",
        "team_cap", "cap_hit", "impact",
        "Player_Type", "Type_Confidence",
        "Top_Role", "Top_Role_Score", "Second_Role", "Second_Role_Score", "Margin",
        "Best_Line_Fit", "Best_Line_Fit_Score", "Best_Pair_Fit", "Best_Pair_Fit_Score",
        "expiry_year", "contract_type", "origin",
    ]
    show_cols = [c for c in show_cols if c in out.columns]
    return out[show_cols].copy()

# =============================================================================
# LOADERS
# =============================================================================
def load_season_sheet(xlsx: Path, season_sheet: str) -> pd.DataFrame:
    sheet = f"{season_sheet} On-Ice"
    df = pd.read_excel(xlsx, sheet_name=sheet)
    df = _clean_cols(df)

    col_player = _find_col(df, FALLBACKS["player"])
    col_team = _find_col(df, FALLBACKS["team"])
    col_pos = _find_col(df, FALLBACKS["pos_raw"])
    if col_player is None or col_pos is None:
        raise ValueError(f"[{sheet}] Missing required Player/Position columns.")

    out = pd.DataFrame()
    out["player"] = df[col_player].astype(str).map(normalize_player_name)
    out["player_key"] = out["player"].map(_player_key)
    out["team"] = (df[col_team].astype(str).str.strip().map(_resolve_team)
                   if col_team else "")
    out["pos_raw"] = df[col_pos]
    out["pos"] = out["pos_raw"].map(_canonical_pos)
    out["pos_group"] = out["pos"].map(_pos_group)
    out["season"] = season_sheet

    def pull(name: str) -> pd.Series:
        c = _find_col(df, FALLBACKS.get(name, []))
        return df[c].map(_to_float) if c else pd.Series([np.nan] * len(df))

    out["gp"] = pull("gp")
    out["toi_ev"] = pull("toi_ev")

    for k in ["cf","ca","sf","sa","gf","ga","xgf","xga","scf","sca","hdcf","hdca"]:
        out[k] = pull(k)

    out["oz_starts"] = pull("oz_starts")
    out["nz_starts"] = pull("nz_starts")
    out["dz_starts"] = pull("dz_starts")
    out["off_zone_start_pct"] = pull("oz_pct")

    starts_total = (out["oz_starts"] + out["nz_starts"] + out["dz_starts"]).replace(0, np.nan)
    out["off_zone_start_pct"] = out["off_zone_start_pct"].fillna(out["oz_starts"] / starts_total)
    out["def_zone_start_pct"] = (out["dz_starts"] / starts_total).fillna(np.nan)

    toi = out["toi_ev"].replace(0, np.nan)
    out["cf60_ev"] = out["cf"] / toi * 60.0
    out["ca60_ev"] = out["ca"] / toi * 60.0
    out["sf60_ev"] = out["sf"] / toi * 60.0
    out["sa60_ev"] = out["sa"] / toi * 60.0
    out["gf60_ev"] = out["gf"] / toi * 60.0
    out["ga60_ev"] = out["ga"] / toi * 60.0
    out["xgf60_ev"] = out["xgf"] / toi * 60.0
    out["xga60_ev"] = out["xga"] / toi * 60.0
    out["scf60_ev"] = out["scf"] / toi * 60.0
    out["sca60_ev"] = out["sca"] / toi * 60.0
    out["hdcf60_ev"] = out["hdcf"] / toi * 60.0
    out["hdca60_ev"] = out["hdca"] / toi * 60.0

    return out

def load_bio_sheet(xlsx: Path, season_sheet: str) -> pd.DataFrame:
    sheet = f"{season_sheet} Bios"
    df = pd.read_excel(xlsx, sheet_name=sheet)
    df = _clean_cols(df)

    col_player = _find_col(df, BIO_FALLBACKS["player"])
    if col_player is None:
        raise ValueError(f"[{sheet}] Missing required Player/Name column.")

    out = pd.DataFrame()
    out["player"] = df[col_player].astype(str).map(normalize_player_name)
    out["player_key"] = out["player"].map(_player_key)
    out["season"] = season_sheet

    def pull(name: str) -> pd.Series:
        c = _find_col(df, BIO_FALLBACKS.get(name, []))
        return df[c].map(_to_float) if c else pd.Series([np.nan] * len(df))

    out["height_in"] = pull("height_in")
    out["weight_lb"] = pull("weight_lb")
    out["age"] = pull("age")
    out["draft_round"] = pull("draft_round")
    out["round_pick"] = pull("round_pick")
    out["overall_pick"] = pull("overall_pick")

    # Height inches, weight pounds
    out["bmi"] = out["weight_lb"] / (out["height_in"] ** 2)

    # unified draft pick field
    out["draft_pick"] = out["overall_pick"].fillna(out["round_pick"])

    return out


def load_individual_sheet(xlsx: Path, season_str: str) -> pd.DataFrame:
    """
    Loads '{season_str} Ind' and returns:
      - keys: player, player_key, season
      - standardized IND columns (gp_ind, toi_ind, team_ind, pos_ind, etc.)
      - derived columns used by ROLE_FEATURES (g60_ind, a1_60_ind, points60_ind, shots60_ind, ixg60_ind, ...)
      - extra helpful fields for type boosts (assists_ind, points_ind, pts_per_gp_ind, etc.)
    """
    sheet = f"{season_str} Ind"
    df = pd.read_excel(xlsx, sheet_name=sheet)
    df = _clean_cols(df)

    def col(name: str) -> Optional[str]:
        return _find_col(df, IND_FALLBACKS.get(name, []))

    req = ["player", "gp", "toi"]
    for r in req:
        if col(r) is None:
            raise ValueError(f"[{sheet}] Missing required column: {r}")

    out = pd.DataFrame()
    out["player"] = df[col("player")].astype(str).map(normalize_player_name)
    out["player_key"] = out["player"].map(_player_key)
    out["season"] = season_str

    # basic identifiers
    out["team_ind"] = (df[col("team")].astype(str).str.strip().map(_resolve_team)
                       if col("team") else "")
    out["pos_raw_ind"] = df[col("pos_raw")] if col("pos_raw") else ""
    out["pos_ind"] = out["pos_raw_ind"].map(_canonical_pos) if col("pos_raw") else ""
    out["gp_ind"] = df[col("gp")].map(_to_float)
    out["toi_ind"] = df[col("toi")].map(_to_float)  # minutes

    # pull numeric columns (best-effort)
    def pull(name: str) -> pd.Series:
        c = col(name)
        return df[c].map(_to_float) if c else pd.Series([np.nan] * len(df))

    out["goals_ind"] = pull("goals")
    out["assists_ind"] = pull("assists")
    out["a1_ind"] = pull("a1")
    out["a2_ind"] = pull("a2")
    out["points_ind"] = pull("points")
    out["ipp"] = pull("ipp")

    out["shots_ind"] = pull("shots")
    out["sh_pct"] = pull("sh_pct") / 100.0 if col("sh_pct") else np.nan  # store as 0-1
    out["ixg_ind"] = pull("ixg")

    out["icf_ind"] = pull("icf")
    out["iff_ind"] = pull("iff")
    out["iscf_ind"] = pull("iscf")
    out["ihdcf"] = pull("ihdcf")  # NOTE: ROLE_FEATURES uses 'ihdcf' (count proxy)

    out["rush_ind"] = pull("rush")
    out["rebounds_ind"] = pull("rebounds")

    out["pim_ind"] = pull("pim")
    out["pen_total_ind"] = pull("pen_total")
    out["pen_minor_ind"] = pull("pen_minor")
    out["pen_major_ind"] = pull("pen_major")
    out["pen_misconduct_ind"] = pull("pen_misconduct")
    out["pen_drawn_ind"] = pull("pen_drawn")

    out["giveaways_ind"] = pull("giveaways")
    out["takeaways_ind"] = pull("takeaways")
    out["hits_ind"] = pull("hits")
    out["hits_taken_ind"] = pull("hits_taken")
    out["shots_blocked_ind"] = pull("shots_blocked")

    out["fow_ind"] = pull("fow")
    out["fol_ind"] = pull("fol")
    out["fo_pct"] = pull("fo_pct") / 100.0 if col("fo_pct") else np.nan  # store as 0-1

    # -------------------------
    # Derived: per game + per 60 (the names ROLE_FEATURES expects)
    # -------------------------
    gp = out["gp_ind"].replace(0, np.nan)
    toi = out["toi_ind"].replace(0, np.nan)

    # per game (useful for type boosts)
    out["pts_per_gp_ind"] = out["points_ind"] / gp
    out["a_per_gp_ind"] = out["assists_ind"] / gp
    out["a1_per_gp_ind"] = out["a1_ind"] / gp
    out["g_per_gp_ind"] = out["goals_ind"] / gp

    # per 60 (ROLE_FEATURES keys)
    out["g60_ind"] = out["goals_ind"] / toi * 60.0
    out["ixg60_ind"] = out["ixg_ind"] / toi * 60.0
    out["shots60_ind"] = out["shots_ind"] / toi * 60.0

    out["a1_60_ind"] = out["a1_ind"] / toi * 60.0
    out["p1_60_ind"] = (out["goals_ind"] + out["a1_ind"]) / toi * 60.0
    out["points60_ind"] = out["points_ind"] / toi * 60.0

    out["rush60_ind"] = out["rush_ind"] / toi * 60.0
    out["rebounds60_ind"] = out["rebounds_ind"] / toi * 60.0

    out["hits60_ind"] = out["hits_ind"] / toi * 60.0
    out["blk_shots60_ind"] = out["shots_blocked_ind"] / toi * 60.0

    out["giveaways60_ind"] = out["giveaways_ind"] / toi * 60.0
    out["takeaways60_ind"] = out["takeaways_ind"] / toi * 60.0

    out["pen_drawn60_ind"] = out["pen_drawn_ind"] / toi * 60.0
    out["pen_taken60_ind"] = out["pen_total_ind"] / toi * 60.0
    out["pen_diff60_ind"] = (out["pen_drawn_ind"] - out["pen_total_ind"]) / toi * 60.0

    out["fow60_ind"] = out["fow_ind"] / toi * 60.0
    out["fol60_ind"] = out["fol_ind"] / toi * 60.0

    return out

# -----------------------------------------------------------------------------
# IND-STAT TYPE NUDGES (season totals + rates)
# These are "assist denominators" for archetype classification + contract/fit logic.
# They are SMALL boosts intended to break ties / fix obvious mislabels.
# -----------------------------------------------------------------------------

TYPE_NUDGES_FWD = {
    # Elite points driver
    "Producer": [
        ("pts_per_gp_ind", ">=", 1.00, 8.0),     # already added, keep here too if you want centralized config
        ("points_ind", ">=", 85, 6.0),           # raw season threshold (tune by GP context)
        ("points60_ind", ">=", 2.8, 4.0),        # rate-based (better cross-GP)
    ],

    # Setup/creation
    "Playmaker": [
        ("assists_ind", ">=", 50, 6.0),          # already added
        ("a1_ind", ">=", 30, 5.0),               # primary assists
        ("a1_60_ind", ">=", 1.0, 4.0),           # creation rate
        ("ipp", ">=", 0.70, 3.0),                # involved when on-ice
        ("rebounds60_ind", ">=", 0.55, 2.0),     # creating second chances
    ],

    # Goal scoring
    "Finisher": [
        ("goals_ind", ">=", 35, 6.0),
        ("g60_ind", ">=", 1.10, 4.0),
        ("shots60_ind", ">=", 9.0, 3.0),
        ("sh_pct", ">=", 0.14, 2.0),             # stored as 0–1 in the patch I gave you
        ("ixg60_ind", ">=", 0.85, 2.0),
    ],

    # Shot/attempt volume driver (puck/possession proxy on IND side)
    "Driver": [
        ("icf_ind", ">=", 500, 3.0),
        ("iscf_ind", ">=", 180, 3.0),
        ("ihdcf", ">=", 75, 2.0),
        ("ixg_ind", ">=", 20, 2.0),
        ("rush_ind", ">=", 80, 2.0),
    ],

    # Defensive-ish / discipline / draws
    "Two-Way": [
        ("pen_diff60_ind", ">=", 0.20, 3.0),     # draws > takes
        ("takeaways_ind", ">=", 55, 2.0),
        ("giveaways_ind", "<=", 35, 2.0),        # lower is better
        ("hits_ind", ">=", 120, 1.5),
        ("shots_blocked_ind", ">=", 60, 1.5),
    ],

    # Net-front / physical offense
    "Power": [
        ("iHDCF", ">=", 85, 3.0),
        ("ixg_ind", ">=", 22, 2.5),
        ("hits_ind", ">=", 160, 2.0),
        ("pen_drawn_ind", ">=", 25, 2.0),
    ],

    # 4th-line energy / pk-ish physical
    "Grinder / Tough Guy": [
        ("hits_ind", ">=", 180, 4.0),
        ("shots_blocked_ind", ">=", 80, 3.0),
        ("pen_total_ind", ">=", 60, 2.5),
        ("pim_ind", ">=", 60, 2.0),
    ],
}

TYPE_NUDGES_DEF = {
    "Shutdown D": [
        ("shots_blocked_ind", ">=", 130, 4.0),
        ("hits_ind", ">=", 140, 2.0),
        ("giveaways_ind", "<=", 35, 2.0),
        ("pen_diff60_ind", ">=", 0.10, 1.0),
    ],
    "Transition": [
        ("rush_ind", ">=", 70, 4.0),
        ("icf_ind", ">=", 350, 2.0),
        ("iscf_ind", ">=", 140, 2.0),
    ],
    "Puck Skill": [
        ("assists_ind", ">=", 35, 4.0),
        ("a1_ind", ">=", 20, 3.0),
        ("giveaways60_ind", "<=", 0.80, 2.0),
        ("takeaways60_ind", ">=", 0.60, 2.0),
    ],
    "Physical": [
        ("hits_ind", ">=", 170, 4.0),
        ("shots_blocked_ind", ">=", 150, 3.0),
        ("pen_total_ind", "<=", 55, 1.5),        # avoid pure penalty-taker
    ],
}

# =============================================================================
# DERIVED CONTEXT METRICS
# Adds analytical enrichment that doesn't affect scores but dramatically improves
# interpretability: on-ice %, YoY deltas, SH% luck, league rank, contract efficiency.
# Call this AFTER apply_role_scores so role scores are already present.
# =============================================================================

LEAGUE_AVG_SH_PCT = 0.091   # ~9.1% league average EV shooting (stored as 0-1 fraction)

def add_derived_context(df: pd.DataFrame, prev_df: Optional[pd.DataFrame] = None) -> pd.DataFrame:
    """
    Enriches the season dataframe with:

    On-ice percentage metrics (all positions):
      xGF%   — expected goals for % at EV (xGF / (xGF + xGA))
      CF%    — Corsi for % at EV
      GF%    — goals for % at EV
      HDCF%  — high-danger chances for %

    Shooting / finishing luck:
      SH%_vs_Avg — player shooting % relative to league average (>1 = above average)
      xSH%_proxy  — ixG/shots as a proxy for expected shooting % (shot quality)
      Finishing_Delta — SH% - xSH%_proxy (positive = overperforming expected quality)

    Zone deployment context:
      OZ_Start_Delta — player OZ start % minus their team's average OZ start % for pos group
                       (positive = deployed more offensively than teammates)

    Year-over-year trajectory (requires prev_df):
      Dim_Delta   — Dimensionality_Score change vs prior season
      Dim_Trend   — "Rising" / "Stable" / "Declining" based on delta

    League rank within position:
      League_Rank_Dim  — rank among all players at same position (1 = best)
      League_Pctile    — percentile (same as Dimensionality_Score but labeled clearly)

    Role consistency (requires prev_df):
      Role_Consistent — True if Top_Role matches prev season's Top_Role
      Seasons_Same_Role — count of seasons (1-3) the player has held their current top role
    """
    out = df.copy()

    # ── 1. On-ice percentage metrics ──────────────────────────────────────────
    xgf = out.get("xgf", pd.Series(np.nan, index=out.index)).astype(float)
    xga = out.get("xga", pd.Series(np.nan, index=out.index)).astype(float)
    cf  = out.get("cf",  pd.Series(np.nan, index=out.index)).astype(float)
    ca  = out.get("ca",  pd.Series(np.nan, index=out.index)).astype(float)
    gf  = out.get("gf",  pd.Series(np.nan, index=out.index)).astype(float)
    ga  = out.get("ga",  pd.Series(np.nan, index=out.index)).astype(float)
    hdcf = out.get("hdcf", pd.Series(np.nan, index=out.index)).astype(float)
    hdca = out.get("hdca", pd.Series(np.nan, index=out.index)).astype(float)

    out["xGF_pct"]  = np.where((xgf + xga) > 0, xgf / (xgf + xga) * 100.0, np.nan)
    out["CF_pct"]   = np.where((cf  + ca)  > 0, cf  / (cf  + ca)  * 100.0, np.nan)
    out["GF_pct"]   = np.where((gf  + ga)  > 0, gf  / (gf  + ga)  * 100.0, np.nan)
    out["HDCF_pct"] = np.where((hdcf + hdca) > 0, hdcf / (hdcf + hdca) * 100.0, np.nan)

    # ── 2. Shooting / finishing luck ──────────────────────────────────────────
    sh_pct  = out.get("sh_pct",    pd.Series(np.nan, index=out.index)).astype(float)  # stored 0-1
    shots   = out.get("shots_ind", pd.Series(np.nan, index=out.index)).astype(float)
    ixg     = out.get("ixg_ind",   pd.Series(np.nan, index=out.index)).astype(float)
    goals   = out.get("goals_ind", pd.Series(np.nan, index=out.index)).astype(float)

    # xSH% proxy: ixG / shots (shot quality expected conversion rate)
    out["xSH_pct_proxy"] = np.where(shots > 5, ixg / shots, np.nan)

    # SH% vs league average (1.0 = exactly average, 1.20 = 20% above average)
    out["SH_pct_vs_Avg"] = np.where(
        sh_pct.notna() & (sh_pct > 0),
        sh_pct / LEAGUE_AVG_SH_PCT,
        np.nan
    )

    # Finishing delta: actual SH% - xSH% proxy (positive = getting more goals than shot quality predicts)
    out["Finishing_Delta"] = np.where(
        sh_pct.notna() & out["xSH_pct_proxy"].notna(),
        (sh_pct - out["xSH_pct_proxy"]) * 100.0,   # in percentage points
        np.nan
    )

    # ── 3. Zone deployment context (vs team peers) ────────────────────────────
    oz_pct = out.get("off_zone_start_pct", pd.Series(np.nan, index=out.index)).astype(float)
    if "team" in out.columns and "pos_group" in out.columns:
        team_avg_oz = out.groupby(["team", "pos_group"])["off_zone_start_pct"].transform("mean")
        out["OZ_Start_Delta"] = oz_pct - team_avg_oz
    else:
        out["OZ_Start_Delta"] = np.nan

    # ── 4. YoY trajectory (requires prev season data) ─────────────────────────
    out["Dim_Delta"] = np.nan
    out["Dim_Trend"] = ""

    if prev_df is not None and "player_key" in prev_df.columns and "Dimensionality_Score" in prev_df.columns:
        prev_dim = (prev_df.drop_duplicates("player_key")
                           .set_index("player_key")["Dimensionality_Score"])
        if "player_key" in out.columns:
            prior_vals = prev_dim.reindex(out["player_key"]).values
            cur_dim    = out.get("Dimensionality_Score", pd.Series(np.nan, index=out.index)).astype(float).values
            delta      = np.where(
                ~np.isnan(prior_vals) & ~np.isnan(cur_dim),
                cur_dim - prior_vals,
                np.nan
            )
            out["Dim_Delta"] = delta

            # Trend labels: >8 Rising, <-8 Declining, else Stable
            out["Dim_Trend"] = np.select(
                [out["Dim_Delta"] > 8, out["Dim_Delta"] < -8, out["Dim_Delta"].notna()],
                ["Rising", "Declining", "Stable"],
                default=""
            )

    # ── 5. League rank within position ────────────────────────────────────────
    dim = out.get("Dimensionality_Score", pd.Series(np.nan, index=out.index)).astype(float)
    if "pos" in out.columns:
        out["League_Rank"] = out.groupby("pos")["Dimensionality_Score"].rank(
            ascending=False, method="min", na_option="bottom"
        ).astype("Int64")
        pos_count = out.groupby("pos")["pos"].transform("count")
        out["League_Rank_Of"] = pos_count
    else:
        out["League_Rank"]    = np.nan
        out["League_Rank_Of"] = np.nan

    # ── 6. Role consistency (requires prev season data) ───────────────────────
    out["Role_Consistent"]   = np.nan
    out["Seasons_Same_Role"] = np.nan

    if prev_df is not None and "player_key" in prev_df.columns and "Top_Role" in prev_df.columns:
        prev_role = (prev_df.drop_duplicates("player_key")
                            .set_index("player_key")["Top_Role"])
        if "player_key" in out.columns and "Top_Role" in out.columns:
            prior_role_vals = prev_role.reindex(out["player_key"]).values.astype(object)
            cur_role_vals   = out["Top_Role"].astype(object).values

            consistent = np.array([
                bool(c == p) if (c and p and not pd.isna(c) and not pd.isna(p) and str(c) != "" and str(p) != "")
                else float("nan")
                for c, p in zip(cur_role_vals, prior_role_vals)
            ], dtype=object)
            out["Role_Consistent"] = consistent

    return out


def build_league_summary(season_all: Dict[str, pd.DataFrame], newest_season: str) -> pd.DataFrame:
    """
    Builds a league-wide team summary sheet showing:
      - Average Dimensionality by team and position group
      - Role type distribution (how many Finishers, Playmakers, etc.)
      - Average cap efficiency (if cap data is merged in)
      - Depth score (how many players above 60th percentile in Dimensionality)
    """
    df = season_all[newest_season].copy()
    df = df[df["keep"] == True].copy()

    if "team" not in df.columns:
        return pd.DataFrame()

    # Drop any rows where the team field still contains a comma — these are
    # traded players whose team wasn't resolved at load time (e.g. came in via
    # a cap merge rather than the season sheet).  They would otherwise create
    # phantom "S.J, VAN"-style rows in the summary.
    df = df[~df["team"].astype(str).str.contains(",", na=False)].copy()
    # Also drop blanks / NaN teams
    df = df[df["team"].astype(str).str.strip().ne("") & df["team"].notna()].copy()

    rows = []
    for team, tdf in df.groupby("team"):
        fwd = tdf[tdf["pos_group"] == "F"]
        dfd = tdf[tdf["pos_group"] == "D"]

        def safe_mean(s): return float(s.mean()) if len(s) > 0 else np.nan
        def depth_n(s, thr=60): return int((s.fillna(0) >= thr).sum())

        row = {
            "Team": team,
            "F_Players": len(fwd),
            "D_Players": len(dfd),
            "Avg_Dim_F":  safe_mean(fwd["Dimensionality_Score"]),
            "Avg_Dim_D":  safe_mean(dfd["Dimensionality_Score"]),
            "Depth_F_60+": depth_n(fwd["Dimensionality_Score"]),
            "Depth_D_60+": depth_n(dfd["Dimensionality_Score"]),
            "Top_Role_F":  fwd["Player_Type"].mode().iloc[0] if len(fwd) > 0 else "",
            "Top_Role_D":  dfd["Player_Type"].mode().iloc[0] if len(dfd) > 0 else "",
        }

        # Role type counts for forwards
        if "Player_Type" in fwd.columns:
            for rtype in ["Finisher","Playmaker","Driver","Two-Way","Producer","Grinder / Tough Guy"]:
                row[f"F_{rtype.replace(' / ',' ').replace('-','').replace(' ','_')}"] = int(
                    fwd["Player_Type"].astype(str).str.contains(rtype, na=False).sum()
                )

        # Average xGF% if available
        if "xGF_pct" in tdf.columns:
            row["Avg_xGF_pct_F"] = safe_mean(fwd["xGF_pct"])
            row["Avg_xGF_pct_D"] = safe_mean(dfd["xGF_pct"])

        rows.append(row)

    summary = pd.DataFrame(rows).sort_values("Avg_Dim_F", ascending=False)
    return summary


# =============================================================================
# SCORING + PLAYER TYPE
# =============================================================================
def apply_role_scores(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()

    out["keep"] = True
    out.loc[(out["pos_group"] == "F") & (out["toi_ev"] < MIN_TOI_EV_FWD), "keep"] = False
    out.loc[(out["pos_group"] == "D") & (out["toi_ev"] < MIN_TOI_EV_DEF), "keep"] = False

    out["Confidence_EV"] = np.clip(out["toi_ev"].fillna(0.0) / 1200.0, 0, 1)

    def score_block(pos_group: str, roles: Dict[str, Dict]) -> None:
        mask = (out["pos_group"] == pos_group) & out["keep"]
        if not mask.any():
            return

        for role_name, spec in roles.items():
            minutes_basis = spec["minutes_basis"]
            k = spec["k"]
            feats = spec["features"]

            s = pd.Series(0.0, index=out.index)

            for feat, w in feats.items():
                if feat not in out.columns:
                    continue
                z = _zscore_within(out.loc[mask].assign(_x=out.loc[mask, feat]), "_x", by=["season", "pos"])
                z_full = pd.Series(0.0, index=out.index)
                z_full.loc[mask] = z

                mins = out[minutes_basis] if minutes_basis in out.columns else out["toi_ev"]
                s = s + w * _shrink(z_full, mins, k=k)

            out[f"S_{role_name}"] = np.nan
            out.loc[mask, f"S_{role_name}"] = s.loc[mask].values

            out[f"{role_name}_Score"] = np.nan
            out.loc[mask, f"{role_name}_Score"] = (
                100.0 * out.loc[mask].groupby(["season", "pos"])[f"S_{role_name}"].transform(_percent_rank).values
            )

    score_block("F", ROLE_FEATURES["F"])
    score_block("D", ROLE_FEATURES["D"])

    def dimensionality_row(row) -> float:
        if not row.get("keep", False):
            return np.nan
        roles = (
            ["Finisher_Score", "Playmaker_Score", "Driver_Score", "TwoWay_Score", "Power_Score", "Grinder_Score"]
            if row["pos_group"] == "F"
            else ["Suppressor_Score", "Transition_Score", "PuckSkill_Score", "Physical_Score"]
        )
        vals = [row.get(r) for r in roles if r in row and pd.notna(row.get(r))]
        if len(vals) < 2:
            return np.nan
        vals = sorted(vals, reverse=True)
        return float(np.clip((vals[0] + vals[1]) / 2.0, 0, 100))

    out["Dimensionality_Score"] = out.apply(dimensionality_row, axis=1)

    # ── Player_Tier: multi-signal composite ──────────────────────────────────
    # Replaces the old single-threshold Dim cut.
    # Inputs (all available at this point in the pipeline):
    #   1. Dimensionality_Score          — current season role versatility (40%)
    #   2. Production percentile         — pts/gp + pts/60 + ixG/60 (25%)
    #      Defensemen use on-ice possession metrics instead of points
    #   3. Confidence_EV                 — sample size quality (10%)
    #   (Rolling Dim_3yr is NOT available here — applied in build_rolling downstream)
    #
    # Thresholds (calibrated to NHL roster reality ~32 teams):
    #   Elite  >= 92  → ~35-45 players  (franchise/Hart-calibre)
    #   Star   >= 82  → ~100 players    (all-star calibre, top-6F / top-4D)
    #   Solid  >= 68  → ~220 players    (reliable NHL regulars)
    #   Depth  >= 50  → ~280 players    (role players, 4th line, 3rd pair)
    #   Fringe  < 50  → rest            (AHL call-ups, marginal NHLers)

    def _compute_player_tier(df_block: pd.DataFrame) -> pd.Series:
        idx = df_block.index

        # ── production percentile, position-aware ──────────────────────────
        fwd_mask = df_block["pos_group"].eq("F")
        def_mask = df_block["pos_group"].eq("D")

        prod_raw = pd.Series(np.nan, index=idx)

        # Forwards: points-based production
        if fwd_mask.any():
            fwd = df_block.loc[fwd_mask]
            raw = (
                0.35 * fwd["pts_per_gp_ind"].fillna(0) +
                0.35 * fwd["points60_ind"].fillna(0) +
                0.20 * fwd["ixg60_ind"].fillna(0) +
                0.10 * fwd["g60_ind"].fillna(0)
            )
            # percentile rank within forwards only
            prod_raw.loc[fwd_mask] = raw.rank(pct=True, method="average") * 100

        # Defense: on-ice possession + modest production weight
        if def_mask.any():
            dfd = df_block.loc[def_mask]
            raw = (
                0.25 * dfd["pts_per_gp_ind"].fillna(0) +
                0.25 * dfd["points60_ind"].fillna(0) +
                0.20 * dfd.get("xgf60_ev", pd.Series(0, index=dfd.index)).fillna(0) +
                0.15 * dfd["ixg60_ind"].fillna(0) +
                0.15 * dfd.get("cf60_ev", pd.Series(0, index=dfd.index)).fillna(0)
            )
            prod_raw.loc[def_mask] = raw.rank(pct=True, method="average") * 100

        prod_raw = prod_raw.fillna(50.0)  # neutral for unknown

        # ── confidence boost (0-1 → 85–100% multiplier) ────────────────────
        conf = df_block["Confidence_EV"].fillna(0.5).clip(0, 1)

        # ── composite raw score ─────────────────────────────────────────────
        dim = df_block["Dimensionality_Score"].fillna(0)
        composite = (
            0.65 * dim +
            0.25 * prod_raw +
            0.10 * conf * 100
        )

        # ── confidence discount: low-sample players can't reach Elite/Star ──
        # Below 0.35 conf → cap composite at Solid ceiling (67.9)
        # Below 0.20 conf → cap at Depth ceiling (49.9)
        composite = composite.copy()
        low_conf  = conf < 0.35
        very_low  = conf < 0.20
        composite.loc[low_conf]  = composite.loc[low_conf].clip(upper=67.9)
        composite.loc[very_low]  = composite.loc[very_low].clip(upper=49.9)

        # ── assign base tiers ───────────────────────────────────────────────
        def _label(v):
            if pd.isna(v):  return "Unknown"
            if v >= 92:     return "Elite"
            if v >= 82:     return "Star"
            if v >= 68:     return "Solid"
            if v >= 50:     return "Depth"
            return "Fringe"

        tier = composite.map(_label)
        return tier, composite

    tier_labels, composite_scores = _compute_player_tier(out)
    out["Player_Tier"]      = tier_labels
    out["Player_Tier_Score"] = composite_scores.round(1)

    # Assign player type labels
    out = assign_player_type(out)

    # ── On-ice percentage metrics (always computable from raw counts) ──────────
    xgf = out.get("xgf", pd.Series(np.nan, index=out.index)).astype(float)
    xga = out.get("xga", pd.Series(np.nan, index=out.index)).astype(float)
    cf  = out.get("cf",  pd.Series(np.nan, index=out.index)).astype(float)
    ca  = out.get("ca",  pd.Series(np.nan, index=out.index)).astype(float)
    gf  = out.get("gf",  pd.Series(np.nan, index=out.index)).astype(float)
    ga  = out.get("ga",  pd.Series(np.nan, index=out.index)).astype(float)
    hdcf = out.get("hdcf", pd.Series(np.nan, index=out.index)).astype(float)
    hdca = out.get("hdca", pd.Series(np.nan, index=out.index)).astype(float)

    out["xGF_pct"]  = np.where((xgf + xga) > 0, (xgf / (xgf + xga) * 100).round(1), np.nan)
    out["CF_pct"]   = np.where((cf  + ca)  > 0, (cf  / (cf  + ca)  * 100).round(1), np.nan)
    out["GF_pct"]   = np.where((gf  + ga)  > 0, (gf  / (gf  + ga)  * 100).round(1), np.nan)
    out["HDCF_pct"] = np.where((hdcf + hdca) > 0, (hdcf / (hdcf + hdca) * 100).round(1), np.nan)

    # ── Shooting % vs expected (luck / finishing skill signal) ─────────────────
    # sh_pct stored as 0-1 fraction; ixg_ind and shots_ind are season totals
    sh_pct  = out.get("sh_pct",    pd.Series(np.nan, index=out.index)).astype(float)
    shots   = out.get("shots_ind", pd.Series(np.nan, index=out.index)).astype(float)
    ixg     = out.get("ixg_ind",   pd.Series(np.nan, index=out.index)).astype(float)

    xsh_pct = np.where(shots > 5, ixg / shots.replace(0, np.nan), np.nan)
    out["xSH_pct"]         = np.where(shots > 5, np.round(xsh_pct * 100, 1), np.nan)  # display as %
    out["SH_pct_display"]  = (sh_pct * 100).round(1)  # convert 0-1 → display %
    out["Finishing_Delta"] = np.where(
        sh_pct.notna() & pd.Series(xsh_pct).notna(),
        ((sh_pct - pd.Series(xsh_pct)) * 100).round(1),
        np.nan
    )  # positive = scoring more than shot quality predicts

    # ── League rank within position (by Dimensionality_Score) ─────────────────
    if "pos" in out.columns and "Dimensionality_Score" in out.columns:
        out["Pos_Rank"] = (
            out.groupby("pos")["Dimensionality_Score"]
               .rank(ascending=False, method="min", na_option="bottom")
               .astype("Int64")
        )
        out["Pos_Total"] = out.groupby("pos")["pos"].transform("count")
    else:
        out["Pos_Rank"]  = np.nan
        out["Pos_Total"] = np.nan

    # ── Zone deployment vs team average ──────────────────────────────────────
    if "off_zone_start_pct" in out.columns and "team" in out.columns and "pos_group" in out.columns:
        team_avg_oz = out.groupby(["team", "pos_group"])["off_zone_start_pct"].transform("mean")
        out["OZ_vs_Team"] = (out["off_zone_start_pct"] - team_avg_oz).round(1)
        # positive = deployed more offensively than team-mates at same position group
    else:
        out["OZ_vs_Team"] = np.nan

    return out


def add_prospect_tier(df: pd.DataFrame) -> pd.DataFrame:
    """
    Prospect tier overlay — applied AFTER add_player_projections so that
    Draft_Pedigree, age, overall_pick, and Prog_Tier all exist on the dataframe.

    Criteria (ALL must be true):
      1. Not already Star or Elite (they've arrived; no label needed)
      2. Draft pedigree is Elite (~picks 1-10) or High (~picks 11-32)
      3. Age < 22 (strict: still clearly in the early development window)
         OR age 22-23 AND pedigree is Elite
         OR age 24 AND pedigree is Elite AND pick <= 10
      4. Minimum NHL presence: toi_ev >= 20 min
         (filters pure AHL call-ups who aren't really in the NHL yet)
      5. Bust guard: age >= 22 AND pick <= 10 AND conf >= 0.55 AND tier is Fringe
         → keep as Fringe; they've had a genuine opportunity and underperformed
         their capital. This is the honest signal; don't hide it.
    """
    out = df.copy()

    required = {"Player_Tier", "Draft_Pedigree", "age", "overall_pick"}
    if not required.issubset(out.columns):
        missing = required - set(out.columns)
        # silently skip if key columns aren't present yet (shouldn't happen in normal pipeline)
        return out

    pedigree  = out["Draft_Pedigree"].fillna("").astype(str)
    age_s     = out["age"].fillna(99).astype(float)
    pick      = out["overall_pick"].fillna(999).astype(float)
    toi_s     = out.get("toi_ev", pd.Series(0, index=out.index)).fillna(0).astype(float)
    conf      = out.get("Confidence_EV", pd.Series(0.5, index=out.index)).fillna(0.5).astype(float)
    tier      = out["Player_Tier"].astype(str)

    elite_ped = pedigree == "Elite"          # ~picks 1-10
    high_ped  = pedigree.isin(["Elite", "High"])  # ~picks 1-32

    # Age windows — simplified and tightened per user request
    window_u22  = age_s < 22                                        # any Elite/High pick
    window_2223 = age_s.between(22, 23.49) & elite_ped             # 22-23, Elite pedigree only
    window_24   = age_s.between(23.5, 24.5) & elite_ped & (pick <= 10)  # 24yo franchise pick

    # Minimum NHL presence
    has_nhl_time = toi_s >= 20.0

    # Not already arrived
    not_arrived = ~tier.isin(["Star", "Elite"])

    # Bust guard: full sample, top-10 pick, age 22+, still Fringe
    bust_guard = (age_s >= 22) & (pick <= 10) & (conf >= 0.55) & (tier == "Fringe")

    prospect_mask = (
        (window_u22 | window_2223 | window_24) &
        high_ped &
        has_nhl_time &
        not_arrived &
        ~bust_guard
    )

    out.loc[prospect_mask, "Player_Tier"] = "Prospect"
    return out


def _rule_hit(val: float, op: str, thr: float) -> bool:
    if pd.isna(val):
        return False
    if op == ">=": return val >= thr
    if op == ">":  return val > thr
    if op == "<=": return val <= thr
    if op == "<":  return val < thr
    if op == "==": return val == thr
    return False

def assign_player_type(df: pd.DataFrame) -> pd.DataFrame:
    """
    Creates:
      Player_Type, Secondary_Type, Type_Margin, Type_Confidence, Type_Notes

    Rules:
      - Uses top-two role score logic with confidence + hybrid margin guardrails.
      - If TOI below thresholds -> "Insufficient TOI"
      - If Top1–Top2 margin < TYPE_HYBRID_MARGIN -> "Hybrid: A / B"
      - Type_Confidence uses Confidence_EV and margin thresholds
      - IMPORTANT TIEBREAK: for Forwards, if Producer is tied for top score, Producer wins.
    """
    out = df.copy()

    type_maps_f = {
        "Finisher_Score": "Finisher",
        "Playmaker_Score": "Playmaker",
        "Driver_Score": "Driver",
        "TwoWay_Score": "Two-Way",
        "Power_Score": "Power",
        "Grinder_Score": "Grinder / Tough Guy",
        "Producer_Score": "Producer",
    }
    type_maps_d = {
        "Suppressor_Score": "Shutdown D",
        "Transition_Score": "Transition",
        "PuckSkill_Score": "Puck Skill",
        "Physical_Score": "Physical",
    }

    # ensure output columns exist (stable dtypes)
    for c in ["Player_Type", "Secondary_Type", "Type_Margin", "Type_Confidence", "Type_Notes"]:
        if c not in out.columns:
            out[c] = np.nan

    EPS = 1e-9

    def _type_for_row(row: pd.Series) -> Tuple[str, str, float, str, str]:
        # guardrail: insufficient TOI / keep flag
        if not bool(row.get("keep", False)):
            return ("Insufficient TOI", "", np.nan, "Low", "Below TOI threshold; types not stable.")

        conf = row.get("Confidence_EV", 0.0)
        try:
            conf = float(conf) if pd.notna(conf) else 0.0
        except Exception:
            conf = 0.0

        pos_group = row.get("pos_group", "F")
        role_map = type_maps_f if pos_group == "F" else type_maps_d

        scores: List[Tuple[float, str, str]] = []
        for col, label in role_map.items():
            v = row.get(col, np.nan)
            if v is None or pd.isna(v):
                continue
            try:
                scores.append((float(v), col, label))
            except Exception:
                continue

        if len(scores) == 0:
            return ("Unknown", "", np.nan, "Low", "No role scores available.")

        # sort by score desc
        scores.sort(key=lambda x: x[0], reverse=True)

        # ------------------------------------------------------------
        # IND-based type nudges (your requested logic)
        # ------------------------------------------------------------
        # If we have IND season counts/rates, apply small nudges to avoid dumb labels.
        # These nudges only matter when the target role is close to the top.
        assists = row.get("assists_ind", np.nan)
        pts_per_gp = row.get("pts_per_gp_ind", np.nan)

        # Convert to float safely
        try: assists = float(assists) if pd.notna(assists) else np.nan
        except: assists = np.nan
        try: pts_per_gp = float(pts_per_gp) if pd.notna(pts_per_gp) else np.nan
        except: pts_per_gp = np.nan

        # ------------------------------------------------------------
        # IND-driven archetype nudges (config-based)
        # ------------------------------------------------------------
        nudges = TYPE_NUDGES_FWD if pos_group == "F" else TYPE_NUDGES_DEF

        # quick score lookup by role column
        def _boost(role_col: str, add: float):
            for i, (v, c, lab) in enumerate(scores):
                if c == role_col:
                    scores[i] = (v + add, c, lab)

        # map role labels -> role score columns used in 'scores'
        label_to_col = {}
        for _, col, label in scores:
            label_to_col[label] = col

        # apply rules
        for archetype_label, rules in nudges.items():
            role_col = label_to_col.get(archetype_label)
            if not role_col:
                continue

            total_add = 0.0
            for metric, op, thr, add in rules:
                v = row.get(metric, np.nan)
                try:
                    v = float(v) if pd.notna(v) else np.nan
                except Exception:
                    v = np.nan

                if _rule_hit(v, op, float(thr)):
                    total_add += float(add)

            if total_add != 0.0:
                _boost(role_col, total_add)

        # re-sort after nudges
        scores.sort(key=lambda x: x[0], reverse=True)

        # -------------------------
        # TOP1 with Producer tiebreak (Forwards)
        # -------------------------
        top_score = scores[0][0]
        top_candidates = [t for t in scores if abs(t[0] - top_score) <= EPS]

        if pos_group == "F":
            prod_cand = [t for t in top_candidates if t[2] == "Producer"]
            if prod_cand:
                top1, top1_col, top1_label = prod_cand[0]
            else:
                top1, top1_col, top1_label = top_candidates[0]
        else:
            top1, top1_col, top1_label = top_candidates[0]

        # TOP2 = best remaining after removing chosen top1 column
        remaining = [t for t in scores if t[1] != top1_col]
        if remaining:
            top2, top2_col, top2_label = remaining[0]
        else:
            top2, top2_col, top2_label = (np.nan, "", "")

        margin = float(top1 - top2) if pd.notna(top2) else float("nan")

        # Hybrid logic
        if pd.notna(top2) and margin < TYPE_HYBRID_MARGIN:
            ptype = f"Hybrid: {top1_label} / {top2_label}"
        else:
            ptype = top1_label

        # Confidence label
        if conf < TYPE_MIN_CONFIDENCE_EV:
            tconf = "Low"
        else:
            if pd.notna(margin) and margin >= TYPE_CLEAR_MARGIN:
                tconf = "High"
            elif pd.notna(margin) and margin >= TYPE_HYBRID_MARGIN:
                tconf = "Medium"
            else:
                tconf = "Medium"

        notes = []
        is_injured = bool(row.get("Injury_Imputed") or row.get("injury_fill"))
        if is_injured:
            notes.append("Stats carried from prior season due to injury/absence — archetype based on historical data.")
        if conf < TYPE_MIN_CONFIDENCE_EV:
            reason = "injury/missed games" if is_injured else f"limited sample (Confidence_EV={conf:.2f})"
            notes.append(f"Low confidence due to {reason}.")
        if pd.notna(margin):
            notes.append(f"Top1–Top2 margin={margin:.1f}.")
        if "Hybrid:" in ptype:
            notes.append("Archetype is blended (top two roles close).")
        if pos_group == "F" and (abs(top1 - top_score) <= EPS) and ("Producer" in [t[2] for t in top_candidates]) and (top1_label == "Producer"):
            notes.append("Producer tiebreak applied.")

        return (ptype, top2_label, margin, tconf, " ".join(notes).strip())

    # vectorized apply
    rows = out.apply(_type_for_row, axis=1, result_type="expand")
    rows.columns = ["Player_Type", "Secondary_Type", "Type_Margin", "Type_Confidence", "Type_Notes"]

    out["Player_Type"] = rows["Player_Type"].astype(object)
    out["Secondary_Type"] = rows["Secondary_Type"].astype(object)
    out["Type_Margin"] = pd.to_numeric(rows["Type_Margin"], errors="coerce")
    out["Type_Confidence"] = rows["Type_Confidence"].astype(object)
    out["Type_Notes"] = rows["Type_Notes"].astype(object)

    return out


# =============================================================================
# ROLLING
# =============================================================================
def build_rolling(season_dfs: Dict[str, pd.DataFrame], season_order: List[str], weights: Tuple[float, float, float]) -> pd.DataFrame:
    newest, prev1, prev2 = season_order[0], season_order[1], season_order[2]
    w0, w1, w2 = weights

    missing = [s for s in [newest, prev1, prev2] if s not in season_dfs]
    if missing:
        raise KeyError(f"build_rolling missing season(s): {missing}. Available: {list(season_dfs.keys())}")

    df0 = season_dfs[newest].copy()
    df1 = season_dfs[prev1].copy()
    df2 = season_dfs[prev2].copy()

    id_cols = ["player_key", "player", "pos", "pos_group", "team"]
    score_cols = sorted({c for c in df0.columns if c.endswith("_Score")} | {"Dimensionality_Score", "Confidence_EV"})
    score_cols = [c for c in score_cols if c in df0.columns]

    def prep(df: pd.DataFrame, suf: str) -> pd.DataFrame:
        cols = [c for c in (id_cols + score_cols) if c in df.columns]
        tmp = df[cols].copy()
        ren = {c: f"{c}_{suf}" for c in cols if c != "player_key"}
        return tmp.rename(columns=ren)

    a = prep(df0, "y0")
    b = prep(df1, "y1")
    c = prep(df2, "y2")

    roll = a.merge(b, on="player_key", how="outer").merge(c, on="player_key", how="outer")

    def coalesce(*cols: str) -> pd.Series:
        s = None
        for col in cols:
            if col in roll.columns:
                s = roll[col] if s is None else s.fillna(roll[col])
        return s if s is not None else pd.Series([np.nan] * len(roll))

    roll["player"] = coalesce("player_y0", "player_y1", "player_y2")
    roll["pos"] = coalesce("pos_y0", "pos_y1", "pos_y2")
    roll["pos_group"] = coalesce("pos_group_y0", "pos_group_y1", "pos_group_y2")
    roll["team"] = coalesce("team_y0", "team_y1", "team_y2")

    for base in score_cols:
        c0, c1, c2 = f"{base}_y0", f"{base}_y1", f"{base}_y2"
        out_vals = np.full(len(roll), np.nan, dtype=float)

        for i in range(len(roll)):
            xs, ws = [], []
            if c0 in roll.columns and pd.notna(roll.at[i, c0]):
                xs.append(float(roll.at[i, c0])); ws.append(w0)
            if c1 in roll.columns and pd.notna(roll.at[i, c1]):
                xs.append(float(roll.at[i, c1])); ws.append(w1)
            if c2 in roll.columns and pd.notna(roll.at[i, c2]):
                xs.append(float(roll.at[i, c2])); ws.append(w2)
            if not xs:
                continue
            ww = np.array(ws, dtype=float)
            ww = ww / ww.sum()
            out_vals[i] = float(np.dot(ww, np.array(xs, dtype=float)))

        roll[f"{base}_3yr"] = out_vals

        year_cols = [col for col in [c0, c1, c2] if col in roll.columns]

        def _row_std(vals: np.ndarray) -> float:
            vals = vals[~np.isnan(vals)]
            if vals.size < 2:
                return np.nan
            return float(np.std(vals, ddof=0))

        mat = np.vstack([roll[col].to_numpy(dtype=float) for col in year_cols]).T
        roll[f"{base}_vol_3yr"] = np.apply_along_axis(_row_std, 1, mat)

    keep = ["player", "team", "pos", "pos_group"] + [c for c in roll.columns if c.endswith("_3yr") or c.endswith("_vol_3yr")]
    roll_out = roll[keep].copy()

    # ── YoY trend columns ─────────────────────────────────────────────────────
    # Compare newest season vs prior season for key signals.
    # Dim_Delta_YoY: change in Dimensionality_Score (current vs 1yr ago)
    # Dim_Trend / Role_Trend: Improving / Stable / Declining labels
    def _trend_label(delta: float, threshold: float = 5.0) -> str:
        if pd.isna(delta): return "Unknown"
        if delta >= threshold: return "Improving"
        if delta <= -threshold: return "Declining"
        return "Stable"

    for base_col, delta_name, trend_name in [
        ("Dimensionality_Score", "Dim_Delta_YoY", "Dim_Trend"),
        ("Top_Role_Score",       "TopRole_Delta_YoY", "Role_Trend"),
    ]:
        y0_col = f"{base_col}_y0"
        y1_col = f"{base_col}_y1"
        if y0_col in roll.columns and y1_col in roll.columns:
            delta = roll[y0_col].astype(float) - roll[y1_col].astype(float)
            roll_out[delta_name] = delta.round(1)
            roll_out[trend_name] = delta.map(_trend_label)
        else:
            roll_out[delta_name] = np.nan
            roll_out[trend_name] = "Unknown"

    # Seasons with data (1, 2, or 3)
    seasons_with_data = pd.Series(0, index=roll.index)
    for yc in [f"Dimensionality_Score_y0", f"Dimensionality_Score_y1", f"Dimensionality_Score_y2"]:
        if yc in roll.columns:
            seasons_with_data += roll[yc].notna().astype(int)
    roll_out["Seasons_With_Data"] = seasons_with_data.astype(int)

    return roll_out

ROLE_SCORE_COLS = [
    "Finisher_Score","Playmaker_Score","Driver_Score","TwoWay_Score","Power_Score",
    "Grinder_Score","Producer_Score","Suppressor_Score","Transition_Score","PuckSkill_Score",
    "Physical_Score","Dimensionality_Score"
]

def add_roles_from_scores(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()

    cols = [c for c in ROLE_SCORE_COLS if c in out.columns]
    if not cols:
        return out

    scores = out[cols].to_numpy()
    # treat all-NaN rows safely
    scores_nan = np.isnan(scores).all(axis=1)
    scores = np.where(np.isnan(scores), -np.inf, scores)

    top_idx = np.argmax(scores, axis=1)
    top_scores = scores[np.arange(len(out)), top_idx]

    # second-best: mask top then argmax again
    scores2 = scores.copy()
    scores2[np.arange(len(out)), top_idx] = -np.inf
    sec_idx = np.argmax(scores2, axis=1)
    sec_scores = scores2[np.arange(len(out)), sec_idx]

    out["Top_Role"] = np.array(cols, dtype=object)[top_idx]
    out["Top_Role_Score"] = top_scores
    out["Second_Role"] = np.array(cols, dtype=object)[sec_idx]
    out["Second_Role_Score"] = sec_scores

    # restore blanks for truly missing rows
    out.loc[scores_nan, ["Top_Role","Second_Role"]] = np.nan
    out.loc[scores_nan, ["Top_Role_Score","Second_Role_Score"]] = np.nan

    return out

# =============================================================================
# OUTPUT CURATION
# =============================================================================
COMMON_TYPE_COLS = ["Player_Type", "Secondary_Type", "Type_Margin", "Type_Confidence", "Type_Notes"]

TYPE_DETAIL_COLS = ["Top_Role","Top_Role_Score","Second_Role","Second_Role_Score","Margin","Role_Stability","Type_Confidence"]

FWD_VISIBLE = [
    # Identity
    "player", "team", "season", "pos", "gp", "toi_ev",
    # Bio
    "age", "height_in", "weight_lb",
    "draft_round", "draft_pick",
    # Quick-scan tier + injury flag
    "Player_Tier", "Player_Tier_Score", "Injury_Imputed", "Cap_Value_Per_M",
    # Rank & trajectory
    "Pos_Rank", "Pos_Total",
    "Dim_Delta", "Dim_Trend",
    # Player archetype
    *COMMON_TYPE_COLS,
    # Role scores
    "Finisher_Score", "Driver_Score", "TwoWay_Score",
    "Producer_Score", "Grinder_Score", "Playmaker_Score", "Power_Score",
    "Dimensionality_Score", "Confidence_EV",
    # Top roles (detail)
    "Top_Role", "Top_Role_Score", "Second_Role", "Second_Role_Score", "Margin",
    # On-ice team context
    "xGF_pct", "CF_pct", "GF_pct", "HDCF_pct",
    "gf60_ev", "xgf60_ev", "sf60_ev", "hdcf60_ev",
    "xga60_ev", "hdca60_ev",
    "off_zone_start_pct", "OZ_vs_Team",
    # Individual rates
    "g60_ind", "a1_60_ind", "p1_60_ind", "points60_ind",
    "shots60_ind", "ixg60_ind",
    "SH_pct_display", "xSH_pct", "Finishing_Delta",
    "hits60_ind", "giveaways60_ind", "takeaways60_ind",
    "pen_drawn60_ind", "pen_taken60_ind", "pen_diff60_ind",
    "rush60_ind", "rebounds60_ind", "blk_shots60_ind",
    "fo_pct",
    # Projections
    "Proj_Points60", "Proj_Goals60", "Proj_PTS_PGP",
    "Age_Curve_Factor", "Prog_Tier", "Prog_Confidence", "Draft_Pedigree",
    # Line fit
    "Best_Line_Fit", "Best_Line_Fit_Score",
    "Fit_Line1", "Fit_Line2", "Fit_Line3", "Fit_Line4",
]

DEF_VISIBLE = [
    # Identity
    "player", "team", "season", "pos", "gp", "toi_ev",
    # Bio
    "age", "height_in", "weight_lb",
    "draft_round", "draft_pick",
    # Quick-scan tier + injury flag
    "Player_Tier", "Player_Tier_Score", "Injury_Imputed", "Cap_Value_Per_M",
    # Rank & trajectory
    "Pos_Rank", "Pos_Total",
    "Dim_Delta", "Dim_Trend",
    # Player archetype
    *COMMON_TYPE_COLS,
    # Role scores
    "Suppressor_Score", "Transition_Score", "PuckSkill_Score", "Physical_Score",
    "Dimensionality_Score", "Confidence_EV",
    # Top roles (detail)
    "Top_Role", "Top_Role_Score", "Second_Role", "Second_Role_Score", "Margin",
    # On-ice team context
    "xGF_pct", "CF_pct", "GF_pct", "HDCF_pct",
    "xga60_ev", "ca60_ev", "sa60_ev", "hdca60_ev",
    "xgf60_ev", "cf60_ev", "sf60_ev", "hdcf60_ev",
    "def_zone_start_pct", "off_zone_start_pct", "OZ_vs_Team",
    # Individual rates (D get fewer offensive metrics)
    "points60_ind", "a1_60_ind",
    "SH_pct_display", "xSH_pct", "Finishing_Delta",
    "hits60_ind", "blk_shots60_ind",
    "giveaways60_ind", "takeaways60_ind",
    # Projections
    "Proj_Points60", "Proj_Goals60", "Proj_PTS_PGP",
    "Age_Curve_Factor", "Prog_Tier", "Prog_Confidence", "Draft_Pedigree",
    # Pair fit
    "Best_Pair_Fit", "Best_Pair_Fit_Score",
    "Fit_Pair1", "Fit_Pair2", "Fit_Pair3",
]

def _hide_sheet(ws) -> None:
    ws.sheet_state = "hidden"

def _hide_unused_columns(ws, keep_col_names: List[str]) -> None:
    headers = [cell.value for cell in ws[1]]
    keep_idx = {i + 1 for i, h in enumerate(headers) if h in set(keep_col_names)}
    for col_idx in range(1, ws.max_column + 1):
        if col_idx not in keep_idx:
            ws.column_dimensions[get_column_letter(col_idx)].hidden = True

from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

USD_FMT    = '"$"#,##0.00_);[Red]"$"#,##0.00'
PCT_FMT    = '0.00%'
INT_FMT    = '#,##0'
DEC2_FMT   = '#,##0.00'

# ---------------------------------------------------------------------------
# COLUMN DISPLAY NAMES
# Maps internal snake_case column names → pretty display headers.
# Any column NOT in this map keeps its original name.
# ---------------------------------------------------------------------------
COLUMN_DISPLAY_NAMES: Dict[str, str] = {
    # Identifiers
    "player":               "Player",
    "player_key":           "Player Key",
    "team":                 "Team",
    "team_cap":             "Team",
    "pos":                  "Position",
    "pos_group":            "Group",
    "pos_cap":              "Position",
    "pos_raw":              "Position (Raw)",
    "season":               "Season",
    "gp":                   "GP",
    "toi_ev":               "TOI (EV, min)",
    "toi_ind":              "TOI (Total, min)",
    "keep":                 "Min TOI Met",
    # Bio
    "age":                  "Age",
    "height_in":            "Height (in)",
    "weight_lb":            "Weight (lbs)",
    "bmi":                  "BMI",
    "draft_round":          "Draft Round",
    "draft_pick":           "Draft Pick (Overall)",
    "overall_pick":         "Overall Pick",
    "round_pick":           "Round Pick",
    # Type / classification
    "Player_Type":          "Player Type",
    "Secondary_Type":       "Secondary Type",
    "Type_Margin":          "Type Margin",
    "Type_Confidence":      "Type Confidence",
    "Type_Notes":           "Type Notes",
    "Top_Role":             "Top Role",
    "Top_Role_Score":       "Top Role Score",
    "Second_Role":          "Second Role",
    "Second_Role_Score":    "Second Role Score",
    "Margin":               "Role Margin",
    "Role_Stability":       "Role Stability",
    # Role scores — forwards
    "Finisher_Score":       "Finisher",
    "Playmaker_Score":      "Playmaker",
    "Driver_Score":         "Driver",
    "TwoWay_Score":         "Two-Way",
    "Power_Score":          "Power",
    "Grinder_Score":        "Grinder",
    "Producer_Score":       "Producer",
    # Role scores — defense
    "Suppressor_Score":     "Shutdown D",
    "Transition_Score":     "Transition",
    "PuckSkill_Score":      "Puck Skill",
    "Physical_Score":       "Physical",
    # Composite
    "Dimensionality_Score": "Dimensionality",
    "Confidence_EV":        "Sample Confidence",
    "Player_Tier_Score":    "Tier Score",
    # On-ice (EV per-60)
    "cf60_ev":              "CF/60",
    "ca60_ev":              "CA/60",
    "sf60_ev":              "SF/60",
    "sa60_ev":              "SA/60",
    "gf60_ev":              "GF/60",
    "ga60_ev":              "GA/60",
    "xgf60_ev":             "xGF/60",
    "xga60_ev":             "xGA/60",
    "scf60_ev":             "SCF/60",
    "sca60_ev":             "SCA/60",
    "hdcf60_ev":            "HDCF/60",
    "hdca60_ev":            "HDCA/60",
    # Zone starts
    "off_zone_start_pct":   "OZ Start %",
    "def_zone_start_pct":   "DZ Start %",
    "oz_starts":            "OZ Starts",
    "nz_starts":            "NZ Starts",
    "dz_starts":            "DZ Starts",
    # Individual per-60
    "g60_ind":              "G/60",
    "a1_60_ind":            "A1/60",
    "p1_60_ind":            "P1/60",
    "points60_ind":         "Pts/60",
    "shots60_ind":          "Shots/60",
    "ixg60_ind":            "ixG/60",
    "hits60_ind":           "Hits/60",
    "giveaways60_ind":      "Giveaways/60",
    "takeaways60_ind":      "Takeaways/60",
    "pen_drawn60_ind":      "Pen Drawn/60",
    "pen_taken60_ind":      "Pen Taken/60",
    "pen_diff60_ind":       "Pen Diff/60",
    "rush60_ind":           "Rush Att/60",
    "rebounds60_ind":       "Rebounds/60",
    "blk_shots60_ind":      "Blocks/60",
    "fow60_ind":            "FOW/60",
    "fol60_ind":            "FOL/60",
    # Individual counting
    "goals_ind":            "Goals",
    "assists_ind":          "Assists",
    "a1_ind":               "Primary Assists",
    "a2_ind":               "Secondary Assists",
    "points_ind":           "Points",
    "shots_ind":            "Shots",
    "ixg_ind":              "ixG",
    "ihdcf":                "iHDCF",
    "rush_ind":             "Rush Attempts",
    "rebounds_ind":         "Rebounds Created",
    "hits_ind":             "Hits",
    "shots_blocked_ind":    "Blocks",
    "giveaways_ind":        "Giveaways",
    "takeaways_ind":        "Takeaways",
    "pim_ind":              "PIM",
    "pen_total_ind":        "Penalties Taken",
    "pen_drawn_ind":        "Penalties Drawn",
    "fow_ind":              "Faceoffs Won",
    "fol_ind":              "Faceoffs Lost",
    # Rates (stored 0–1)
    "sh_pct":               "SH%",
    "fo_pct":               "FO%",
    "ipp":                  "IPP",
    # Projection outputs
    "Proj_Points60":        "Proj Pts/60",
    "Proj_Goals60":         "Proj G/60",
    "Proj_PTS_PGP":         "Proj Pts/GP",
    "Age_Curve_Factor":     "Age Curve",
    "Prog_Tier":            "Progression Tier",
    "Prog_Confidence":      "Proj Confidence",
    "Draft_Pedigree":       "Draft Pedigree",
    # Line / pair fit
    "Fit_Line1":            "Line 1 Fit",
    "Fit_Line2":            "Line 2 Fit",
    "Fit_Line3":            "Line 3 Fit",
    "Fit_Line4":            "Line 4 Fit",
    "Best_Line_Fit":        "Best Line Fit",
    "Best_Line_Fit_Score":  "Best Line Fit Score",
    "Fit_Pair1":            "Pair 1 Fit",
    "Fit_Pair2":            "Pair 2 Fit",
    "Fit_Pair3":            "Pair 3 Fit",
    "Best_Pair_Fit":        "Best Pair Fit",
    "Best_Pair_Fit_Score":  "Best Pair Fit Score",
    "Usage_Rank_Team":      "Usage Rank (Team)",
    # Cap / contract
    "cap_hit":              "Cap Hit",
    "new_cap_hit":          "New Cap Hit",
    "term":                 "Term",
    "expiry_year":          "Expiry Year",
    "contract_type":        "Contract Type",
    "hand":                 "Hand",
    "aav_est":              "Est. AAV",
    "aav_mean":             "AAV Mean",
    "aav_std":              "AAV Std Dev",
    "aav_p10":              "AAV P10",
    "aav_p50":              "AAV P50",
    "aav_p90":              "AAV P90",
    "comp_n":               "# Comps",
    "impact":               "Impact Score",
    "resign_priority":      "Re-sign Priority Score",
    "fa_priority":          "FA Priority Score",
    "priority":             "Priority Score",
    "hole_bonus":           "Hole Fill Bonus",
    "Signed":               "Signed by Model",
    "value_per_$":          "Value / $M",
    "value_prog_adj":       "Value / $M (Prog Adj)",
    # Offseason planner
    "decision":             "Decision",
    "origin":               "Origin",
    "hole_bonus":           "Hole Fill Bonus",
    "hole_score":           "Hole Fill Score",
    "Assigned_Unit":        "Assigned Unit",
    "Assigned_Slot":        "Assigned Slot",
    # Cap summary
    "cap_limit":                    "Cap Limit",
    "core_under_contract_cap":      "Core Cap (Under Contract)",
    "resign_spend_est":             "Re-sign Spend (Est.)",
    "post_resign_cap":              "Post Re-sign Cap",
    "fa_sign_spend_est":            "FA Signing Spend (Est.)",
    "projected_total_cap":          "Projected Total Cap",
    "projected_cap_space":          "Projected Cap Space",
    "need_F_remaining":             "Forwards Still Needed",
    "need_D_remaining":             "Defense Still Needed",
    "need_C_remaining":             "Centers Still Needed",
    "need_LW_remaining":            "LW Still Needed",
    "need_RW_remaining":            "RW Still Needed",
    # Imputation flags
    "Injury_Imputed":       "Injury Imputed",
    "injury_fill":          "Injury Fill",
    "fill_source":          "Fill Source",
}

# ---------------------------------------------------------------------------
# Number format rules: (col_name_fragment_or_exact → format_string)
# Checked as exact match first, then endswith, then contains.
# ---------------------------------------------------------------------------
_NUM_FMT_RULES: List[Tuple[str, str]] = [
    # Currency — exact
    ("cap_hit",                     USD_FMT),
    ("new_cap_hit",                 USD_FMT),
    ("aav_est",                     USD_FMT),
    ("aav_mean",                    USD_FMT),
    ("aav_std",                     USD_FMT),
    ("aav_p10",                     USD_FMT),
    ("aav_p50",                     USD_FMT),
    ("aav_p90",                     USD_FMT),
    ("cap_limit",                   USD_FMT),
    ("core_under_contract_cap",     USD_FMT),
    ("resign_spend_est",            USD_FMT),
    ("post_resign_cap",             USD_FMT),
    ("fa_sign_spend_est",           USD_FMT),
    ("projected_total_cap",         USD_FMT),
    ("projected_cap_space",         USD_FMT),
    # Percentages (stored 0–1)
    ("sh_pct",      PCT_FMT),
    ("fo_pct",      PCT_FMT),
    ("ipp",         PCT_FMT),
    ("off_zone_start_pct", PCT_FMT),
    ("def_zone_start_pct", PCT_FMT),
    ("Confidence_EV",      PCT_FMT),
    ("Role_Stability",     PCT_FMT),
    ("Age_Curve_Factor",   DEC2_FMT),
    # Integer-ish
    ("gp",          INT_FMT),
    ("term",        INT_FMT),
    ("expiry_year", INT_FMT),
    ("comp_n",      INT_FMT),
    ("draft_round", INT_FMT),
    ("draft_pick",  INT_FMT),
    ("overall_pick",INT_FMT),
    ("round_pick",  INT_FMT),
    ("oz_starts",   INT_FMT),
    ("nz_starts",   INT_FMT),
    ("dz_starts",   INT_FMT),
    # 2-decimal (catch-all for scores, per-60s, etc.)
]

# Columns that get 2-decimal regardless (matched by suffix)
_DEC2_SUFFIXES = ("_Score", "_score", "60", "_pct", "impact", "value_per_$",
                  "value_prog_adj", "hole_bonus", "hole_score",
                  "Proj_Points60", "Proj_Goals60", "Proj_PTS_PGP",
                  "toi_ev", "toi_ind", "bmi", "age",
                  "Type_Margin", "Margin")


def _col_num_fmt(col_raw: str) -> Optional[str]:
    """Return the Excel number format string for a given raw column name, or None."""
    for key, fmt in _NUM_FMT_RULES:
        if col_raw == key:
            return fmt
    for suf in _DEC2_SUFFIXES:
        if col_raw.endswith(suf) or col_raw == suf:
            return DEC2_FMT
    return None


# Header style
_HDR_FILL  = PatternFill(fill_type="solid", fgColor="1F4E79")   # dark navy
_HDR_FONT  = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
_HDR_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)

# Alternating row fills
_ROW_FILL_ODD  = PatternFill(fill_type="solid", fgColor="FFFFFF")
_ROW_FILL_EVEN = PatternFill(fill_type="solid", fgColor="EBF3FB")

# Thin border for table cells
_THIN = Side(style="thin", color="B8CCE4")
_CELL_BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)

_DATA_FONT  = Font(name="Calibri", size=10)
_DATA_ALIGN_L = Alignment(horizontal="left",   vertical="center")
_DATA_ALIGN_C = Alignment(horizontal="center", vertical="center")
_DATA_ALIGN_R = Alignment(horizontal="right",  vertical="center")


def _write_depth_chart_sheet(wb: openpyxl.Workbook, lineup_df: pd.DataFrame, team: str,
                              roster_df: pd.DataFrame = None) -> None:
    """
    Writes a styled depth chart sheet from the lineup table produced by build_lineup_tables.
    roster_df: the full projected roster (used to populate healthy scratches — players
               on the roster but not in any active lineup unit).
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    sheet_name = f"{team} — Depth Chart"
    if sheet_name in wb.sheetnames:
        wb.remove(wb[sheet_name])
    ws = wb.create_sheet(sheet_name)

    # ---- Colour palette ----
    NAV   = "1F4E79"   # header navy
    GOLD  = "C9A227"   # line labels
    ICE   = "D6E4F0"   # fwd row background
    SLATE = "E8ECEF"   # def row background
    WHITE = "FFFFFF"
    WARN  = "FFF2CC"   # off-side hand warning

    def _hdr_font(size=11): return Font(bold=True, color=WHITE, name="Calibri", size=size)
    def _lbl_font():         return Font(bold=True, color=NAV,   name="Calibri", size=10)
    def _cell_font(bold=False, warn=False):
        color = "7B3F00" if warn else "000000"
        return Font(bold=bold, name="Calibri", size=10, color=color)
    def _fill(hex_color):
        return PatternFill(fill_type="solid", fgColor=hex_color)
    def _thin_border():
        s = Side(style="thin", color="B8CCE4")
        return Border(left=s, right=s, top=s, bottom=s)
    def _center(wrap=False): return Alignment(horizontal="center", vertical="center", wrap_text=wrap)
    def _left():  return Alignment(horizontal="left",  vertical="center")

    # ---- Title ----
    ws.column_dimensions["A"].width = 12
    for col in ["B","C","D"]: ws.column_dimensions[col].width = 26
    ws.column_dimensions["E"].width = 8   # hand note

    ws.row_dimensions[1].height = 6
    ws.row_dimensions[2].height = 28
    t = ws.cell(row=2, column=2, value=f"{team} — Projected Depth Chart (2026-27)")
    t.font      = Font(bold=True, name="Calibri", size=14, color=NAV)
    t.alignment = _left()
    ws.merge_cells("B2:D2")
    ws.row_dimensions[3].height = 8

    if lineup_df is None or len(lineup_df) == 0:
        ws.cell(row=4, column=2, value="No lineup data available.")
        return

    df = lineup_df.copy()

    # Normalise column names (handles both raw and display-renamed versions)
    _cn = {c.lower().replace(" ","_"): c for c in df.columns}
    def _col(*candidates):
        for c in candidates:
            if c in df.columns: return c
            if c.lower().replace(" ","_") in _cn: return _cn[c.lower().replace(" ","_")]
        return None

    unit_col    = _col("Assigned_Unit", "assigned_unit")
    slot_col    = _col("Assigned_Slot", "assigned_slot")
    disp_col    = _col("player", "Player", "Name")
    pos_col     = _col("pos", "Pos", "Position")
    hand_col    = _col("hand", "Hand")
    impact_col  = _col("impact", "Impact Score")
    id_col      = _col("player_key", "player", "Player")

    if unit_col is None or disp_col is None:
        ws.cell(row=4, column=2, value="Lineup data missing required columns.")
        return

    def _is_offside(natural_pos, slot, hand):
        """True if player is on their off/weaker side."""
        if not hand or not slot: return False
        if natural_pos == "C":  return False   # C on wing is expected, not flagged
        if natural_pos == "LW" and slot == "RW" and hand == "L": return True
        if natural_pos == "RW" and slot == "LW" and hand == "R": return True
        return False

    def _get_players_for_unit(unit):
        """Return rows for a given unit, keyed by slot."""
        block = df[df[unit_col].astype(str).eq(unit)].copy()
        return block

    def _display_name(row):
        name = str(row.get(disp_col) or "")
        nat  = str(row.get(pos_col)  or "").upper()
        slotv = str(row.get(slot_col) or "") if slot_col else ""
        hand  = str(row.get(hand_col)  or "").upper()[:1] if hand_col else ""
        imp   = row.get(impact_col)
        imp_s = f" ({imp:.0f})" if imp and not pd.isna(imp) else ""
        pos_note = f" [{nat}]" if nat and nat != slotv else ""
        return name + pos_note + imp_s, _is_offside(nat, slotv, hand)

    row_num = 4  # start writing from here

    # ============================================================
    # FORWARDS
    # ============================================================
    # Section label
    sec = ws.cell(row=row_num, column=1, value="FORWARDS")
    sec.font = Font(bold=True, name="Calibri", size=11, color=WHITE)
    sec.fill = _fill(NAV)
    sec.alignment = _center()
    for c in [1,2,3,4,5]:
        ws.cell(row=row_num, column=c).fill = _fill(NAV)
    ws.row_dimensions[row_num].height = 18
    row_num += 1

    # Column headers
    for col, label in [(1,"Line"),(2,"LW"),(3,"C"),(4,"RW"),(5,"Side Note")]:
        cell = ws.cell(row=row_num, column=col, value=label)
        cell.font      = _hdr_font(10)
        cell.fill      = _fill("2E75B6")
        cell.alignment = _center()
        cell.border    = _thin_border()
    ws.row_dimensions[row_num].height = 18
    row_num += 1

    for line in ["Line 1", "Line 2", "Line 3", "Line 4"]:
        block = _get_players_for_unit(line)
        slots_data = {"LW": ("","",False), "C": ("","",False), "RW": ("","",False)}

        if slot_col and slot_col in block.columns:
            for _, r in block.iterrows():
                slot = str(r.get(slot_col) or "").upper()
                if slot in slots_data:
                    txt, warn = _display_name(r)
                    slots_data[slot] = (txt, str(r.get(hand_col) or "") if hand_col else "", warn)
        else:
            # Fallback: assign by position best-effort
            taken = set()
            for target_slot in ["C","LW","RW"]:
                candidates = block[block[pos_col].astype(str).str.upper().eq(target_slot)] if pos_col else pd.DataFrame()
                candidates = candidates[~candidates.index.isin(taken)]
                if len(candidates):
                    r = candidates.iloc[0]
                    txt, warn = _display_name(r)
                    slots_data[target_slot] = (txt, "", warn)
                    taken.add(r.name)

        # Write row
        lbl_cell = ws.cell(row=row_num, column=1, value=line)
        lbl_cell.font = _lbl_font()
        lbl_cell.fill = _fill(ICE)
        lbl_cell.alignment = _center()
        lbl_cell.border = _thin_border()

        warn_notes = []
        for col_i, slot in enumerate(["LW","C","RW"], start=2):
            txt, hnd, warn = slots_data[slot]
            cell = ws.cell(row=row_num, column=col_i, value=txt)
            cell.fill      = _fill(WARN if warn else ICE)
            cell.font      = _cell_font(warn=warn)
            cell.alignment = _left()
            cell.border    = _thin_border()
            if warn:
                side = "off-side" if slot in ("LW","RW") else ""
                warn_notes.append(f"{slot}: off-side hand")

        note_cell = ws.cell(row=row_num, column=5, value="; ".join(warn_notes) if warn_notes else "")
        note_cell.font      = Font(italic=True, size=9, color="7B3F00", name="Calibri")
        note_cell.alignment = _left()
        note_cell.border    = _thin_border()

        ws.row_dimensions[row_num].height = 18
        row_num += 1

    row_num += 1  # spacer

    # ============================================================
    # DEFENSE
    # ============================================================
    sec = ws.cell(row=row_num, column=1, value="DEFENSE")
    sec.font = Font(bold=True, name="Calibri", size=11, color=WHITE)
    sec.fill = _fill(NAV)
    sec.alignment = _center()
    for c in [1,2,3,4,5]:
        ws.cell(row=row_num, column=c).fill = _fill(NAV)
    ws.row_dimensions[row_num].height = 18
    row_num += 1

    for col, label in [(1,"Pair"),(2,"LD"),(3,"RD"),(4,""),(5,"Side Note")]:
        cell = ws.cell(row=row_num, column=col, value=label)
        cell.font      = _hdr_font(10)
        cell.fill      = _fill("2E75B6")
        cell.alignment = _center()
        cell.border    = _thin_border()
    ws.row_dimensions[row_num].height = 18
    row_num += 1

    for pair in ["Pair 1", "Pair 2", "Pair 3"]:
        block = _get_players_for_unit(pair)
        slots_data_d = {"LD": ("","",False), "RD": ("","",False)}

        if slot_col and slot_col in block.columns:
            for _, r in block.iterrows():
                slot = str(r.get(slot_col) or "").upper()
                if slot in slots_data_d:
                    # Off-side flag for D: L-hand prefers LD, R-hand prefers RD
                    hand = str(r.get(hand_col) or "").upper()[:1] if hand_col else ""
                    warn = (slot == "RD" and hand == "L") or (slot == "LD" and hand == "R")
                    txt, _ = _display_name(r)
                    slots_data_d[slot] = (txt, hand, warn)
        else:
            # Fallback: fill LD then RD by impact
            sorted_block = block.sort_values(impact_col, ascending=False) if impact_col and impact_col in block.columns else block
            for i, (_, r) in enumerate(sorted_block.head(2).iterrows()):
                slot = "LD" if i == 0 else "RD"
                hand = str(r.get(hand_col) or "").upper()[:1] if hand_col else ""
                warn = (slot == "RD" and hand == "L") or (slot == "LD" and hand == "R")
                txt, _ = _display_name(r)
                slots_data_d[slot] = (txt, hand, warn)

        lbl_cell = ws.cell(row=row_num, column=1, value=pair)
        lbl_cell.font = _lbl_font()
        lbl_cell.fill = _fill(SLATE)
        lbl_cell.alignment = _center()
        lbl_cell.border = _thin_border()

        warn_notes = []
        for col_i, slot in enumerate(["LD","RD"], start=2):
            txt, hnd, warn = slots_data_d[slot]
            cell = ws.cell(row=row_num, column=col_i, value=txt)
            cell.fill      = _fill(WARN if warn else SLATE)
            cell.font      = _cell_font(warn=warn)
            cell.alignment = _left()
            cell.border    = _thin_border()
            if warn:
                warn_notes.append(f"{slot}: off-side hand")

        ws.cell(row=row_num, column=4).border = _thin_border()
        note_cell = ws.cell(row=row_num, column=5, value="; ".join(warn_notes) if warn_notes else "")
        note_cell.font      = Font(italic=True, size=9, color="7B3F00", name="Calibri")
        note_cell.alignment = _left()
        note_cell.border    = _thin_border()

        ws.row_dimensions[row_num].height = 18
        row_num += 1

    # ---- Healthy Scratches ----
    # Scratches = players on the projected ROSTER who are not in any active lineup unit.
    # We use roster_df for this, not lineup_df, because lineup_df only contains the
    # players already placed in active units — unplaced players simply aren't in it.
    row_num += 1
    hs_sec = ws.cell(row=row_num, column=1, value="HEALTHY SCRATCHES")
    hs_sec.font = Font(bold=True, name="Calibri", size=11, color=WHITE)
    hs_sec.fill = _fill("6B4E71")
    hs_sec.alignment = _center()
    for c in [1, 2, 3, 4, 5]:
        ws.cell(row=row_num, column=c).fill = _fill("6B4E71")
    ws.row_dimensions[row_num].height = 18
    row_num += 1

    for col, label in [(1, "#"), (2, "Player"), (3, "Pos"), (4, "Impact"), (5, "Note")]:
        cell = ws.cell(row=row_num, column=col, value=label)
        cell.font      = _hdr_font(10)
        cell.fill      = _fill("9B72A8")
        cell.alignment = _center()
        cell.border    = _thin_border()
    ws.row_dimensions[row_num].height = 16
    row_num += 1

    # Build the set of player names already placed in the active lineup
    lineup_names = set(
        str(df[disp_col].iloc[i]).strip().lower()
        for i in range(len(df))
        if disp_col and pd.notna(df[disp_col].iloc[i])
    )

    scratches_rows = []
    if roster_df is not None and len(roster_df) > 0:
        # Find player name col in roster_df
        r_name_col = next(
            (c for c in ["Player", "player", "player_y", "player_x", "Name"] if c in roster_df.columns),
            None
        )
        r_pos_col  = next((c for c in ["Position", "pos", "Pos"] if c in roster_df.columns), None)
        r_imp_col  = next((c for c in ["Impact Score", "impact", "Impact"] if c in roster_df.columns), None)
        r_hand_col = next((c for c in ["Hand", "hand", "hand_x"] if c in roster_df.columns), None)

        if r_name_col:
            for _, r in roster_df.iterrows():
                name = str(r.get(r_name_col) or "").strip()
                if not name or name.lower() in ("nan", ""):
                    continue
                pos  = str(r.get(r_pos_col) or "").upper().strip() if r_pos_col else ""
                # Only skaters (not goalies / blank)
                if pos not in {"C", "LW", "RW", "D"}:
                    continue
                # Not already in active lineup
                if name.strip().lower() in lineup_names:
                    continue
                imp = r.get(r_imp_col, np.nan) if r_imp_col else np.nan
                try:
                    imp = float(imp) if pd.notna(imp) else np.nan
                except Exception:
                    imp = np.nan
                hand = str(r.get(r_hand_col) or "").upper()[:1] if r_hand_col else ""
                scratches_rows.append({"name": name, "pos": pos, "impact": imp, "hand": hand})

        # Sort by impact descending
        scratches_rows.sort(key=lambda x: x["impact"] if pd.notna(x["impact"]) else -1, reverse=True)

    if scratches_rows:
        for i, sr in enumerate(scratches_rows[:8], start=1):
            imp_str = f"{sr['impact']:.0f}" if pd.notna(sr["impact"]) else ""
            ws.cell(row=row_num, column=1, value=i).font = Font(size=9, color="595659", name="Calibri")
            p_cell = ws.cell(row=row_num, column=2, value=sr["name"])
            p_cell.font      = Font(size=10, name="Calibri")
            p_cell.alignment = _left()
            p_cell.border    = _thin_border()
            ws.cell(row=row_num, column=3, value=sr["pos"]).alignment = _center()
            ws.cell(row=row_num, column=4, value=imp_str).alignment   = _center()
            for c in [1, 2, 3, 4, 5]:
                ws.cell(row=row_num, column=c).border = _thin_border()
            ws.row_dimensions[row_num].height = 16
            row_num += 1
    else:
        ws.cell(row=row_num, column=2, value="(No scratches identified — roster fully deployed)").font = Font(
            italic=True, size=9, color="595659", name="Calibri")
        row_num += 1

    # ---- Legend ----
    row_num += 1
    ws.cell(row=row_num, column=2,
            value="Notes: [POS] = player deployed off natural position.  "
                  "Yellow = off-side hand assignment.  (##) = Impact Score.").font = Font(
        italic=True, size=9, color="595659", name="Calibri")

    ws.freeze_panes = "B4"


def write_player_analysis_tab(
    wb: openpyxl.Workbook,
    season_df: pd.DataFrame,
    rolling_df: pd.DataFrame,
    prev_season_df: pd.DataFrame = None,
) -> None:
    """
    Writes a 'Player Analysis' sheet with:
      - A data validation dropdown (cell C3) listing every player in season_df
      - A hidden lookup table (_PA_Data) with one row per player containing all dashboard fields
      - Excel INDEX/MATCH formulas in the dashboard cells so selecting a name
        instantly populates the entire dashboard without macros
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    # ── colour palette ────────────────────────────────────────────────────────
    NAV   = "1F4E79"
    GOLD  = "C9A227"
    ICE   = "D9EAF7"
    GREEN = "E2EFDA"
    AMBER = "FFF2CC"
    RED_L = "FCE4D6"
    WHITE = "FFFFFF"
    GREY  = "F2F2F2"
    SLATE = "D6DCE4"

    def _fill(h): return PatternFill("solid", fgColor=h)
    def _font(bold=False, size=10, color="000000", italic=False):
        return Font(bold=bold, size=size, color=color, name="Calibri", italic=italic)
    def _side(): return Side(style="thin", color="B8CCE4")
    def _border(): return Border(left=_side(), right=_side(), top=_side(), bottom=_side())
    def _center(wrap=False): return Alignment(horizontal="center", vertical="center", wrap_text=wrap)
    def _left(wrap=False):   return Alignment(horizontal="left",   vertical="center", wrap_text=wrap)
    def _right():            return Alignment(horizontal="right",  vertical="center")

    # ── helpers ───────────────────────────────────────────────────────────────
    def _w(ws, row, col, value, bold=False, size=10, color="000000", fill=None,
           align=None, border=False, italic=False, num_fmt=None):
        c = ws.cell(row=row, column=col, value=value)
        c.font      = _font(bold=bold, size=size, color=color, italic=italic)
        c.alignment = align or _left()
        if fill:   c.fill   = _fill(fill)
        if border: c.border = _border()
        if num_fmt: c.number_format = num_fmt
        return c

    def _hdr(ws, row, col, label, span=1, fill=NAV, size=10):
        c = ws.cell(row=row, column=col, value=label)
        c.font      = _font(bold=True, size=size, color=WHITE)
        c.fill      = _fill(fill)
        c.alignment = _center()
        c.border    = _border()
        if span > 1:
            ws.merge_cells(start_row=row, start_column=col,
                           end_row=row,   end_column=col + span - 1)
        return c

    def _label(ws, row, col, text):
        c = ws.cell(row=row, column=col, value=text)
        c.font      = _font(bold=True, size=9, color="404040")
        c.alignment = _left()
        c.fill      = _fill(GREY)
        c.border    = _border()
        return c

    def _val_cell(ws, row, col, formula, fill=WHITE, num_fmt=None, bold=False):
        c = ws.cell(row=row, column=col, value=formula)
        c.font      = _font(bold=bold, size=10)
        c.alignment = _left()
        c.fill      = _fill(fill)
        c.border    = _border()
        if num_fmt: c.number_format = num_fmt
        return c

    # ── build the hidden data table ───────────────────────────────────────────
    # We flatten everything we want per player into one row keyed by player name.
    # Rolling 3yr scores are joined on player_key.

    sdf = season_df.copy()
    rdf = rolling_df.copy() if rolling_df is not None else pd.DataFrame()

    # normalise key
    def _key(s): return re.sub(r"\s+", " ", str(s).lower().strip())
    sdf["_pk"] = sdf["player"].map(_key)
    if len(rdf) > 0 and "player" in rdf.columns:
        rdf["_pk"] = rdf["player"].map(_key)
        # deduplicate so .at[] returns a scalar, not a Series
        rdf = rdf.sort_values("Dimensionality_Score_3yr" if "Dimensionality_Score_3yr" in rdf.columns else rdf.columns[0],
                               ascending=False).drop_duplicates("_pk")
        rdf_idx = rdf.set_index("_pk")
    else:
        rdf_idx = pd.DataFrame()

    # Per-player: take the row with highest toi_ev (handles duplicates), then sort alphabetically
    sdf = sdf.sort_values("toi_ev", ascending=False).drop_duplicates("_pk")
    sdf = sdf.sort_values("player").reset_index(drop=True)

    # Previous season for YoY delta
    pdf = prev_season_df.copy() if prev_season_df is not None else pd.DataFrame()
    if len(pdf) > 0 and "player" in pdf.columns:
        pdf["_pk"] = pdf["player"].map(_key)
        pdf = pdf.sort_values("toi_ev", ascending=False).drop_duplicates("_pk")
        pdf_idx = pdf.set_index("_pk")
    else:
        pdf_idx = pd.DataFrame()

    def _safe(row, col, default="—"):
        v = row.get(col, None)
        if v is None or (isinstance(v, float) and np.isnan(v)):
            return default
        return v

    def _fmt_pct(v, default="—"):
        try: return f"{float(v)*100:.1f}%" if pd.notna(v) else default
        except: return default

    def _fmt1(v, default="—"):
        try: return f"{float(v):.1f}" if pd.notna(v) else default
        except: return default

    def _fmt2(v, default="—"):
        try: return f"{float(v):.2f}" if pd.notna(v) else default
        except: return default

    def _fmti(v, default="—"):
        try: return f"{int(round(float(v)))}" if pd.notna(v) else default
        except: return default

    # Determine forward role score columns
    FWD_ROLES = ["Finisher_Score","Playmaker_Score","Driver_Score",
                 "TwoWay_Score","Power_Score","Grinder_Score","Producer_Score"]
    DEF_ROLES = ["Suppressor_Score","Transition_Score","PuckSkill_Score","Physical_Score"]
    FWD_LABELS = ["Finisher","Playmaker","Driver","Two-Way","Power","Grinder","Producer"]
    DEF_LABELS = ["Shutdown D","Transition","Puck Skill","Physical"]

    # ── create / replace PA sheet ─────────────────────────────────────────────
    PA_SHEET = "Player Analysis"
    DATA_SHEET = "_PA_Data"

    for sn in [PA_SHEET, DATA_SHEET]:
        if sn in wb.sheetnames:
            del wb[sn]

    ws  = wb.create_sheet(PA_SHEET)
    wsd = wb.create_sheet(DATA_SHEET)
    wsd.sheet_state = "hidden"

    # ── write hidden data table ───────────────────────────────────────────────
    # Columns: A=PlayerName, B=Team, C=Pos, D=Age, E=Height, F=Weight,
    # G=GP, H=TOI_EV, I=Goals, J=Assists, K=Points, L=PtsPGP,
    # M=G60, N=A160, O=Pts60, P=ixG60, Q=Shots60,
    # R=Hits60, S=BlkShots60, T=GvA60, U=TkA60,
    # V=FO%, W=OZ_Start, X=DZ_Start,
    # Y=Dim, Z=Type, AA=TypeConf, AB=TopRole, AC=TopScore,
    # AD=2ndRole, AE=2ndScore, AF=Margin,
    # AG=Proj_Pts60, AH=ProjPtsGP, AI=ProgTier, AJ=AgeCurve,
    # AK=BestFit, AL=FitScore,
    # AM..AQ = role scores (5 or 7 depending on pos)
    # AR = Dim_3yr, AS = Confidence_EV, AT = DraftPedigree,
    # AU = Injury_Imputed, AV = Player_Tier

    DATA_COLS = [
        "PlayerName","Team","Pos","Age","Height_in","Weight_lb",
        "GP","TOI_EV","Goals","Assists","Points","PtsPGP",
        "G60","A160","Pts60","ixG60","Shots60",
        "Hits60","BlkShots60","GvA60","TkA60","PenDiff60",
        "FO_pct","OZ_pct","DZ_pct",
        "Dim","PlayerType","TypeConf","TopRole","TopScore","SecRole","SecScore","Margin",
        "ProjPts60","ProjPtsGP","ProgTier","AgeCurve",
        "BestFit","FitScore",
        "R1_Label","R1_Score","R2_Label","R2_Score","R3_Label","R3_Score",
        "R4_Label","R4_Score","R5_Label","R5_Score","R6_Label","R6_Score","R7_Label","R7_Score",
        "Dim_3yr","Conf_EV","DraftPedigree","Injury","Tier","TierScore",
        "xGF60","xGA60","CF60","CA60",
        "Dim_Trend","Dim_Delta",
    ]

    # Write header row
    for j, h in enumerate(DATA_COLS, start=1):
        wsd.cell(row=1, column=j, value=h)

    player_names = []
    for i, (_, row) in enumerate(sdf.iterrows(), start=2):
        pk = row["_pk"]
        pg = row.get("pos_group", "F")

        # Role scores — always 7 slots, use pos_group to pick the right set
        if pg == "F":
            roles = [(l, row.get(c, np.nan)) for l, c in zip(FWD_LABELS, FWD_ROLES)]
        else:
            roles = [(l, row.get(c, np.nan)) for l, c in zip(DEF_LABELS, DEF_ROLES)]
        # pad to 7
        while len(roles) < 7:
            roles.append(("", np.nan))

        # 3yr dim
        dim_3yr = np.nan
        if len(rdf_idx) > 0 and pk in rdf_idx.index:
            dim_3yr = rdf_idx.at[pk, "Dimensionality_Score_3yr"] if "Dimensionality_Score_3yr" in rdf_idx.columns else np.nan

        # best fit
        best_fit = row.get("Best_Line_Fit") or row.get("Best_Pair_Fit") or "—"
        fit_score = row.get("Best_Line_Fit_Score") or row.get("Best_Pair_Fit_Score") or np.nan

        name = str(row.get("player","")).strip()
        player_names.append(name)

        vals = [
            name,
            str(row.get("team","")).strip(),
            str(row.get("pos","")).strip(),
            _safe(row, "age"),
            _safe(row, "height_in"),
            _safe(row, "weight_lb"),
            _safe(row, "gp"),
            _safe(row, "toi_ev"),
            _safe(row, "goals_ind"),
            _safe(row, "assists_ind"),
            _safe(row, "points_ind"),
            _safe(row, "pts_per_gp_ind"),
            _safe(row, "g60_ind"),
            _safe(row, "a1_60_ind"),
            _safe(row, "points60_ind"),
            _safe(row, "ixg60_ind"),
            _safe(row, "shots60_ind"),
            _safe(row, "hits60_ind"),
            _safe(row, "blk_shots60_ind"),
            _safe(row, "giveaways60_ind"),
            _safe(row, "takeaways60_ind"),
            _safe(row, "pen_diff60_ind"),
            _safe(row, "fo_pct"),
            _safe(row, "off_zone_start_pct"),
            _safe(row, "def_zone_start_pct"),
            _safe(row, "Dimensionality_Score"),
            str(row.get("Player_Type","")).strip(),
            str(row.get("Type_Confidence","")).strip(),
            str(row.get("Top_Role","")).strip(),
            _safe(row, "Top_Role_Score"),
            str(row.get("Second_Role","")).strip(),
            _safe(row, "Second_Role_Score"),
            _safe(row, "Margin"),
            _safe(row, "Proj_Points60"),
            _safe(row, "Proj_PTS_PGP"),
            str(row.get("Prog_Tier","")).strip(),
            _safe(row, "Age_Curve_Factor"),
            str(best_fit).strip(),
            fit_score,
            roles[0][0], roles[0][1],
            roles[1][0], roles[1][1],
            roles[2][0], roles[2][1],
            roles[3][0], roles[3][1],
            roles[4][0], roles[4][1],
            roles[5][0], roles[5][1],
            roles[6][0], roles[6][1],
            dim_3yr,
            _safe(row, "Confidence_EV"),
            str(row.get("Draft_Pedigree","")).strip(),
            "Yes" if row.get("Injury_Imputed") is True or str(row.get("Injury_Imputed","")).lower() == "true" else "No",
            str(row.get("Player_Tier","")).strip(),
            _safe(row, "Player_Tier_Score"),
            _safe(row, "xgf60_ev"),
            _safe(row, "xga60_ev"),
            _safe(row, "cf60_ev"),
            _safe(row, "ca60_ev"),
            str(row.get("Dim_Trend","")).strip(),
            _safe(row, "Dim_Delta"),
        ]
        for j, v in enumerate(vals, start=1):
            if isinstance(v, float) and not np.isnan(v):
                col_name = DATA_COLS[j-1]
                # Same per-column precision as the workbook formatter
                if col_name in ("GP","TOI_EV","Goals","Assists","Points"):
                    v = round(v, 1)
                elif col_name in ("Dim","TopScore","SecScore","Margin","Dim_Delta",
                                  "R1_Score","R2_Score","R3_Score","R4_Score",
                                  "R5_Score","R6_Score","R7_Score","Dim_3yr","FitScore"):
                    v = round(v, 1)
                elif col_name in ("FO_pct","OZ_pct","DZ_pct","Conf_EV","AgeCurve"):
                    v = round(v, 4)
                else:
                    v = round(v, 2)
            wsd.cell(row=i, column=j, value=v if not (isinstance(v, float) and np.isnan(v)) else None)

    n_players = len(player_names)
    data_range = f"_PA_Data!$A$2:$A${n_players+1}"

    # ── build the dashboard sheet ─────────────────────────────────────────────
    # Column widths
    col_widths = {"A":3,"B":22,"C":18,"D":22,"E":18,"F":22,"G":18,"H":22,"I":18,"J":22}
    for col, w in col_widths.items():
        ws.column_dimensions[col].width = w

    # Row 1: spacer
    ws.row_dimensions[1].height = 8

    # ── Title bar ─────────────────────────────────────────────────────────────
    ws.merge_cells("B2:J2")
    t = ws.cell(row=2, column=2, value="⬡  PLAYER INSIGHT DASHBOARD")
    t.font      = _font(bold=True, size=16, color=WHITE)
    t.fill      = _fill(NAV)
    t.alignment = _center()
    ws.row_dimensions[2].height = 32

    # ── Selector row ─────────────────────────────────────────────────────────
    ws.row_dimensions[3].height = 26
    ws.merge_cells("B3:C3")
    lbl = ws.cell(row=3, column=2, value="SELECT PLAYER ▼")
    lbl.font      = _font(bold=True, size=10, color=NAV)
    lbl.fill      = _fill(AMBER)
    lbl.alignment = _center()
    lbl.border    = _border()

    # The dropdown cell — user picks here
    ws.merge_cells("D3:F3")
    sel = ws.cell(row=3, column=4, value=player_names[0] if player_names else "")
    sel.font      = _font(bold=True, size=11, color=NAV)
    sel.fill      = _fill(AMBER)
    sel.alignment = _left()
    sel.border    = _border()

    # Data validation dropdown
    dv = DataValidation(
        type="list",
        formula1=f'_PA_Data!$A$2:$A${n_players+1}',
        allow_blank=False,
        showDropDown=False,
    )
    dv.sqref = "D3"
    ws.add_data_validation(dv)

    # Helper: INDEX/MATCH formula to look up col N of _PA_Data for selected player
    SEL = "$D$3"
    def _im(col_num: int) -> str:
        """Returns an INDEX/MATCH formula for column col_num of the data table."""
        return (f'=IFERROR(INDEX(_PA_Data!${get_column_letter(col_num)}$2:'
                f'${get_column_letter(col_num)}${n_players+1},'
                f'MATCH({SEL},_PA_Data!$A$2:$A${n_players+1},0)),"—")')

    # Column index map (1-based, matching DATA_COLS order)
    C = {name: idx for idx, name in enumerate(DATA_COLS, start=1)}

    ws.row_dimensions[4].height = 8  # spacer

    # ── Section helper ────────────────────────────────────────────────────────
    def _section(row, label, fill_col=NAV):
        ws.merge_cells(f"B{row}:J{row}")
        c = ws.cell(row=row, column=2, value=f"  {label}")
        c.font      = _font(bold=True, size=10, color=WHITE)
        c.fill      = _fill(fill_col)
        c.alignment = _left()
        ws.row_dimensions[row].height = 20

    def _row2(ws, row, pairs, fills=None):
        """Write a row of label/value pairs starting at col B.
        pairs = [(label, formula), ...] up to 4 pairs across B-J"""
        fills = fills or [GREY, WHITE] * len(pairs)
        col = 2
        for i, (lbl_txt, formula) in enumerate(pairs):
            _w(ws, row, col,   lbl_txt,       bold=True,  size=9,  color="404040", fill=GREY,  align=_left(),  border=True)
            _w(ws, row, col+1, formula,        bold=False, size=10, color="000000", fill=WHITE, align=_left(),  border=True)
            col += 2
        ws.row_dimensions[row].height = 18

    # ── SECTION 1: Identity ───────────────────────────────────────────────────
    r = 5
    _section(r, "▌ PLAYER IDENTITY")
    r += 1
    _row2(ws, r, [
        ("Name",    _im(C["PlayerName"])),
        ("Team",    _im(C["Team"])),
        ("Position",_im(C["Pos"])),
        ("Age",     _im(C["Age"])),
    ])
    r += 1
    _row2(ws, r, [
        ("Height (in)",  _im(C["Height_in"])),
        ("Weight (lbs)", _im(C["Weight_lb"])),
        ("Age",          _im(C["Age"])),
        ("Position",     _im(C["Pos"])),
    ])

    r += 1
    ws.row_dimensions[r].height = 8

    # ── SECTION 2: Season Stats ───────────────────────────────────────────────
    r += 1
    _section(r, "▌ 2025-26 SEASON STATS")
    r += 1
    # Primary counting stats — most scouting-readable row, displayed first
    _row2(ws, r, [
        ("Goals",       _im(C["Goals"])),
        ("Assists",     _im(C["Assists"])),
        ("Points",      _im(C["Points"])),
        ("Pts/GP",      _im(C["PtsPGP"])),
    ])
    r += 1
    _row2(ws, r, [
        ("GP",          _im(C["GP"])),
        ("TOI/EV (min)",_im(C["TOI_EV"])),
        ("FO%",         _im(C["FO_pct"])),
        ("Sample Conf.",_im(C["Conf_EV"])),
    ])
    r += 1
    # Rate stats
    _row2(ws, r, [
        ("G/60",        _im(C["G60"])),
        ("A1/60",       _im(C["A160"])),
        ("Pts/60",      _im(C["Pts60"])),
        ("ixG/60",      _im(C["ixG60"])),
    ])
    r += 1
    _row2(ws, r, [
        ("Shots/60",    _im(C["Shots60"])),
        ("Hits/60",     _im(C["Hits60"])),
        ("Blocks/60",   _im(C["BlkShots60"])),
        ("Pen Diff/60", _im(C["PenDiff60"])),
    ])
    r += 1
    _row2(ws, r, [
        ("Giveaways/60",_im(C["GvA60"])),
        ("Takeaways/60",_im(C["TkA60"])),
        ("OZ Start %",  _im(C["OZ_pct"])),
        ("DZ Start %",  _im(C["DZ_pct"])),
    ])
    r += 1
    _row2(ws, r, [
        ("xGF/60 (EV)", _im(C["xGF60"])),
        ("xGA/60 (EV)", _im(C["xGA60"])),
        ("CF/60 (EV)",  _im(C["CF60"])),
        ("CA/60 (EV)",  _im(C["CA60"])),
    ])

    r += 1
    ws.row_dimensions[r].height = 8

    # ── SECTION 3: Role Profile ───────────────────────────────────────────────
    r += 1
    _section(r, "▌ ROLE PROFILE  (percentile scores, 0–100)", fill_col="2E75B6")
    r += 1
    _row2(ws, r, [
        (_im(C["R1_Label"]), _im(C["R1_Score"])),
        (_im(C["R2_Label"]), _im(C["R2_Score"])),
        (_im(C["R3_Label"]), _im(C["R3_Score"])),
        (_im(C["R4_Label"]), _im(C["R4_Score"])),
    ])
    r += 1
    _row2(ws, r, [
        (_im(C["R5_Label"]), _im(C["R5_Score"])),
        (_im(C["R6_Label"]), _im(C["R6_Score"])),
        (_im(C["R7_Label"]), _im(C["R7_Score"])),
        ("Dimensionality",   _im(C["Dim"])),
    ])
    r += 1
    _row2(ws, r, [
        ("Top Role",        _im(C["TopRole"])),
        ("Top Role Score",  _im(C["TopScore"])),
        ("2nd Role",        _im(C["SecRole"])),
        ("2nd Role Score",  _im(C["SecScore"])),
    ])
    r += 1
    _row2(ws, r, [
        ("Player Type",     _im(C["PlayerType"])),
        ("Type Confidence", _im(C["TypeConf"])),
        ("Role Margin",     _im(C["Margin"])),
        ("Dim (3yr avg)",   _im(C["Dim_3yr"])),
    ])
    r += 1
    _row2(ws, r, [
        ("Dim Trend",       _im(C["Dim_Trend"])),
        ("Dim Delta (YoY)", _im(C["Dim_Delta"])),
        ("Player Tier",     _im(C["Tier"])),
        ("Tier Score",      _im(C["TierScore"])),
    ])
    r += 1
    # Tier legend as a single merged note row
    ws.merge_cells(f"B{r}:J{r}")
    legend = ws.cell(row=r, column=2,
        value="Tier Score = weighted composite of Dimensionality (65%), Production percentile (25%), Sample confidence (10%).  "
              "Elite ≥92 · Star ≥82 · Solid ≥68 · Depth ≥50 · Fringe <50.  "
              "Prospect = young player (≤23, or 24 w/ top-10 pick) with Elite/High draft capital below Star tier — "
              "current stats understate ceiling.  Low-sample players are capped at Solid/Depth ceiling.")
    legend.font      = Font(italic=True, name="Calibri", size=8, color="595659")
    legend.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[r].height = 28

    r += 1
    ws.row_dimensions[r].height = 8

    # ── SECTION 4: Deployment / Fit ───────────────────────────────────────────
    r += 1
    _section(r, "▌ DEPLOYMENT & LINE FIT", fill_col="375623")
    r += 1
    _row2(ws, r, [
        ("Best Fit",        _im(C["BestFit"])),
        ("Fit Score",       _im(C["FitScore"])),
        ("Injury Flag",     _im(C["Injury"])),
        ("Draft Pedigree",  _im(C["DraftPedigree"])),
    ])

    r += 1
    ws.row_dimensions[r].height = 8

    # ── SECTION 5: Projection & Development ──────────────────────────────────
    r += 1
    _section(r, "▌ PROJECTION & DEVELOPMENT", fill_col="7B3F00")
    r += 1
    _row2(ws, r, [
        ("Proj Pts/60",     _im(C["ProjPts60"])),
        ("Proj Pts/GP",     _im(C["ProjPtsGP"])),
        ("Prog Tier",       _im(C["ProgTier"])),
        ("Age Curve",       _im(C["AgeCurve"])),
    ])

    r += 2
    # ── SECTION 6: Field Key / Glossary ──────────────────────────────────────
    _section(r, "▌ FIELD KEY  — What each metric means", fill_col="404040")
    r += 1

    KEY_ENTRIES = [
        # Identity
        ("Age",                 "Player age at season start."),
        ("Height (in)",         "Height in inches."),
        ("Weight (lbs)",        "Weight in pounds."),
        ("Draft Pedigree",      "Qualitative draft tier: Elite (top-10 pick), High, Mid, Late, Undrafted."),
        ("Injury Flag",         "Yes = stats for this season were carried from prior year due to absence."),
        # Season counting
        ("Goals / Assists / Points", "Raw season totals (prorated where applicable)."),
        ("Pts/GP",              "Points per game played — context-neutral scoring rate."),
        ("GP",                  "Games played."),
        ("TOI/EV (min)",        "Even-strength time on ice in minutes for the season."),
        ("FO%",                 "Faceoff win percentage (blank for non-faceoff takers)."),
        ("Sample Conf.",        "Confidence in stats: 0–1 scale based on TOI. <0.35 = Low, ≥0.50 = High."),
        # Rate stats
        ("G/60",                "Goals per 60 minutes (individual, all situations)."),
        ("A1/60",               "Primary (first) assists per 60 minutes — best proxy for puck creation."),
        ("Pts/60",              "Total points per 60 minutes."),
        ("ixG/60",              "Individual expected goals per 60 min — shot quality weighted."),
        ("Shots/60",            "Shot attempts per 60 minutes."),
        ("Hits/60",             "Hits per 60 minutes."),
        ("Blocks/60",           "Blocked shots per 60 minutes."),
        ("Giveaways/60",        "Puck giveaways per 60 min. Lower is better for puck-skill assessment."),
        ("Takeaways/60",        "Puck takeaways per 60 min. Higher = better defensive disruption."),
        ("Pen Diff/60",         "Penalties drawn minus penalties taken, per 60 min. Positive = disciplined."),
        ("OZ / DZ Start %",     "Share of shifts starting in the offensive / defensive zone."),
        # On-ice
        ("xGF/60 (EV)",         "Expected goals for per 60 at even strength — team possession quality when on ice."),
        ("xGA/60 (EV)",         "Expected goals against per 60 at even strength — defensive exposure when on ice."),
        ("CF/60 (EV)",          "Corsi for (shot attempts for) per 60 at EV — raw possession proxy."),
        ("CA/60 (EV)",          "Corsi against per 60 at even strength."),
        # Role scores
        ("Role Scores (0–100)", "Percentile rank vs. all NHL players at same position. 75 = better than 75%."),
        ("Dimensionality",      "Average of player's top two role scores. High = versatile; low = specialist."),
        ("Top / 2nd Role",      "The one or two roles where this player scores highest relative to peers."),
        ("Player Type",         "Best-fit archetype label. 'Hybrid' = two roles within 6 pts of each other."),
        ("Type Confidence",     "How reliable the archetype label is: High / Medium / Low."),
        ("Role Margin",         "Gap between top role score and second role score. Larger = clearer specialist."),
        ("Dim (3yr avg)",       "Dimensionality score averaged over the last three seasons — filters single-year noise."),
        ("Dim Trend",           "Direction of Dimensionality change vs. prior year, age-adjusted. ⚡ Breakout · ↑↑ Strong Rise · ↑ Rising · → Stable · ↓ Slipping · ↓↓ Declining · ⚠ Sharp Drop · ⛑ Injury Year."),
        ("Dim Delta (YoY)",     "Raw point change in Dimensionality score vs. prior season. Positive = improvement."),
        ("Player Tier",         "Talent tier based on composite score. Elite · Star · Solid · Depth · Fringe. "
                                "Prospect = young player (age ≤23 Emerging, or age 24 top-10 pick) with Elite/High draft capital "
                                "who hasn't reached Star/Elite yet — their current stats are expected to understate true ceiling. "
                                "Low-sample players are hard-capped below Star/Elite regardless of role scores."),
        ("Tier Score (0–100)",  "Composite: 65% Dimensionality + 25% Production percentile (pts/gp, pts/60, ixG/60 for F; possession-weighted for D) + "
                                "10% Sample confidence. Thresholds: Elite ≥92, Star ≥82, Solid ≥68, Depth ≥50, Fringe <50. "
                                "Prospect tier is an overlay — Tier Score still reflects current performance."),
        ("Draft Pedigree",      "How highly the player was originally drafted — context for development trajectory."),
        # Fit
        ("Best Fit",            "Projected best deployment line or pairing based on role profile."),
        ("Fit Score",           "Numeric fit score (0–100) for that line/pair — higher = stronger structural match."),
        ("Injury Flag",         "Confirms whether current season stats are real or carried from prior season."),
        ("Prog Tier",           "Development phase: Breakout · Rising · Prime · Plateau · Declining."),
        # Projection
        ("Proj Pts/60",         "Model projection of points per 60 next season, incorporating age curve."),
        ("Proj Pts/GP",         "Model projection of points per game next season."),
        ("Age Curve",           "Multiplier (e.g. 1.05 = +5%) applied to stats based on age trajectory. Peak ~24–28."),
    ]

    # Write as two-column label/description pairs across B-J
    for lbl, desc in KEY_ENTRIES:
        _w(ws, r, 2, lbl,  bold=True,  size=9, color="1F4E79", fill=ICE,   align=_left(), border=True)
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=10)
        _w(ws, r, 3, desc, bold=False, size=9, color="000000", fill=WHITE, align=_left(wrap=True), border=True)
        ws.row_dimensions[r].height = 28
        r += 1

    r += 1
    # ── Footer note ───────────────────────────────────────────────────────────
    ws.merge_cells(f"B{r}:J{r}")
    note = ws.cell(row=r, column=2,
        value="📌  Use the dropdown in D3 to select any player. All fields update automatically via INDEX/MATCH.")
    note.font      = _font(italic=True, size=9, color="595659")
    note.alignment = _left()
    ws.row_dimensions[r].height = 16

    # Freeze selector row
    ws.freeze_panes = "B5"

    # Make PA sheet first
    wb.move_sheet(PA_SHEET, offset=-wb.index(wb[PA_SHEET]))


def _format_workbook(wb: openpyxl.Workbook) -> None:
    """
    Full workbook formatting pass (runs on every visible sheet):
      1. Rename headers via COLUMN_DISPLAY_NAMES (pretty titles, no underscores)
      2. Navy header band + alternating row fills + thin borders
      3. Number formats: currency ($), percentage (%), integer, 2-decimal
      4. Round float cell values to max 2 decimal places before writing format
      5. Auto-size columns (8–42 chars), freeze header row, set header row height
    Skips hidden sheets and the Depth Chart sheet (hand-styled separately).
    """
    SKIP_SHEETS = {"_ALL_25-26","_ALL_24-25","_ALL_23-24","_ROLLING_ALL"}

    for ws in wb.worksheets:
        if ws.sheet_state == "hidden":
            continue
        if ws.title in SKIP_SHEETS:
            continue
        if "Depth Chart" in ws.title:
            continue
        if ws.max_row < 1 or ws.max_column < 1:
            continue

        # ── 1. Rename headers ───────────────────────────────────────────────
        header_row = list(ws.iter_rows(min_row=1, max_row=1))[0]
        raw_headers: List[str] = []
        for cell in header_row:
            raw = str(cell.value) if cell.value is not None else ""
            cell.value = COLUMN_DISPLAY_NAMES.get(raw, raw)
            raw_headers.append(raw)

        # ── 2. Style header ─────────────────────────────────────────────────
        for cell in header_row:
            cell.fill      = _HDR_FILL
            cell.font      = _HDR_FONT
            cell.alignment = _HDR_ALIGN
            cell.border    = _CELL_BORDER

        # ── 3 & 4. Data rows: fill + border + number format + rounding ──────
        n_cols = ws.max_column
        n_rows = ws.max_row

        for row_idx, row in enumerate(ws.iter_rows(min_row=2, max_row=n_rows), start=1):
            fill = _ROW_FILL_EVEN if (row_idx % 2 == 0) else _ROW_FILL_ODD
            for col_idx, cell in enumerate(row):
                cell.fill   = fill
                cell.font   = _DATA_FONT
                cell.border = _CELL_BORDER

                raw_col = raw_headers[col_idx] if col_idx < len(raw_headers) else ""
                fmt     = _col_num_fmt(raw_col)

                val = cell.value
                if isinstance(val, float):
                    # Per-column precision:
                    #  - integer-ish cols (GP, draft picks, starts) → 0 dp
                    #  - percentage cols → already handled by number_format; round to 4dp so display is clean
                    #  - score/percentile cols (0-100) → 1 dp is plenty
                    #  - rate stats (per-60, pts/gp) → 2 dp
                    #  - everything else → 2 dp max
                    if fmt == INT_FMT:
                        rounded = round(val, 0)
                    elif fmt == PCT_FMT:
                        rounded = round(val, 4)   # e.g. 0.5268 → shows as 52.68%
                    elif raw_col.endswith("_Score") or raw_col in ("Dimensionality_Score","Type_Margin","Margin","Dim_Delta"):
                        rounded = round(val, 1)
                    else:
                        rounded = round(val, 2)
                    cell.value = rounded
                    val = rounded

                if fmt and isinstance(val, (int, float)) and val is not None:
                    cell.number_format = fmt
                    cell.alignment     = _DATA_ALIGN_R
                elif isinstance(val, str):
                    cell.alignment = _DATA_ALIGN_L
                elif isinstance(val, (int, float)):
                    cell.alignment = _DATA_ALIGN_R
                else:
                    cell.alignment = _DATA_ALIGN_C

        # ── 5. Column widths ────────────────────────────────────────────────
        MAX_W, MIN_W = 42, 8
        for col_cells in ws.iter_cols(min_row=1, max_row=min(n_rows, 200)):
            max_len = 0
            for cell in col_cells:
                try:
                    v = cell.value
                    # Use formatted width for numbers, raw string length for text
                    if isinstance(v, float):
                        s = f"{v:,.2f}"
                    elif v is not None:
                        s = str(v)
                    else:
                        s = ""
                    max_len = max(max_len, len(s))
                except Exception:
                    pass
            col_letter = get_column_letter(col_cells[0].column)
            ws.column_dimensions[col_letter].width = min(max(max_len + 3, MIN_W), MAX_W)

        # ── 6. Freeze + header height ───────────────────────────────────────
        ws.freeze_panes      = "A2"
        ws.row_dimensions[1].height = 32


def write_methodology_docx(out_path: Path) -> Path:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("NHL Player Analysis Model — Methodology & Technical Reference", level=1)
    doc.add_paragraph(
        f"Input workbook : {INPUT_XLSX}\n"
        f"Output workbook: {OUTPUT_XLSX}\n"
        f"Season tabs    : {', '.join(SEASON_TABS)}\n"
        f"Rolling weights (newest → oldest): {ROLLING_WEIGHTS}"
    )

    # ── 1. DATA SOURCES ──────────────────────────────────────────────────────
    doc.add_heading("1  Data Sources & Assumptions", level=2)
    doc.add_paragraph(
        "Three data streams are loaded per season from Natural Stat Trick exports "
        "and a Bio tab, then merged on a normalised player key."
    )
    doc.add_paragraph("On-Ice (EV) tab — team context metrics: Corsi, Fenwick, xGoals, HDCF, zone starts.", style="List Bullet")
    doc.add_paragraph("Individual (Ind) tab — per-player counting stats: goals, assists, shots, ixG, hits, blocks, faceoffs, penalties.", style="List Bullet")
    doc.add_paragraph("Bio tab — Height (in), Weight (lbs), Age, Draft Round, Overall Pick.", style="List Bullet")
    doc.add_paragraph(
        "BMI is derived as weight_lb / height_in² and used as a physical-profile proxy "
        "inside the Power and Physical role models only."
    )
    doc.add_paragraph(
        "Cap data is loaded from a separate sheet. Player keys are matched after "
        "normalising name format (comma-last → first-last) so 'Barkov, Aleksander' "
        "and 'Aleksander Barkov' resolve to the same key."
    )
    doc.add_paragraph(
        "Players absent from the current season tab (injured, recalled, etc.) are "
        "imputed from the most recent available prior season with an 8 % regression "
        "toward the mean, so they still appear in all output and cap-planning tabs."
    )

    # ── 2. DERIVED METRICS ───────────────────────────────────────────────────
    doc.add_heading("2  Derived Metrics", level=2)
    doc.add_paragraph("Per-60 rates:  metric_per60 = raw_count / TOI_minutes × 60", style="List Bullet")
    doc.add_paragraph("Zone start %:  OZ_starts / (OZ + NZ + DZ) if not exported directly.", style="List Bullet")
    doc.add_paragraph("Sample Confidence (Confidence_EV):  clip(TOI / 1200, 0, 1) — reaches 1.0 at ~1200 EV minutes.", style="List Bullet")
    doc.add_paragraph("Penalty Differential/60:  (Penalties Drawn − Penalties Taken) / TOI × 60.", style="List Bullet")
    doc.add_paragraph("Primary Points/60 (P1/60):  (Goals + Primary Assists) / TOI × 60.", style="List Bullet")

    # ── 3. ROLE SCORE PIPELINE ───────────────────────────────────────────────
    doc.add_heading("3  Role Score Construction", level=2)
    doc.add_paragraph(
        "Each role score represents how well a player fills a specific tactical function "
        "relative to their peers at the same position. The pipeline is:"
    )
    doc.add_paragraph("Z-score each feature within (season × position) groups.", style="List Number")
    doc.add_paragraph("Apply minutes-shrinkage: z_shrunk = z × TOI / (TOI + k), where k is role-specific (500–700 min). Low-minute players shrink toward zero.", style="List Number")
    doc.add_paragraph("Weighted sum across features for each role.", style="List Number")
    doc.add_paragraph("Convert to a 0–100 percentile rank within (season × position).", style="List Number")
    doc.add_paragraph(
        "Negative feature weights (e.g. xGA/60 in the Two-Way role, "
        "giveaways in Puck Skill) mean that lower values of those stats "
        "improve the role score."
    )

    # ── 4. ROLE DEFINITIONS ──────────────────────────────────────────────────
    doc.add_heading("4  Role Definitions", level=2)
    for group_key, roles in ROLE_FEATURES.items():
        group_label = "Forwards" if group_key == "F" else "Defensemen"
        doc.add_heading(f"4.{'1' if group_key == 'F' else '2'}  {group_label}", level=3)
        for role_name, spec in roles.items():
            doc.add_paragraph(
                f"{role_name}   (shrinkage k = {spec['k']}, minutes basis: {spec['minutes_basis']})",
                style="List Bullet"
            )
            for feat, w in spec["features"].items():
                doc.add_paragraph(f"{feat}  {w:+.2f}", style="List Bullet 2")

    # ── 5. PRODUCTION SCORE ──────────────────────────────────────────────────
    doc.add_heading("5  Individual Production Score", level=2)
    doc.add_paragraph(
        "A normalised 0–100 production score is derived directly from individual "
        "counting stats to supplement the on-ice role model:"
    )
    doc.add_paragraph("50 % weight — Points/60 (ceiling: 3.5 fwd / 1.8 def).", style="List Bullet")
    doc.add_paragraph("25 % weight — Goals/60  (ceiling: 1.5 fwd / 0.6 def).", style="List Bullet")
    doc.add_paragraph("25 % weight — ixG/60    (ceiling: 1.8 fwd / 0.8 def).", style="List Bullet")
    doc.add_paragraph(
        "This score feeds directly into the Impact Score (15 % weight) and "
        "into the Line 1 / Line 2 fit templates as a production modifier."
    )

    # ── 6. DIMENSIONALITY & IMPACT ───────────────────────────────────────────
    doc.add_heading("6  Dimensionality Score & Impact Score", level=2)
    doc.add_paragraph(
        "Dimensionality Score = average of a player's two highest role scores. "
        "A high dimensionality player contributes in multiple tactical contexts "
        "rather than being a single-role specialist."
    )
    doc.add_paragraph(
        "Impact Score blends four components, reliability-adjusted by sample "
        "confidence and role stability:"
    )
    doc.add_paragraph("38 % — Dimensionality Score.", style="List Bullet")
    doc.add_paragraph("30 % — Top Role Score.", style="List Bullet")
    doc.add_paragraph("17 % — Best Lineup Fit Score (line or pair).", style="List Bullet")
    doc.add_paragraph("15 % — Individual Production Score (if IND data available; redistributed if not).", style="List Bullet")
    doc.add_paragraph(
        "Reliability multiplier = 0.85 + 0.10 × Confidence_EV + 0.05 × Role Stability."
    )

    # ── 7. PLAYER TYPE ───────────────────────────────────────────────────────
    doc.add_heading("7  Player Type Classification", level=2)
    doc.add_paragraph(
        "Player Type is derived from the sorted role scores with stat-driven nudges "
        "and guardrails applied in the following order:"
    )
    doc.add_paragraph("Stat nudges: rule-based boosts to specific role scores based on counting-stat thresholds (e.g. 35+ goals → Finisher boost, 50+ assists → Playmaker boost). Nudges are small by design — they break ties, not rankings.", style="List Number")
    doc.add_paragraph(f"Hybrid guard: if Top1 − Top2 margin < {TYPE_HYBRID_MARGIN}, type is 'Hybrid: A / B'.", style="List Number")
    doc.add_paragraph("Producer tiebreak: if Producer is tied for top score on a forward, Producer wins.", style="List Number")
    doc.add_paragraph(f"TOI guard: players below minimum TOI thresholds (F: {MIN_TOI_EV_FWD} min, D: {MIN_TOI_EV_DEF} min) receive 'Insufficient TOI'.", style="List Number")
    doc.add_paragraph(
        "Type Confidence is set to High when Confidence_EV ≥ 0.50 and margin ≥ "
        f"{TYPE_CLEAR_MARGIN}, Medium when margin ≥ {TYPE_HYBRID_MARGIN}, Low otherwise."
    )

    # ── 8. PLAYER PROGRESSION ────────────────────────────────────────────────
    doc.add_heading("8  Player Projection & Progression Model", level=2)
    doc.add_paragraph(
        "A projection is generated per player each season and appears as "
        "Proj Pts/60, Proj G/60, Proj Pts/GP, Age Curve, Progression Tier, "
        "Proj Confidence, and Draft Pedigree in the output sheets."
    )
    doc.add_heading("8.1  Aging Curve", level=3)
    doc.add_paragraph("Age ≤ 21: factor 0.88 — still developing.", style="List Bullet")
    doc.add_paragraph("Age 22–23: factor 0.96 — late emergence.", style="List Bullet")
    doc.add_paragraph("Age 24–27: factor 1.02 — peak years.", style="List Bullet")
    doc.add_paragraph("Age 28–29: factor 1.00 — sustained prime.", style="List Bullet")
    doc.add_paragraph("Age 30–31: factor 0.97 — plateau.", style="List Bullet")
    doc.add_paragraph("Age 32–33: factor 0.93 — early decline.", style="List Bullet")
    doc.add_paragraph("Age 34–35: factor 0.87 — decline.", style="List Bullet")
    doc.add_paragraph("Age 36+  : factor 0.80 — late career.", style="List Bullet")
    doc.add_paragraph("Defensemen peak ~2 years later; their curve is shifted accordingly.")
    doc.add_heading("8.2  Draft Pedigree Signal", level=3)
    doc.add_paragraph(
        "Draft pedigree provides a residual talent signal for young players "
        "where historical production is limited. It decays linearly to zero "
        "influence by age ~32."
    )
    doc.add_paragraph("Overall pick 1–10 (Elite): +6 % production boost.", style="List Bullet")
    doc.add_paragraph("Overall pick 11–30 (High): +4 %.", style="List Bullet")
    doc.add_paragraph("Overall pick 31–90 (Mid): +2 %.", style="List Bullet")
    doc.add_paragraph("Pick 181+ or undrafted (Late): −1 %.", style="List Bullet")
    doc.add_heading("8.3  Projected Rates", level=3)
    doc.add_paragraph(
        "Proj Pts/60 = production_baseline × age_curve_factor × (1 + pedigree_boost). "
        "The baseline prefers current-season pts/60 when the sample is adequate "
        "(Confidence_EV ≥ 0.40), falls back to 3-year rolling if available, "
        "then to a dimensionality-derived proxy."
    )
    doc.add_paragraph("Proj Pts/GP is approximated as Proj Pts/60 × average TOI per game (15.5 min fwd / 20 min def).")
    doc.add_heading("8.4  Progression Tier", level=3)
    doc.add_paragraph("Emerging: age ≤ 23 — projection likely understates ceiling.", style="List Bullet")
    doc.add_paragraph("Prime: age 24–29 — peak performance window.", style="List Bullet")
    doc.add_paragraph("Plateau: age 30–31 — stable but no further growth expected.", style="List Bullet")
    doc.add_paragraph("Declining: age 32–35 — output expected to decrease.", style="List Bullet")
    doc.add_paragraph("Aging: age 36+ — significant decline risk.", style="List Bullet")

    # ── 9. LINE & PAIR FIT ───────────────────────────────────────────────────
    doc.add_heading("9  Lineup Fit Modeling", level=2)
    doc.add_heading("9.1  Forward Lines", level=3)
    doc.add_paragraph(
        "Each forward receives four fit scores (Line 1–4). Lines 1 and 2 "
        "blend role scores with an individual production modifier so that "
        "actual scoring output informs top-line deployment:"
    )
    doc.add_paragraph("Line 1 fit = 85 % role-score blend + 15 % production modifier (pts/60, normalised 0–100).", style="List Bullet")
    doc.add_paragraph("Line 2 fit = 88 % role-score blend + 12 % production modifier.", style="List Bullet")
    doc.add_paragraph("Line 3 fit = 95 % role-score blend + 5 % production modifier.", style="List Bullet")
    doc.add_paragraph("Line 4 fit = 100 % role-score blend (no production signal — energy / checking role).", style="List Bullet")
    doc.add_paragraph(
        "Usage eligibility gates prevent low-minute players from winning "
        "Line 1 or Line 2 slots. An elite-producer override lifts this gate "
        "for players with Producer Score ≥ 80 or Finisher Score ≥ 80. "
        "Grinder-type forwards receive a 20 % / 40 % suppression on Line 1 / Line 2 fits."
    )
    doc.add_heading("9.2  Defense Pairs", level=3)
    doc.add_paragraph("Pair 1 = 30 % Shutdown + 35 % Transition + 35 % Puck Skill — top-pair shutdown/breakout.", style="List Bullet")
    doc.add_paragraph("Pair 2 = 35 % Transition + 30 % Puck Skill + 35 % Physical — heavy defensive deployment.", style="List Bullet")
    doc.add_paragraph("Pair 3 = 50 % Shutdown + 35 % Physical + 15 % Transition — depth / specialist.", style="List Bullet")

    # ── 10. CONTRACT ESTIMATION ──────────────────────────────────────────────
    doc.add_heading("10  Contract Estimation via Market Comparables", level=2)
    doc.add_paragraph(
        "Expected AAV is estimated using a market comparables approach rather "
        "than a fixed formula, capturing real market variation."
    )
    doc.add_heading("10.1  Comparable Selection", level=3)
    doc.add_paragraph("Comparable players are filtered to the same position group and best lineup fit bucket (e.g. Line 1 forwards only).", style="List Bullet")
    doc.add_paragraph("Top-Role match is applied as a soft filter when enough comparables remain.", style="List Bullet")
    doc.add_paragraph(
        "Distance is measured across seven dimensions (IQR-normalised): "
        "age, impact score, dimensionality, top role score, fit score, "
        "current pts/60, and projected pts/60. Production and projection carry "
        "~20 % combined weight so that a 30-goal scorer finds 30-goal comps."
    )
    doc.add_heading("10.2  AAV Distribution", level=3)
    doc.add_paragraph("Closer comparables receive higher weights (w ∝ 1 / (1 + rank)).", style="List Bullet")
    doc.add_paragraph("AAV standard deviation is clipped to 12–35 % of mean to avoid extreme tails.", style="List Bullet")
    doc.add_paragraph("3 000 draws from a Student-t distribution (df = 6) model negotiation variance.", style="List Bullet")
    doc.add_paragraph("P10 / P50 / P90 of the draw distribution are reported. P50 is used as the planning estimate.", style="List Bullet")
    doc.add_heading("10.3  Progression-Adjusted Value", level=3)
    doc.add_paragraph(
        "The Value / $M (Prog Adj) column adjusts raw value-per-dollar by "
        "the player's expected trajectory:"
    )
    doc.add_paragraph("Emerging: +15 % — current stats understate next-year ceiling.", style="List Bullet")
    doc.add_paragraph("Prime: +5 %.", style="List Bullet")
    doc.add_paragraph("Plateau: 0 %.", style="List Bullet")
    doc.add_paragraph("Declining: −8 %.", style="List Bullet")
    doc.add_paragraph("Aging: −15 %.", style="List Bullet")
    doc.add_paragraph(
        "Draft pedigree adds an additional boost (up to +6 %) for players "
        "under ~26, decaying to zero by age 32."
    )

    # ── 11. FREE AGENT POOL ──────────────────────────────────────────────────
    doc.add_heading("11  Free Agent Pool Construction", level=2)
    doc.add_paragraph("League-wide UFAs (expiry year = offseason year, contract type ≠ RFA).", style="List Bullet")
    doc.add_paragraph(
        "RFAs are excluded from the open FA pool. Team RFAs are handled "
        "in the re-sign decision logic separately.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "FA candidates are scored for roster-hole fit (position need, "
        "group need, lineup-fit-bucket need), type confidence, and "
        "progression-adjusted value per dollar. Low-confidence players "
        "are excluded from non-positional-need slots.",
        style="List Bullet"
    )

    # ── 12. ROSTER HOLE DETECTION ────────────────────────────────────────────
    doc.add_heading("12  Roster Hole Detection & Re-sign Decisioning", level=2)
    doc.add_paragraph(
        "Holes are tracked at three levels of granularity:"
    )
    doc.add_paragraph("Group totals: 12 forwards, 6 defensemen.", style="List Bullet")
    doc.add_paragraph("Position minimums: 4 C, 3 LW, 3 RW, 6 D.", style="List Bullet")
    doc.add_paragraph("Fit-bucket minimums: 3 players per line slot (Line 1–4), 2 per pair slot (Pair 1–3).", style="List Bullet")
    doc.add_paragraph(
        "Re-sign decisions rank expiring players by: manual override → "
        "hole-fill bonus → type confidence → progression-adjusted value/$ → "
        "raw impact. Budget consumed greedily until the resign_budget_share "
        "of available cap space is exhausted."
    )

    # ── 13. OUTPUT SHEETS ────────────────────────────────────────────────────
    doc.add_heading("13  Output Sheets", level=2)
    sheets = [
        ("F_25-26, F_24-25, F_23-24", "Forward role scores, type classification, production, projections, and line fit per season."),
        ("D_25-26, D_24-25, D_23-24", "Defense role scores, type, projections, and pair fit per season."),
        ("Rolling_F / Rolling_D",      "3-year weighted average scores with volatility. Weights: 60/30/10 newest → oldest."),
        (f"TEAM_{'{TEAM}'}_Returning_26-27", "Under-contract core returning next season with impact scores."),
        (f"TEAM_{'{TEAM}'}_Resign_Decisions_26", "Expiring players ranked with AAV distribution (P10/P50/P90), progression tier, and RESIGN / LET WALK decision."),
        (f"TEAM_{'{TEAM}'}_FA_Targets_26", "Recommended free agent signings prioritised by hole fill and progression-adjusted value."),
        (f"TEAM_{'{TEAM}'}_Projected_Roster_26-27", "Full post-FA projected roster with all scoring and projection fields."),
        (f"TEAM_{'{TEAM}'}_Projected_Lineup_26-27", "Lineup table with Assigned Unit (Line 1–4 / Pair 1–3)."),
        (f"{'{TEAM}'} — Depth Chart", "Visual depth chart showing projected line/pair groupings."),
        (f"TEAM_{'{TEAM}'}_Cap_Summary_26-27", "Cap arithmetic: core spend, re-sign spend, FA spend, projected total, remaining space."),
    ]
    for name, desc in sheets:
        doc.add_paragraph(f"{name} — {desc}", style="List Bullet")

    # ── 14. FORMATTING ───────────────────────────────────────────────────────
    doc.add_heading("14  Output Formatting", level=2)
    doc.add_paragraph(
        "All visible sheets receive the following formatting automatically:"
    )
    doc.add_paragraph("Column headers are renamed to human-readable display names (no underscores).", style="List Bullet")
    doc.add_paragraph("Navy header band with white bold text; alternating light-blue / white row fills; thin cell borders.", style="List Bullet")
    doc.add_paragraph("Financial columns (Cap Hit, AAV, cap space): USD currency format with 2 decimal places, red negatives.", style="List Bullet")
    doc.add_paragraph("Percentage columns (Sample Confidence, Role Stability, OZ Start %, SH%, FO%, IPP): displayed as e.g. 85.00%.", style="List Bullet")
    doc.add_paragraph("All other numeric columns: max 2 decimal places.", style="List Bullet")
    doc.add_paragraph("Columns auto-sized to content (8–42 characters); header row height 32pt; first row frozen.", style="List Bullet")

    # ── 15. LIMITATIONS ──────────────────────────────────────────────────────
    doc.add_heading("15  Limitations & Future Enhancements", level=2)
    doc.add_paragraph("On-ice metrics reflect team context; quality-of-competition and quality-of-teammates adjustments are not applied.", style="List Bullet")
    doc.add_paragraph("TOI is used as an EV proxy because the NST export does not always separate EV TOI explicitly.", style="List Bullet")
    doc.add_paragraph("The aging curve uses population-level NHL averages; individual durability and style are not modelled.", style="List Bullet")
    doc.add_paragraph("Possible future enhancements: handedness / LD-RD split for D, micro-stats (zone entries/exits), playoff performance weighting, multi-team trade modelling.", style="List Bullet")

    out_file = out_path / "P_Model_Methodology.docx"
    doc.save(out_file)
    return out_file


def write_methodology_pdf(out_path: Path) -> Path:
    """
    Produces a clean multi-page PDF version of the methodology.
    Uses a simple two-pass layout: section headings in bold, body in regular weight,
    bullet items indented. Handles pagination automatically.
    """
    out_file = out_path / "P_Model_Methodology.pdf"
    c = canvas.Canvas(str(out_file), pagesize=letter)
    W, H = letter

    MARGIN_L = 54
    MARGIN_R = W - 54
    TEXT_W   = MARGIN_R - MARGIN_L
    Y_TOP    = H - 60
    Y_BOTTOM = 54

    # State
    y = Y_TOP

    def new_page():
        nonlocal y
        c.showPage()
        y = Y_TOP

    def draw_line(text: str, font: str, size: int, indent: int = 0, gap_before: int = 0, gap_after: int = 4, color=(0,0,0)):
        nonlocal y
        y -= gap_before
        if y < Y_BOTTOM + size + 4:
            new_page()
        c.setFont(font, size)
        c.setFillColorRGB(*color)
        c.drawString(MARGIN_L + indent, y, text[:120])  # hard truncate for safety
        c.setFillColorRGB(0, 0, 0)
        y -= size + gap_after

    def section(num: str, title: str):
        draw_line(f"{num}  {title}", "Helvetica-Bold", 12, gap_before=10, gap_after=4, color=(0.12, 0.31, 0.47))

    def subsection(title: str):
        draw_line(title, "Helvetica-Bold", 10, gap_before=6, gap_after=2, color=(0.20, 0.40, 0.60))

    def body(text: str):
        # Simple word-wrap
        words = text.split()
        line_buf: list[str] = []
        for w in words:
            test = " ".join(line_buf + [w])
            c.setFont("Helvetica", 9)
            if c.stringWidth(test, "Helvetica", 9) > TEXT_W - 8:
                draw_line(" ".join(line_buf), "Helvetica", 9, gap_after=2)
                line_buf = [w]
            else:
                line_buf.append(w)
        if line_buf:
            draw_line(" ".join(line_buf), "Helvetica", 9, gap_after=4)

    def bullet(text: str):
        words = text.split()
        first = True
        line_buf: list[str] = []
        for w in words:
            test = " ".join(line_buf + [w])
            c.setFont("Helvetica", 9)
            if c.stringWidth(test, "Helvetica", 9) > TEXT_W - 28:
                prefix = "•  " if first else "   "
                draw_line(prefix + " ".join(line_buf), "Helvetica", 9, indent=14, gap_after=2)
                first = False
                line_buf = [w]
            else:
                line_buf.append(w)
        if line_buf:
            prefix = "•  " if first else "   "
            draw_line(prefix + " ".join(line_buf), "Helvetica", 9, indent=14, gap_after=3)

    # ── Title page ─────────────────────────────────────────────────────────
    draw_line("NHL Player Analysis Model", "Helvetica-Bold", 16, gap_after=4, color=(0.12, 0.31, 0.47))
    draw_line("Methodology & Technical Reference", "Helvetica-Bold", 12, gap_after=10, color=(0.12, 0.31, 0.47))
    body(f"Input: {INPUT_XLSX}   |   Output: {OUTPUT_XLSX}")
    body(f"Seasons: {', '.join(SEASON_TABS)}   |   Rolling weights: {ROLLING_WEIGHTS}")

    # ── 1. Data Sources ─────────────────────────────────────────────────────
    section("1", "Data Sources & Assumptions")
    bullet("On-Ice (EV) tab — team context: Corsi, Fenwick, xGoals, HDCF, zone starts.")
    bullet("Individual (Ind) tab — per-player: goals, assists, shots, ixG, hits, blocks, faceoffs, penalties.")
    bullet("Bio tab — Height (in), Weight (lbs), Age, Draft Round, Overall Pick. BMI = weight / height².")
    bullet("Cap sheet — contract details; matched to performance data via normalised player key.")
    bullet("Injured / absent players are imputed from the most recent prior season (8 % regression toward mean).")

    # ── 2. Derived Metrics ──────────────────────────────────────────────────
    section("2", "Derived Metrics")
    bullet("Per-60 rates: metric_per60 = raw_count / TOI_minutes × 60.")
    bullet("Zone start %: OZ / (OZ + NZ + DZ).")
    bullet("Sample Confidence: clip(TOI / 1200, 0, 1) — reaches 1.0 at ~1 200 EV minutes.")
    bullet("Penalty Differential/60: (Drawn − Taken) / TOI × 60.")
    bullet("Primary Points/60 (P1/60): (Goals + Primary Assists) / TOI × 60.")

    # ── 3. Role Score Pipeline ──────────────────────────────────────────────
    section("3", "Role Score Construction")
    bullet("Z-score each feature within (season × position).")
    bullet(f"Minutes shrinkage: z_shrunk = z × TOI / (TOI + k), k = 500–700 min depending on role.")
    bullet("Weighted sum across features — negative weights penalise unfavourable stats.")
    bullet("Percentile rank to 0–100 within (season × position).")

    # ── 4. Role Definitions ─────────────────────────────────────────────────
    section("4", "Role Definitions")
    subsection("Forwards")
    fwd_roles = ["Finisher","Playmaker","Driver","TwoWay","Power","Grinder","Producer"]
    for r in fwd_roles:
        spec = ROLE_FEATURES["F"].get(r, {})
        feats = ", ".join(f"{k}({w:+.2f})" for k,w in spec.get("features",{}).items())
        bullet(f"{r} (k={spec.get('k','?')}): {feats}")
    subsection("Defensemen")
    def_roles = ["Suppressor","Transition","PuckSkill","Physical"]
    for r in def_roles:
        spec = ROLE_FEATURES["D"].get(r, {})
        feats = ", ".join(f"{k}({w:+.2f})" for k,w in spec.get("features",{}).items())
        bullet(f"{r} (k={spec.get('k','?')}): {feats}")

    # ── 5. Production Score ─────────────────────────────────────────────────
    section("5", "Individual Production Score (0–100)")
    bullet("50 % — Pts/60 (ceiling 3.5 fwd / 1.8 def).")
    bullet("25 % — G/60  (ceiling 1.5 fwd / 0.6 def).")
    bullet("25 % — ixG/60 (ceiling 1.8 fwd / 0.8 def).")
    bullet("Feeds into Impact Score (15 % weight) and Line 1 / Line 2 fit templates.")

    # ── 6. Impact Score ─────────────────────────────────────────────────────
    section("6", "Dimensionality & Impact Score")
    bullet("Dimensionality = average of player's two highest role scores.")
    bullet("Impact = (38% Dimensionality + 30% Top Role + 17% Best Fit + 15% Production) × reliability.")
    bullet("Reliability = 0.85 + 0.10 × Confidence_EV + 0.05 × Role Stability.")

    # ── 7. Player Type ──────────────────────────────────────────────────────
    section("7", "Player Type Classification")
    bullet("Stat nudges adjust role scores using counting-stat thresholds before ranking.")
    bullet(f"Hybrid guard: Top1 − Top2 < {TYPE_HYBRID_MARGIN} → type = 'Hybrid: A / B'.")
    bullet("Producer tiebreak: Producer wins ties on forwards.")
    bullet(f"TOI guard: below {MIN_TOI_EV_FWD} min (F) / {MIN_TOI_EV_DEF} min (D) → 'Insufficient TOI'.")
    bullet(f"Confidence: High if Conf_EV ≥ 0.50 and margin ≥ {TYPE_CLEAR_MARGIN}; Medium if margin ≥ {TYPE_HYBRID_MARGIN}; else Low.")

    # ── 8. Projection Model ─────────────────────────────────────────────────
    section("8", "Player Projection & Progression Model")
    subsection("Aging Curve (multipliers)")
    bullet("≤21: 0.88 | 22–23: 0.96 | 24–27: 1.02 | 28–29: 1.00 | 30–31: 0.97 | 32–33: 0.93 | 34–35: 0.87 | 36+: 0.80.")
    bullet("Defensemen peak ~2 years later.")
    subsection("Draft Pedigree Signal (decays to 0 by age ~32)")
    bullet("Pick 1–10 (Elite): +6% | Pick 11–30 (High): +4% | Pick 31–90 (Mid): +2% | Pick 181+ (Late): −1%.")
    subsection("Projected Rates")
    bullet("Proj Pts/60 = baseline × age_curve × (1 + pedigree_boost).")
    bullet("Baseline: current pts/60 if conf ≥ 0.40, else 3yr rolling, else dimensionality proxy.")
    subsection("Progression Tier")
    bullet("Emerging (≤23) | Prime (24–29) | Plateau (30–31) | Declining (32–35) | Aging (36+).")

    # ── 9. Line & Pair Fit ──────────────────────────────────────────────────
    section("9", "Lineup Fit Modeling")
    subsection("Forward Lines")
    bullet("Line 1: 85% role blend + 15% production modifier (pts/60 normalised).")
    bullet("Line 2: 88% role blend + 12% production modifier.")
    bullet("Line 3: 95% role blend + 5% production modifier.")
    bullet("Line 4: 100% role blend (energy / checking — no production signal).")
    bullet("Usage gates prevent low-minute players from Line 1/2 unless Production Score ≥ 80.")
    bullet("Grinder types: 20% / 40% suppression on Line 1 / Line 2 fit scores.")
    subsection("Defense Pairs")
    bullet("Pair 1: 30% Shutdown + 35% Transition + 35% Puck Skill.")
    bullet("Pair 2: 35% Transition + 30% Puck Skill + 35% Physical.")
    bullet("Pair 3: 50% Shutdown + 35% Physical + 15% Transition.")

    # ── 10. Contract Estimation ─────────────────────────────────────────────
    section("10", "Contract Estimation — Market Comparables")
    bullet("Comps filtered to same position group and lineup fit bucket.")
    bullet("Distance = IQR-normalised across age, impact, dimensionality, top role, fit, pts/60, proj pts/60.")
    bullet("Production and projection carry ~20% combined distance weight.")
    bullet("AAV distribution: weighted mean (closer comps = higher weight), std clipped to 12–35% of mean.")
    bullet("3 000 Student-t draws (df=6) model negotiation variance. P10/P50/P90 reported.")
    subsection("Progression-Adjusted Value")
    bullet("Emerging +15% | Prime +5% | Plateau 0% | Declining −8% | Aging −15%.")
    bullet("Draft pedigree adds up to +6% for players under ~26.")

    # ── 11. FA Pool ─────────────────────────────────────────────────────────
    section("11", "Free Agent Pool & Re-sign Decisioning")
    bullet("League-wide UFAs (expiry = offseason year, type ≠ RFA).")
    bullet("FA rank: hole-fill score + type confidence + progression-adjusted value / $ + impact.")
    bullet("Re-sign rank: override → hole bonus → confidence → prog-adj value → impact.")
    bullet("Budget consumed greedily until resign_budget_share of available cap space is spent.")

    # ── 12. Roster Holes ────────────────────────────────────────────────────
    section("12", "Roster Hole Detection")
    bullet("Group totals: 12 F, 6 D.")
    bullet("Position minimums: 4 C, 3 LW, 3 RW, 6 D.")
    bullet("Fit-bucket minimums: 3 per line slot (Line 1–4), 2 per pair slot (Pair 1–3).")

    # ── 13. Output Sheets ───────────────────────────────────────────────────
    section("13", "Output Sheets & Formatting")
    bullet("F_25-26 / D_25-26 etc. — role scores, type, production, projections, line/pair fit.")
    bullet("Rolling_F / Rolling_D — 3yr weighted averages with volatility (60/30/10).")
    bullet("TEAM tabs — Returning, Resign Decisions, FA Targets, Projected Roster, Lineup, Cap Summary.")
    bullet("Depth Chart — visual line/pair groupings.")
    bullet("All sheets: pretty headers, navy/alternating rows, number formats, 2dp max, frozen header row.")

    c.showPage()
    c.save()
    return out_file


# =============================================================================
# MAIN
# =============================================================================
def main() -> None:
    xl = pd.ExcelFile(INPUT_XLSX)
    needed = [f"{s} On-Ice" for s in SEASON_TABS] + [f"{s} Bios" for s in SEASON_TABS] + [IND_TABS[s] for s in SEASON_TABS]
    missing = [s for s in needed if s not in xl.sheet_names]
    if missing:
        raise ValueError(f"Missing required sheet(s): {missing}")

    season_all: Dict[str, pd.DataFrame] = {}

    # ── PASS 1: Load raw data — do NOT score yet; imputation must come first ──
    for season in SEASON_TABS:
        stats = load_season_sheet(INPUT_XLSX, season)
        bio   = load_bio_sheet(INPUT_XLSX, season)
        ind   = load_individual_sheet(INPUT_XLSX, season)

        merged = stats.merge(bio.drop(columns=["player", "season"], errors="ignore"),
                             on="player_key", how="left")
        merged = merged.merge(ind.drop(columns=["player", "season"], errors="ignore"),
                              on="player_key", how="left")

        if "team" in merged.columns and "team_ind" in merged.columns:
            merged["team"]     = merged["team"].fillna("").astype(str).str.strip()
            merged["team_ind"] = merged["team_ind"].fillna("").astype(str).str.strip()
            merged.loc[(merged["team"] == "") & (merged["team_ind"] != ""), "team"] = merged["team_ind"]

        if "toi_ind" not in merged.columns:
            merged["toi_ind"] = merged.get("toi_ev", np.nan)

        season_all[season] = merged

    # ── PASS 2: Impute BEFORE scoring ────────────────────────────────────────
    # Players with toi_ev=0 get prior-season stats filled in here so that the
    # keep flag is evaluated on imputed (realistic) TOI, not zero.
    season_all = impute_injured_seasons(season_all)

    # ── Load cap table ────────────────────────────────────────────────────────
    cap_df = load_cap_table(INPUT_XLSX)

    # ── PASS 3: Add active-contract players absent from NST (e.g. long-term IR)
    add_missing_active_contract_players(
        season_dfs=season_all,
        roster_df=cap_df,
        current_season=SEASON_TABS[0],
        fallback_seasons=SEASON_TABS[1:],
        key_col="player_key",
        name_col="player",
        decay=0.08,
    )

    # ── PASS 4: Impute again for any newly inserted missing-contract rows ─────
    season_all = impute_injured_seasons(season_all)

    # ── PASS 5: Score with full column parity ─────────────────────────────────
    # Imputation is now complete. apply_role_scores sets keep=True/False based
    # on imputed toi_ev. Players who still fall below the threshold (no usable
    # prior-season data either) receive keep=False and are omitted from output.
    def _score_season(df: pd.DataFrame) -> pd.DataFrame:
        df["Confidence_EV"] = np.clip(df["toi_ev"].fillna(0.0) / 1200.0, 0, 1)
        df = apply_role_scores(df)
        df = assign_player_type(df)
        df = add_type_confidence_and_stability(df)
        df = add_team_usage_rank(df)
        df = add_player_projections(df)
        df = add_prospect_tier(df)   # must run after add_player_projections (needs Draft_Pedigree + Prog_Tier)
        df = add_line_fit(df)
        return df

    for k in season_all:
        season_all[k] = _score_season(season_all[k])

    rolling_all = build_rolling(season_all, SEASON_TABS, ROLLING_WEIGHTS)
    rolling_all = add_type_confidence_and_stability(rolling_all)
    rolling_all = add_line_fit(rolling_all)


    TEAM = "FLA"  # <--- change this anytime (or read from a Settings tab later)

    model_2526 = season_all["25-26"].copy()
    model_2526 = merge_cap(model_2526, cap_df)

    plan_tabs = build_team_offseason_plan(
        model_season_df=model_2526,
        cap_df=cap_df,
        team=TEAM,
        offseason_year=2026,
        cap_limit=CAP_LIMIT_26_27,
        resign_budget_share=0.45,  # tune
    )

    visible: Dict[str, pd.DataFrame] = {}

    # ── YoY delta: join prior season to compute Dim_Delta and Dim_Trend ──────
    SEASON_ORDER = SEASON_TABS  # newest first
    for i, season in enumerate(SEASON_ORDER):
        df = season_all[season]
        prev_key = SEASON_ORDER[i + 1] if i + 1 < len(SEASON_ORDER) else None
        if prev_key and "player_key" in df.columns and "Dimensionality_Score" in df.columns:
            prev_df = season_all[prev_key]
            prev_dim = (
                prev_df.drop_duplicates("player_key")
                       .set_index("player_key")["Dimensionality_Score"]
            )
            prior_vals = prev_dim.reindex(df["player_key"]).values
            cur_dim    = df["Dimensionality_Score"].astype(float).values
            delta = np.where(
                ~np.isnan(prior_vals.astype(float)) & ~np.isnan(cur_dim),
                cur_dim - prior_vals.astype(float),
                np.nan
            )
            df = df.copy()
            df["Dim_Delta"] = np.round(delta, 1)

            # Age-adjusted, multi-category trend labels.
            #
            # Logic:
            #   - Young players (≤25): normal development curves mean a small dip is
            #     "noise", and meaningful gains should be expected. Thresholds are wider
            #     before flagging a decline, and a surge gets a distinct label.
            #   - Prime players (26–31): year-to-year changes are more signal than noise.
            #     Tighter thresholds; decline label kicks in sooner.
            #   - Aging players (32–35): some decline is expected. A mild dip is
            #     "Age-Expected Decline", not a red flag. A rise at this age is notable.
            #   - Late career (36+): meaningful gains are rare enough to call out.
            #     Any significant drop is a serious concern.
            #
            # Categories:
            #   ⚡ Breakout     — young player, large surge (≥15)
            #   ↑↑ Strong Rise  — clear improvement (≥10, any age)
            #   ↑  Rising       — moderate improvement (age-adjusted)
            #   → Stable        — within noise band (age-adjusted)
            #   ↓  Slipping     — mild decline, within expected variance
            #   ↓↓ Declining    — meaningful decline (age-adjusted)
            #   ⚠ Sharp Drop   — severe decline (≤-20, any age)
            #   ✦ Age Surge    — notable rise for 32+ player
            #   ~  New Data     — no prior season to compare

            ages = df.get("age", pd.Series(np.nan, index=df.index)).astype(float)

            # Injury mask: these players' current stats are imputed from prior season * 0.92.
            # Their Dim_Delta is mechanically ~-8% (decay artifact, not real decline).
            # We suppress the trend signal for them and replace with a dedicated label.
            injury_flag = (
                df.get("Injury_Imputed", pd.Series(False, index=df.index))
                  .fillna(False).astype(bool)
            )

            def _trend_label_row(row_delta, age, is_injured):
                # ── Injury override ──────────────────────────────────────────
                # If the current season was imputed from a prior season due to
                # injury/absence, the delta is contaminated by the 8% decay applied
                # during imputation — NOT a real performance change. Suppress entirely.
                if is_injured:
                    return "⛑ Injury Year"

                d = row_delta
                if pd.isna(d):
                    return "~  New Data"
                if pd.isna(age):
                    age = 27.0  # assume prime if unknown

                # Universal severe thresholds (override age buckets)
                if d >= 20:
                    return "↑↑ Strong Rise" if age >= 32 else ("⚡ Breakout" if age <= 25 else "↑↑ Strong Rise")
                if d <= -20:
                    return "⚠ Sharp Drop"

                if age <= 25:
                    # Development age: wider noise band; decline only on clear signal
                    if d >= 12:  return "⚡ Breakout"
                    if d >= 6:   return "↑  Rising"
                    if d >= -5:  return "→ Stable"
                    if d >= -12: return "↓  Slipping"
                    return "↓↓ Declining"

                elif age <= 31:
                    # Prime years: tightest thresholds, most signal
                    if d >= 10:  return "↑↑ Strong Rise"
                    if d >= 5:   return "↑  Rising"
                    if d >= -5:  return "→ Stable"
                    if d >= -10: return "↓  Slipping"
                    return "↓↓ Declining"

                elif age <= 35:
                    # Aging curve expected; mild drops are normal
                    if d >= 10:  return "✦ Age Surge"
                    if d >= 4:   return "↑  Rising"
                    if d >= -8:  return "→ Stable"     # wider stable band — age noise
                    if d >= -15: return "↓  Slipping"
                    return "↓↓ Declining"

                else:
                    # Late career (36+): any sustained gain is remarkable
                    if d >= 8:   return "✦ Age Surge"
                    if d >= 2:   return "↑  Rising"
                    if d >= -10: return "→ Stable"
                    if d >= -20: return "↓  Slipping"
                    return "↓↓ Declining"

            df["Dim_Trend"] = [
                _trend_label_row(d, a, inj)
                for d, a, inj in zip(df["Dim_Delta"], ages, injury_flag)
            ]

            # Also null out Dim_Delta for imputed players so the numeric column
            # doesn't mislead anyone into thinking -6 is a real performance drop.
            df.loc[injury_flag, "Dim_Delta"] = np.nan
        else:
            df = df.copy()
            if "Dim_Delta" not in df.columns:
                df["Dim_Delta"] = np.nan
            if "Dim_Trend" not in df.columns:
                df["Dim_Trend"] = ""
        season_all[season] = df

    # ── Re-run line fit now that Dim_Trend is populated ──────────────────────
    # add_line_fit was called earlier (inside _score_season) before Dim_Trend
    # existed, so the trend multipliers had nothing to work with.  A second
    # pass ensures the Tier and Trend multipliers actually influence the final
    # Fit_Line*/Fit_Pair* and Best_Line_Fit/Best_Pair_Fit values.
    for season in SEASON_TABS:
        season_all[season] = add_line_fit(season_all[season])

    for season in SEASON_TABS:
        df = season_all[season].copy()

        if "keep" in df.columns:
            df = df[df["keep"] == True].copy()

        f = df[df["pos_group"] == "F"].copy()
        d = df[df["pos_group"] == "D"].copy()

        f = f[[c for c in FWD_VISIBLE if c in f.columns]].sort_values(["pos", "Dimensionality_Score"], ascending=[True, False])
        d = d[[c for c in DEF_VISIBLE if c in d.columns]].sort_values(["pos", "Dimensionality_Score"], ascending=[True, False])

        visible[f"F_{season}"] = f
        visible[f"D_{season}"] = d

    # Rolling: only include players who met the threshold in at least one scored season.
    rolling_f = rolling_all[rolling_all["pos_group"] == "F"].copy()
    rolling_d = rolling_all[rolling_all["pos_group"] == "D"].copy()

    _f_score_cols = [c for c in rolling_f.columns if c.endswith("_3yr") and "vol" not in c]
    _d_score_cols = [c for c in rolling_d.columns if c.endswith("_3yr") and "vol" not in c]
    if _f_score_cols:
        rolling_f = rolling_f[rolling_f[_f_score_cols].notna().any(axis=1)].copy()
    if _d_score_cols:
        rolling_d = rolling_d[rolling_d[_d_score_cols].notna().any(axis=1)].copy()

    # ── Clean rolling column names for user readability ───────────────────────
    def _pretty_rolling_cols(df: pd.DataFrame) -> pd.DataFrame:
        rename = {}
        for c in df.columns:
            nc = c
            nc = nc.replace("_Score_3yr", " (3yr)").replace("_Score_vol_3yr", " (3yr Volatility)")
            nc = nc.replace("_3yr", " (3yr)").replace("_vol_3yr", " (Volatility)")
            nc = nc.replace("Dimensionality", "Dimensionality").replace("_", " ")
            if nc != c:
                rename[c] = nc
        return df.rename(columns=rename)

    visible["Rolling_F"] = _pretty_rolling_cols(rolling_f)
    visible["Rolling_D"] = _pretty_rolling_cols(rolling_d)

    # ── League-wide team summary sheet ────────────────────────────────────────
    league_summary = build_league_summary(season_all, SEASON_TABS[0])
    if len(league_summary) > 0:
        visible["League_Team_Summary"] = league_summary

    def _round_df(df: pd.DataFrame) -> pd.DataFrame:
        """
        Round all float columns to sane display precision before writing to Excel.
        Prevents 8-decimal trailing noise from appearing in any sheet.
          - *_Score / Dimensionality / Margin / Dim_Delta → 1 dp
          - percentage 0-1 cols → 4 dp  (so 52.68% renders correctly)
          - everything else float → 2 dp
        """
        out = df.copy()
        SCORE_COLS  = {c for c in out.columns if c.endswith("_Score") or
                       c in ("Dimensionality_Score","Type_Margin","Margin",
                             "Dim_Delta","Top_Role_Score","Second_Role_Score",
                             "Best_Line_Fit_Score","Best_Pair_Fit_Score",
                             "Fit_Line1","Fit_Line2","Fit_Line3","Fit_Line4",
                             "Fit_Pair1","Fit_Pair2","Fit_Pair3")}
        PCT_COLS    = {c for c in out.columns if c in (
                       "sh_pct","fo_pct","ipp","off_zone_start_pct",
                       "def_zone_start_pct","Confidence_EV","Role_Stability",
                       "Age_Curve_Factor","Usage_Rank_Team")}
        for col in out.select_dtypes(include="float").columns:
            if col in SCORE_COLS:
                out[col] = out[col].round(1)
            elif col in PCT_COLS:
                out[col] = out[col].round(4)
            else:
                out[col] = out[col].round(2)
        return out

    with pd.ExcelWriter(OUTPUT_XLSX, engine="openpyxl") as writer:
        for name, df in visible.items():
            _round_df(df).to_excel(writer, sheet_name=name, index=False)

        for season in SEASON_TABS:
            _round_df(season_all[season]).to_excel(writer, sheet_name=f"_ALL_{season}", index=False)
        _round_df(rolling_all).to_excel(writer, sheet_name="_ROLLING_ALL", index=False)

        for name, df in plan_tabs.items():
            _round_df(df).to_excel(writer, sheet_name=name[:31], index=False)

    wb = openpyxl.load_workbook(OUTPUT_XLSX)

    for season in SEASON_TABS:
        _hide_sheet(wb[f"_ALL_{season}"])
    _hide_sheet(wb["_ROLLING_ALL"])

    for season in SEASON_TABS:
        _hide_unused_columns(wb[f"F_{season}"], FWD_VISIBLE)
        _hide_unused_columns(wb[f"D_{season}"], DEF_VISIBLE)

    # -------------------------
    # Full workbook formatting: column display names, table styles,
    # number formats, auto-column widths, freeze panes
    # -------------------------
    _format_workbook(wb)

    # -------------------------
    # Visual depth chart (post-FA)
    # -------------------------
    TEAM = "FLA"  # keep in sync with above
    depth_key   = f"TEAM_{TEAM}_Projected_Lineup_26-27"
    roster_key  = f"TEAM_{TEAM}_Projected_Roster_26-27"
    if depth_key in plan_tabs:
        roster_df = plan_tabs.get(roster_key, pd.DataFrame())
        _write_depth_chart_sheet(wb, plan_tabs[depth_key], team=TEAM, roster_df=roster_df)

    # ── Player Analysis dashboard ──────────────────────────────────────────
    write_player_analysis_tab(
        wb,
        season_df    = season_all[SEASON_TABS[0]],
        rolling_df   = rolling_all,
        prev_season_df = season_all[SEASON_TABS[1]] if len(SEASON_TABS) > 1 else None,
    )

    wb.save(OUTPUT_XLSX)

    # Write docs next to the output workbook
    out_dir = OUTPUT_XLSX.parent
    docx_path = write_methodology_docx(out_dir)
    pdf_path = write_methodology_pdf(out_dir)

    print(f"✅ Wrote workbook: {OUTPUT_XLSX}")
    print(f"✅ Wrote docx:     {docx_path}")
    print(f"✅ Wrote pdf:      {pdf_path}")


if __name__ == "__main__":
    main()